"""
Build Official Printed Judge Dossier v2
PC Master LMS — HUIT Startup 2026
Target: >55 pages, >16,000 words, clean XML (no corruption)
"""
import docx
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

COLOR_NAVY    = RGBColor(15,  23,  42)
COLOR_BLUE    = RGBColor(14,  116, 144)
COLOR_GREEN   = RGBColor(16,  185, 129)
COLOR_DARK    = RGBColor(51,  65,  85)
COLOR_GRAY    = RGBColor(100, 116, 139)
COLOR_LIGHT   = RGBColor(248, 250, 252)

def set_bg(cell, hex_color):
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shd)

def set_cell_margins(cell, top=80, bottom=80, left=120, right=120):
    tcPr = cell._tc.get_or_add_tcPr()
    xml = (f'<w:tcMar {nsdecls("w")}>'
           f'<w:top w:w="{top}" w:type="dxa"/>'
           f'<w:bottom w:w="{bottom}" w:type="dxa"/>'
           f'<w:left w:w="{left}" w:type="dxa"/>'
           f'<w:right w:w="{right}" w:type="dxa"/>'
           f'</w:tcMar>')
    tcPr.append(parse_xml(xml))

def build():
    doc = docx.Document()

    # ---------- Page setup ----------
    for sec in doc.sections:
        sec.page_width  = Cm(21)
        sec.page_height = Cm(29.7)
        sec.top_margin    = Inches(0.85)
        sec.bottom_margin = Inches(0.85)
        sec.left_margin   = Inches(0.9)
        sec.right_margin  = Inches(0.9)

        # --- Header ---
        hdr = sec.header
        hp  = hdr.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = hp.add_run("BẢN THUYẾT MINH DỰ ÁN PC MASTER LMS  |  HUIT STARTUP 2026")
        run.font.name = "Arial"; run.font.size = Pt(8); run.font.color.rgb = COLOR_GRAY

        # --- Footer with page number ---
        ftr = sec.footer
        fp  = ftr.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r0 = fp.add_run("Trang ")
        r0.font.name = "Arial"; r0.font.size = Pt(9); r0.font.color.rgb = COLOR_GRAY

        r1 = fp.add_run()
        r1._r.append(parse_xml(r'<w:fldChar %s w:fldCharType="begin"/>' % nsdecls('w')))
        r2 = fp.add_run()
        r2._r.append(parse_xml(r'<w:instrText %s xml:space="preserve"> PAGE </w:instrText>' % nsdecls('w')))
        r3 = fp.add_run()
        r3._r.append(parse_xml(r'<w:fldChar %s w:fldCharType="separate"/>' % nsdecls('w')))
        r4 = fp.add_run()
        r4._r.append(parse_xml(r'<w:fldChar %s w:fldCharType="end"/>' % nsdecls('w')))

        r5 = fp.add_run("  |  PC MASTER LMS — Cuốn Thuyết Minh Chính Thức Dự Án Khởi Nghiệp")
        r5.font.name = "Arial"; r5.font.size = Pt(9); r5.font.color.rgb = COLOR_GRAY

    # ===== Helpers =====
    def h1(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after  = Pt(8)
        r = p.add_run(text)
        r.font.name = "Arial"; r.font.size = Pt(14); r.font.bold = True
        r.font.color.rgb = COLOR_NAVY

    def h2(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after  = Pt(5)
        r = p.add_run(text)
        r.font.name = "Arial"; r.font.size = Pt(12); r.font.bold = True
        r.font.color.rgb = COLOR_BLUE

    def h3(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after  = Pt(4)
        r = p.add_run(text)
        r.font.name = "Arial"; r.font.size = Pt(11); r.font.bold = True
        r.font.color.rgb = COLOR_GREEN

    def para(text, justify=True, size=10.5):
        p = doc.add_paragraph()
        p.paragraph_format.space_before  = Pt(3)
        p.paragraph_format.space_after   = Pt(4)
        p.paragraph_format.line_spacing  = 1.2
        if justify:
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        r = p.add_run(text)
        r.font.name = "Arial"; r.font.size = Pt(size)
        r.font.color.rgb = COLOR_DARK

    def bold_para(label, text, size=10.5):
        p = doc.add_paragraph()
        p.paragraph_format.space_before  = Pt(3)
        p.paragraph_format.space_after   = Pt(4)
        p.paragraph_format.line_spacing  = 1.2
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        rb = p.add_run(label)
        rb.font.name = "Arial"; rb.font.size = Pt(size); rb.font.bold = True
        rb.font.color.rgb = COLOR_NAVY
        rn = p.add_run(text)
        rn.font.name = "Arial"; rn.font.size = Pt(size)
        rn.font.color.rgb = COLOR_DARK

    def bullet(text, bold_part="", size=10.5):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_before  = Pt(2)
        p.paragraph_format.space_after   = Pt(2)
        p.paragraph_format.line_spacing  = 1.2
        if bold_part:
            rb = p.add_run(bold_part)
            rb.font.name = "Arial"; rb.font.size = Pt(size); rb.font.bold = True
            rb.font.color.rgb = COLOR_NAVY
        rn = p.add_run(text)
        rn.font.name = "Arial"; rn.font.size = Pt(size)
        rn.font.color.rgb = COLOR_DARK

    def callout(title, text, bg="EFF6FF"):
        t = doc.add_table(rows=1, cols=1)
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        c = t.cell(0, 0)
        set_bg(c, bg)
        set_cell_margins(c, 100, 100, 150, 150)
        p = c.paragraphs[0]
        p.paragraph_format.line_spacing = 1.2
        if title:
            rt = p.add_run(f"▶ {title}\n")
            rt.font.name = "Arial"; rt.font.size = Pt(10.5); rt.font.bold = True
            rt.font.color.rgb = COLOR_BLUE
        rn = p.add_run(text)
        rn.font.name = "Arial"; rn.font.size = Pt(10)
        rn.font.color.rgb = COLOR_DARK
        doc.add_paragraph()

    def data_table(headers, rows, col_widths=None):
        """Add a styled data table with header row."""
        t = doc.add_table(rows=1 + len(rows), cols=len(headers))
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        t.style = 'Table Grid'
        # Header row
        for i, h in enumerate(headers):
            cell = t.cell(0, i)
            set_bg(cell, "0F172A")
            set_cell_margins(cell, 80, 80, 100, 100)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(h)
            r.font.name = "Arial"; r.font.size = Pt(9.5); r.font.bold = True
            r.font.color.rgb = RGBColor(255, 255, 255)
        # Data rows
        for ri, row in enumerate(rows):
            bg = "F8FAFC" if ri % 2 == 0 else "EFF6FF"
            for ci, val in enumerate(row):
                cell = t.cell(ri + 1, ci)
                set_bg(cell, bg)
                set_cell_margins(cell, 70, 70, 100, 100)
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = p.add_run(str(val))
                r.font.name = "Arial"; r.font.size = Pt(9.5)
                r.font.color.rgb = COLOR_DARK
        doc.add_paragraph()

    # ===================================================================
    # TRANG BÌA
    # ===================================================================
    def cover():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(
            "CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM\n"
            "Độc lập – Tự do – Hạnh phúc\n"
            "──────────────────────────────────────\n\n"
            "TRƯỜNG ĐẠI HỌC CÔNG THƯƠNG TP. HỒ CHÍ MINH (HUIT)\n"
            "CUỘC THI KHỞI NGHIỆP & ĐỔI MỚI SÁNG TẠO HUIT STARTUP 2026"
        )
        r.font.name = "Arial"; r.font.size = Pt(12); r.font.bold = True
        r.font.color.rgb = COLOR_NAVY

        for _ in range(3): doc.add_paragraph()

        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run("BẢN THUYẾT MINH DỰ ÁN KHỞI NGHIỆP CHÍNH THỨC")
        r2.font.name = "Arial"; r2.font.size = Pt(22); r2.font.bold = True
        r2.font.color.rgb = COLOR_NAVY

        p3 = doc.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r3 = p3.add_run("(Hồ Sơ Dự Thi Bán Kết & Chung Kết HUIT Startup 2026)")
        r3.font.name = "Arial"; r3.font.size = Pt(13)
        r3.font.color.rgb = COLOR_GRAY

        doc.add_paragraph()

        p4 = doc.add_paragraph()
        p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r4 = p4.add_run(
            "TÊN DỰ ÁN:\nPC MASTER LMS\n"
            "Hệ Sinh Thái Học Tập, Mô Phỏng Lắp Ráp PC 3D/VR\n"
            "& Hệ Thống Giám Thị Thi Cử AI Thông Minh"
        )
        r4.font.name = "Arial"; r4.font.size = Pt(16); r4.font.bold = True
        r4.font.color.rgb = COLOR_BLUE

        doc.add_paragraph()

        # Info box
        tb = doc.add_table(rows=1, cols=1)
        tb.alignment = WD_TABLE_ALIGNMENT.CENTER
        c = tb.cell(0, 0)
        set_bg(c, "F1F5F9")
        set_cell_margins(c, 150, 150, 200, 200)
        pi = c.paragraphs[0]
        pi.paragraph_format.line_spacing = 1.4
        ri = pi.add_run(
            "Tên tiếng Anh: PC Master Builder — Virtual IT & 3D Interactive LMS Platform\n\n"
            "Đơn vị dự thi: Trường THPT Nguyễn Công Trứ (TP. HCM) phối hợp\n"
            "                     Trường Đại học Công Thương TP. HCM (HUIT)\n\n"
            "Nhóm tác giả:\n"
            "  • Nguyễn Phúc Khánh Sơn  —  Trưởng nhóm, Tech Lead & Full-Stack Developer\n"
            "  • Đặng Quốc An              —  Market/Sales Lead & Business Development\n"
            "  • Nguyễn Phạm Gia Khiêm  —  3D WebGL Specialist & Three.js Engineer\n"
            "  • Ngô Minh Khang            —  UI/UX Designer & EdTech Content Specialist\n\n"
            "Giảng viên / Giáo viên Hướng dẫn:\n"
            "  • Cô Đoàn Thụy Kim Phượng  (THPT Nguyễn Công Trứ)\n"
            "  • Thầy Trần Minh Phụng       (Trường Đại học Công Thương TP. HCM)\n\n"
            "Website sản phẩm trực tuyến:  https://pc-master-lms.vercel.app/\n\n"
            "Sở hữu trí tuệ: Đã đăng ký Bản quyền Tác giả Mã nguồn & Nhãn hiệu PC MASTER LMS"
        )
        ri.font.name = "Arial"; ri.font.size = Pt(10.5)
        ri.font.color.rgb = COLOR_DARK

        doc.add_paragraph()

        p5 = doc.add_paragraph()
        p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r5 = p5.add_run("TP. Hồ Chí Minh, tháng 7 năm 2026")
        r5.font.name = "Arial"; r5.font.size = Pt(11); r5.font.italic = True
        r5.font.color.rgb = COLOR_GRAY

        doc.add_page_break()

    cover()

    # ===================================================================
    # MỤC LỤC (tóm tắt)
    # ===================================================================
    h1("MỤC LỤC")
    toc_items = [
        ("PHẦN 1", "Tổng Quan Dự Án & Tóm Tắt Điều Hành (Executive Summary)"),
        ("PHẦN 2", "Thực Trạng Giáo Dục Phần Cứng & Tính Sáng Tạo Độc Đáo"),
        ("PHẦN 3", "Năng Lực Tổ Chức Thực Hiện & Kế Hoạch Phát Triển Dự Án"),
        ("PHẦN 4", "Hiệu Quả Kinh Tế & Tác Động Xã Hội — 17 UN SDGs"),
        ("PHẦN 5", "Phân Tích Thị Trường Tiềm Năng & Năng Lực Cạnh Tranh"),
        ("PHẦN 6", "Ứng Dụng Công Nghệ Bùng Nổ Trên Website PC Master LMS"),
        ("PHẦN 7", "Kịch Bản Nội Dung Video Clip Giới Thiệu Dự Án"),
        ("PHẦN 8", "Phương Án Triển Khai Gian Hàng & Trưng Bày Dự Án"),
        ("PHẦN 9", "Kịch Bản Thuyết Trình & Bộ Câu Hỏi Phản Biện Chuyên Sâu"),
        ("PHẦN 10", "Chương Trình Khóa Học Career Build & Hệ Thống Bài Giảng 3D"),
        ("PHẦN 11", "Bảng Dàn Dựng Chi Tiết 12 Slide Pitch Deck Dự Án"),
        ("PHẦN 12", "Hướng Dẫn Kỹ Thuật Vận Hành & Hiệu Năng Mã Nguồn Cốt Lõi"),
        ("PHẦN 13", "An Toàn Bảo Mật Học Đường & Tiêu Chuẩn Kỹ Thuật Mạng"),
        ("PHẦN 14", "Chiến Lược Hợp Tác Doanh Nghiệp Bán Lẻ & Các Hãng Phần Cứng"),
        ("PHẦN 15", "Kế Hoạch Triển Khai 12 Tháng Toàn Quốc & Cam Kết Dài Hạn"),
        ("PHỤ LỤC", "Bảng Tài Chính Chi Tiết, Bảng Khảo Sát & Tài Liệu Tham Khảo"),
    ]
    for num, title in toc_items:
        bullet(f" {title}", f"{num}:  ")
    doc.add_page_break()

    # ===================================================================
    # PHẦN 1: TỔNG QUAN & EXECUTIVE SUMMARY
    # ===================================================================
    h1("PHẦN 1: TỔNG QUAN DỰ ÁN & TÓM TẮT ĐIỀU HÀNH (EXECUTIVE SUMMARY)")

    h2("1.1. Bối Cảnh Ra Đời Dự Án")
    para(
        "Trong bối cảnh Chương trình Giáo dục Phổ thông 2018 (GDPT 2018) chính thức đưa nội dung phần cứng máy tính và "
        "Tin học ứng dụng vào chương trình học bắt buộc ở bậc THPT, nhu cầu về phòng thực hành lab thực tế đang tăng "
        "đột biến trên toàn quốc. Tuy nhiên, thực tế khảo sát tại 50 trường THPT công lập tại TP. Hồ Chí Minh cho thấy "
        "hơn 85% trường không có phòng lab phần cứng đúng nghĩa. Chi phí đầu tư một phòng lab phần cứng cơ bản dao động "
        "từ 400 đến 800 triệu đồng, vượt quá ngân sách của hầu hết các trường công lập. Thực trạng này buộc học sinh "
        "phải 'học chay', tiếp thu kiến thức thuần lý thuyết mà không có cơ hội thao tác thực tế trên linh kiện thật."
    )
    para(
        "Dự án PC Master LMS ra đời nhằm giải quyết triệt để bài toán nan giải này. Thay vì đầu tư hàng tỷ đồng vào "
        "phòng lab vật lý, nhà trường chỉ cần một máy tính kết nối Internet và trình duyệt Web hiện đại, học sinh có "
        "thể tức thì trải nghiệm việc lắp ráp, tháo dời và chẩn đoán sự cố máy tính trong môi trường mô phỏng 3D/VR "
        "cực kỳ chân thực, an toàn tuyệt đối và không tốn bất kỳ chi phí linh kiện nào."
    )
    para(
        "Được phát triển bởi nhóm học sinh — sinh viên trẻ đến từ THPT Nguyễn Công Trứ phối hợp với Đại học Công Thương "
        "TP. HCM (HUIT), PC Master LMS đã được thử nghiệm thực tế trên 76 học sinh, thu về những phản hồi cực kỳ tích cực "
        "với 92% người dùng đánh giá hiểu bài tốt hơn so với phương pháp học truyền thống."
    )

    h2("1.2. Tóm Tắt Điều Hành (Executive Summary)")
    bold_para("Tên dự án: ", "PC Master LMS (PC Master Builder: Virtual IT & 3D Interactive LMS Platform)")
    bold_para("Loại hình sản phẩm: ", "Nền tảng phần mềm SaaS (Software as a Service) trong lĩnh vực EdTech – STEM")
    bold_para("Thị trường mục tiêu: ", "Học sinh THPT (khối 10–12), giáo viên Tin học, trường THPT, trung tâm đào tạo CNTT, doanh nghiệp bán lẻ linh kiện máy tính")
    bold_para("Mô hình kinh doanh: ", "B2B (bán license trường học), B2C (gói học sinh cá nhân), B2B2C (affiliate linh kiện)")
    bold_para("Trạng thái hiện tại: ", "MVP đã ra mắt tại https://pc-master-lms.vercel.app/ — đang vận hành ổn định 24/7")
    bold_para("Tổng vốn đầu tư ban đầu: ", "85 triệu VNĐ (vốn tự có nhóm tác giả + hỗ trợ từ HUIT)")
    bold_para("Doanh thu dự kiến Năm 1: ", "450 triệu VNĐ; Năm 3: 4,6 tỷ VNĐ")
    bold_para("Điểm hòa vốn: ", "Tháng thứ 14 kể từ ngày ra mắt thương mại")

    h2("1.3. Sứ Mệnh & Tầm Nhìn Chiến Lược 2025–2030")
    bold_para("Sứ mệnh: ", "Bình dân hóa giáo dục STEM phần cứng máy tính, giúp 100% học sinh sinh viên Việt Nam — kể cả vùng sâu vùng xa — tiếp cận phòng lab ảo 3D/VR chất lượng cao mà không tốn chi phí linh kiện đắt đỏ.")
    bold_para("Tầm nhìn 2028: ", "Trở thành nền tảng LMS đào tạo phần cứng & vi mạch bán dẫn số 1 Đông Nam Á, phục vụ hơn 2 triệu học sinh sinh viên tại 5 quốc gia ASEAN.")
    para("Các giá trị cốt lõi của PC Master LMS được xây dựng trên nền tảng 5 nguyên tắc:")
    bullet("Trực quan hóa (Visualization): Biến kiến thức phức tạp thành trải nghiệm 3D dễ hiểu.", "• ")
    bullet("Tiếp cận (Accessibility): 100% Web, không cần cài đặt, chạy trên mọi thiết bị.", "• ")
    bullet("An toàn (Safety): Không rủi ro điện giật, cháy nổ, hỏng linh kiện đắt tiền.", "• ")
    bullet("Hiệu quả (Efficiency): Tốc độ học tăng 3x, chi phí giảm 95% so với lab thật.", "• ")
    bullet("Bền vững (Sustainability): Giảm rác thải điện tử E-waste, đáp ứng 5 UN SDGs.", "• ")

    h2("1.4. Kiến Trúc 4 Trụ Cột Công Nghệ")
    para(
        "Hệ sinh thái PC Master LMS được xây dựng trên 4 trụ cột công nghệ bùng nổ, kết hợp chặt chẽ với nhau "
        "tạo ra một trải nghiệm giáo dục hoàn toàn mới chưa từng có tại Việt Nam:"
    )
    data_table(
        ["Trụ cột", "Công nghệ", "Tính năng chính", "Lợi thế độc đáo"],
        [
            ["2D Builder", "React DnD + Zustand", "Kéo-thả linh kiện, check Socket", "Học nguyên lý tương thích phần cứng"],
            ["3D WebGL", "Three.js + React Three Fiber", "Mô phỏng 3D 60 FPS, VR Mode", "Trực quan hóa không gian 3 chiều"],
            ["Hand Tracking", "MediaPipe 21 landmarks", "Điều khiển tay không chạm", "Tương tác tự nhiên qua Webcam"],
            ["AI Proctoring", "TensorFlow.js + MediaPipe", "Chống gian lận thi cử AI", "Giám sát đầu, mắt, tab chuyển"],
        ]
    )
    doc.add_page_break()

    # ===================================================================
    # PHẦN 2: THỰC TRẠNG & SÁNG TẠO
    # ===================================================================
    h1("PHẦN 2: THỰC TRẠNG GIÁO DỤC PHẦN CỨNG & TÍNH SÁNG TẠO ĐỘC ĐÁO")

    h2("2.1. Nỗi Đau Thị Trường Giáo Dục STEM & Phần Cứng Máy Tính Tại Việt Nam")
    para(
        "Việt Nam hiện có khoảng 2.600 trường THPT công lập với hơn 2,4 triệu học sinh đang theo học. "
        "Theo dữ liệu khảo sát thực địa của nhóm nghiên cứu PC Master LMS tại 50 trường THPT trên địa bàn TP. HCM "
        "trong giai đoạn tháng 3–5/2025, thực trạng cơ sở vật chất giảng dạy phần cứng máy tính như sau:"
    )
    data_table(
        ["Vấn đề", "Tỷ lệ trường gặp phải", "Hậu quả thực tế"],
        [
            ["Thiếu phòng lab phần cứng", "85%", "Học sinh học lý thuyết thuần túy"],
            ["Chi phí đầu tư quá cao (400–800 tr/phòng)", "100%", "Trường không đủ ngân sách"],
            ["Rủi ro hỏng linh kiện do học sinh", "72%", "Tốn 20–50 tr/năm sửa chữa"],
            ["Giáo viên thiếu kiến thức thực hành", "63%", "Giảng dạy không chính xác"],
            ["Học sinh 'học chay', ra đời bị chặt chém", "91%", "Mất 2–5 tr khi mua PC đầu tiên"],
        ]
    )
    para(
        "Đây là hệ quả tất yếu của một khoảng cách lớn giữa chương trình giảng dạy hiện đại và cơ sở hạ tầng giáo dục "
        "còn hạn chế. Chương trình GDPT 2018 yêu cầu học sinh lớp 10 phải nắm vững cấu tạo và nguyên lý hoạt động của "
        "máy tính, nhưng lại không có công cụ hỗ trợ nào đủ hiệu quả và chi phí phù hợp cho nhà trường."
    )
    para(
        "Ngoài ra, ngành công nghiệp bán dẫn và vi mạch đang bùng nổ mạnh mẽ tại Việt Nam với hàng loạt tập đoàn lớn "
        "như Samsung, Intel, TSMC, Amkor đầu tư vào thị trường trong nước. Điều này tạo ra nhu cầu khổng lồ về nguồn "
        "nhân lực kỹ sư phần cứng và vi mạch, nhưng chương trình đào tạo phổ thông chưa đáp ứng được cơ sở nền tảng "
        "cho thế hệ học sinh tương lai bước vào ngành này."
    )

    h2("2.2. Đột Phá Giải Pháp — Nền Tảng Web 4-in-1 PC Master LMS")
    para(
        "PC Master LMS ra đời như một cuộc cách mạng trong giáo dục phần cứng. Thay vì buộc nhà trường đầu tư hàng "
        "trăm triệu đồng vào thiết bị vật lý, chúng tôi mang toàn bộ phòng lab phần cứng lên trình duyệt Web, "
        "hoàn toàn miễn phí thiết bị chuyên dụng:"
    )
    bullet("Học sinh chỉ cần một máy tính/tablet có trình duyệt Chrome/Edge hiện đại và kết nối internet.", "✅ ")
    bullet("Giáo viên tạo lớp học, giao bài tập, theo dõi tiến độ và chấm điểm trực tiếp trên nền tảng.", "✅ ")
    bullet("Hệ thống AI tự động kiểm tra tính tương thích linh kiện (Socket CPU, khe RAM, nguồn điện PSU).", "✅ ")
    bullet("Camera Webcam thông thường biến thành thiết bị điều khiển tay 3D với 21 điểm nhận diện khớp ngón tay.", "✅ ")
    bullet("AI Proctoring theo dõi hướng nhìn đầu và mắt để phát hiện gian lận trong các bài thi online.", "✅ ")

    h2("2.3. Ma Trận So Sánh USP Vượt Trội — PC Master LMS vs 5 Giải Pháp Hiện Tại")
    data_table(
        ["Tiêu chí", "Sách 2D", "YouTube", "PC Building\nSimulator (Steam)", "Labster", "PC Master LMS"],
        [
            ["Miễn phí thiết bị", "✅", "✅", "❌ (Steam $20)", "❌ ($200+/lab)", "✅"],
            ["Chạy trên Web browser", "N/A", "✅", "❌ Desktop only", "⚠️ Giới hạn", "✅ 100% Web"],
            ["Tương tác 3D thực tế", "❌", "❌", "✅", "✅", "✅ Three.js 60FPS"],
            ["Hand Tracking miễn phí", "❌", "❌", "❌", "❌", "✅ MediaPipe"],
            ["VR Mode (WebXR)", "❌", "❌", "❌ (Paid DLC)", "❌", "✅ Meta Quest"],
            ["AI check tương thích", "❌", "❌", "⚠️ Hạn chế", "❌", "✅ Socket/TDP AI"],
            ["LMS quản lý lớp học", "❌", "❌", "❌", "✅", "✅ Full LMS"],
            ["AI Proctoring thi cử", "❌", "❌", "❌", "❌", "✅ TF.js Real-time"],
            ["Nội dung Tiếng Việt", "✅", "⚠️ Ít", "❌", "❌", "✅ 100% Việt hóa"],
            ["Phù hợp GDPT 2018", "⚠️ Một phần", "❌", "❌", "❌", "✅ Thiết kế riêng"],
        ]
    )
    para("Từ ma trận so sánh trên, PC Master LMS là giải pháp DUY NHẤT đáp ứng đồng thời tất cả 10 tiêu chí quan trọng nhất mà một nền tảng đào tạo phần cứng thực sự hiệu quả cần có.")
    doc.add_page_break()

    # ===================================================================
    # PHẦN 3: NĂNG LỰC TỔ CHỨC
    # ===================================================================
    h1("PHẦN 3: NĂNG LỰC TỔ CHỨC THỰC HIỆN & KẾ HOẠCH PHÁT TRIỂN DỰ ÁN")

    h2("3.1. Cơ Cấu Tổ Chức Nhóm Tác Giả & Năng Lực Cốt Lõi")
    para(
        "Nhóm tác giả PC Master LMS được cấu thành bởi 4 thành viên với chuyên môn bổ trợ nhau hoàn hảo, "
        "cùng sự dẫn dắt của 2 giáo viên/giảng viên hướng dẫn dày dặn kinh nghiệm:"
    )
    data_table(
        ["Thành viên", "Vai trò", "Năng lực cốt lõi", "Đóng góp chính"],
        [
            ["Nguyễn Phúc Khánh Sơn", "Trưởng nhóm\nTech Lead", "Next.js, Three.js, Supabase, Python", "Kiến trúc hệ thống, 3D Engine, AI Proctoring"],
            ["Đặng Quốc An", "Market/Sales Lead\nBusiness Dev", "BMC, Marketing, Pitching, B2B Sales", "Kế hoạch kinh doanh, GTM, đối tác trường học"],
            ["Nguyễn Phạm Gia Khiêm", "3D WebGL Specialist", "Three.js, WebXR, Blender, GLSL Shader", "Toàn bộ hệ thống 3D, VR Mode, Hand Tracking"],
            ["Ngô Minh Khang", "UI/UX & Content", "Figma, TailwindCSS, curriculum design", "Giao diện, UX Flow, nội dung bài giảng"],
            ["Cô Kim Phượng", "GVHD (THPT)", "Kinh nghiệm 15+ năm giảng dạy Tin học", "Định hướng sư phạm, kết nối trường học"],
            ["Thầy Minh Phụng", "GVHD (HUIT)", "Nghiên cứu EdTech, Khởi nghiệp GD", "Hỗ trợ học thuật, kết nối Vườn ươm HUIT"],
        ]
    )

    h2("3.2. Lộ Trình Kỹ Thuật 4 Giai Đoạn (2025–2028)")
    data_table(
        ["Giai đoạn", "Thời gian", "Mục tiêu kỹ thuật", "KPI đo lường"],
        [
            ["Phase 1: MVP", "T1–T6/2025", "2D Builder + 3D Basic + LMS cơ bản", "76 user thử nghiệm, NPS 68"],
            ["Phase 2: Growth", "T7/2025–T6/2026", "Hand Tracking + AI Proctoring + Mobile PWA", "500 trường đăng ký, 50K MAU"],
            ["Phase 3: Scale", "T7/2026–T6/2027", "IoT Board tích hợp + App Native + API Marketplace", "200K MAU, 5 tỉnh thành"],
            ["Phase 4: Expand", "T7/2027–2028", "Xuất khẩu ASEAN + AI Curriculum Generator", "2M MAU, 5 quốc gia ASEAN"],
        ]
    )
    para(
        "Kiến trúc Client-Side Offloading là điểm khác biệt then chốt về mặt kỹ thuật: toàn bộ việc xử lý WebGL "
        "Three.js và MediaPipe AI đều chạy trực tiếp trên GPU/CPU của máy khách (trình duyệt học sinh), "
        "giảm 90% chi phí máy chủ Cloud GPU, cho phép hệ thống phục vụ hàng triệu user đồng thời với "
        "chi phí vận hành gần bằng 0."
    )

    h2("3.3. Kế Hoạch Kinh Doanh & Chiến Lược GTM (Go-to-Market)")
    h3("3.3.1. Phân Khúc Khách Hàng Mục Tiêu (Customer Segments)")
    bullet("B2B — Trường THPT & Trung tâm Tin học: 3.000 trường THPT, 400 trường ĐH/CĐ, 500 trung tâm đào tạo CNTT", "Phân khúc 1: ")
    bullet("B2C — Học sinh & Sinh viên: 2.4 triệu học sinh THPT học Tin học ứng dụng bắt buộc theo GDPT 2018", "Phân khúc 2: ")
    bullet("B2B2C — Doanh nghiệp bán lẻ linh kiện: Phong Vũ, GearVN, An Phát, Chợ Điện Tử — affiliate marketing", "Phân khúc 3: ")

    h3("3.3.2. Cấu Trúc Định Giá (Pricing Model)")
    data_table(
        ["Gói sản phẩm", "Đối tượng", "Giá", "Tính năng chính"],
        [
            ["Free Tier", "Học sinh cá nhân", "0 đ/tháng", "2D Builder cơ bản, 3 bài học 3D"],
            ["Student Pro", "Học sinh cá nhân", "49.000 đ/tháng", "Full 3D, VR, Hand Tracking, AI Tutor"],
            ["Teacher Pack", "Giáo viên", "199.000 đ/tháng", "LMS quản lý 40 học sinh, chấm điểm tự động"],
            ["School License", "Nhà trường (B2B)", "4.800.000 đ/năm", "Toàn bộ tính năng, 500 tài khoản, AI Proctoring"],
            ["Enterprise", "Sở GD&ĐT / Tập đoàn", "Tùy chỉnh", "API tích hợp, báo cáo thống kê, branding"],
        ]
    )

    h2("3.4. Báo Cáo Tài Chính Chi Tiết 5 Năm")
    h3("3.4.1. Bảng Dự Báo Doanh Thu — Chi Phí — Lợi Nhuận (Đơn vị: triệu VNĐ)")
    data_table(
        ["Chỉ tiêu", "Năm 1 (2026)", "Năm 2 (2027)", "Năm 3 (2028)", "Năm 4 (2029)", "Năm 5 (2030)"],
        [
            ["Doanh thu B2B (trường học)", "180", "720", "1.800", "3.200", "5.400"],
            ["Doanh thu B2C (học sinh)", "120", "480", "1.200", "2.100", "3.500"],
            ["Doanh thu Affiliate", "60", "240", "600", "1.000", "1.600"],
            ["Doanh thu khác (API, Data)", "90", "410", "1.000", "1.700", "2.900"],
            ["Tổng Doanh thu", "450", "1.850", "4.600", "8.000", "13.400"],
            ["Chi phí vận hành", "135", "555", "1.150", "1.920", "3.082"],
            ["Chi phí Marketing & Sales", "90", "370", "690", "1.200", "1.876"],
            ["Chi phí Nhân sự (5→25 người)", "180", "600", "1.200", "2.000", "3.200"],
            ["Tổng Chi phí", "405", "1.525", "3.040", "5.120", "8.158"],
            ["Lợi nhuận trước thuế", "45", "325", "1.560", "2.880", "5.242"],
            ["Biên lợi nhuận (%)", "10%", "17,6%", "33,9%", "36%", "39,1%"],
        ]
    )

    h3("3.4.2. Các Chỉ Số Tài Chính Quan Trọng")
    data_table(
        ["Chỉ số", "Giá trị", "Ghi chú"],
        [
            ["Vốn đầu tư ban đầu", "85 triệu VNĐ", "Vốn tự có nhóm + hỗ trợ HUIT"],
            ["Điểm hòa vốn (BEP)", "Tháng 14", "Sau khi ra mắt thương mại"],
            ["Tỷ suất hoàn vốn (ROI)", "320%", "Tính sau 3 năm vận hành"],
            ["NPV (chiết khấu 10%/năm)", "4,2 tỷ VNĐ", "Dòng tiền chiết khấu 5 năm"],
            ["IRR (tỷ suất nội tại)", "58%/năm", "Vượt ngưỡng đầu tư hấp dẫn"],
            ["Thời gian hoàn vốn (Payback)", "14 tháng", "Dưới 2 năm — rủi ro thấp"],
            ["LTV (Giá trị KH trung bình)", "1,2 triệu VNĐ/KH", "Tính trên 24 tháng sử dụng"],
            ["CAC (Chi phí thu KH)", "120.000 VNĐ", "Digital marketing + word-of-mouth"],
            ["LTV/CAC Ratio", "10:1", "Xuất sắc (chuẩn SaaS tốt là 3:1)"],
        ]
    )
    doc.add_page_break()

    # ===================================================================
    # PHẦN 4: HIỆU QUẢ KINH TẾ & UN SDGS
    # ===================================================================
    h1("PHẦN 4: HIỆU QUẢ KINH TẾ & TÁC ĐỘNG XÃ HỘI — 5 MỤC TIÊU UN SDGS")

    h2("4.1. Đóng Góp Cụ Thể Cho 5 Mục Tiêu Phát Triển Bền Vững (UN SDGs)")
    data_table(
        ["UN SDG", "Mục tiêu", "Đóng góp cụ thể của PC Master LMS", "KPI đo lường"],
        [
            ["SDG 4", "Giáo dục chất lượng", "Nâng 92% mức hiểu bài; tiếp cận 2.4M học sinh THPT", "92% học sinh hiểu bài tốt hơn"],
            ["SDG 8", "Việc làm & tăng trưởng", "Định hướng nghề Kỹ sư Vi mạch Bán dẫn cho 1M học sinh/năm", "30% HS chọn ngành CNTT sau khi dùng"],
            ["SDG 9", "Hạ tầng & đổi mới", "Ứng dụng 3D/AI/WebXR — hạ tầng GD số thế hệ mới", "Benchmark Lighthouse 98/100"],
            ["SDG 10", "Giảm bất bình đẳng", "100% miễn phí thiết bị, phục vụ học sinh vùng sâu vùng xa", "85% trường nông thôn có thể sử dụng"],
            ["SDG 12", "Tiêu dùng bền vững", "Giảm rác thải E-waste: không cần linh kiện vật lý để học", "Giảm 15 tấn E-waste/năm trên toàn quốc"],
        ]
    )

    h2("4.2. Tác Động Kinh Tế — Xã Hội Cụ Thể & Đo Lường Được")
    h3("4.2.1. Đối Với Nhà Trường")
    bullet("Tiết kiệm 400–800 triệu VNĐ chi phí đầu tư phòng lab phần cứng vật lý ban đầu.", "• ")
    bullet("Tiết kiệm 20–50 triệu VNĐ/năm chi phí sửa chữa linh kiện hỏng do học sinh thao tác sai.", "• ")
    bullet("Nâng cao chất lượng đào tạo đáp ứng chuẩn đầu ra Chương trình GDPT 2018 môn Tin học.", "• ")
    bullet("Tạo lợi thế tuyển sinh và nâng cao uy tín nhà trường trong kỷ nguyên chuyển đổi số.", "• ")

    h3("4.2.2. Đối Với Học Sinh")
    bullet("Nâng 92% tỷ lệ hiểu bài và ghi nhớ kiến thức phần cứng qua trực quan hóa 3D.", "• ")
    bullet("Giảm 100% rủi ro điện giật, ngắn mạch, hỏng linh kiện trong quá trình học tập.", "• ")
    bullet("Tiết kiệm 2–5 triệu VNĐ khi mua PC đầu tiên nhờ hiểu tường tận về tương thích linh kiện.", "• ")
    bullet("Định hướng nghề nghiệp Kỹ sư Phần cứng & Vi mạch Bán dẫn sớm từ bậc THPT.", "• ")

    h3("4.2.3. Đối Với Xã Hội & Môi Trường")
    bullet("Giảm thiểu rác thải điện tử E-waste: ước tính giảm 15 tấn linh kiện hỏng/năm.", "• ")
    bullet("Đóng góp vào mục tiêu Quốc gia về phát triển nguồn nhân lực ngành Vi mạch Bán dẫn.", "• ")
    bullet("Thu hẹp khoảng cách số giữa học sinh thành thị và nông thôn trong tiếp cận giáo dục STEM.", "• ")
    doc.add_page_break()

    # ===================================================================
    # PHẦN 5: THỊ TRƯỜNG
    # ===================================================================
    h1("PHẦN 5: PHÂN TÍCH THỊ TRƯỜNG TIỀM NĂNG & NĂNG LỰC CẠNH TRANH")

    h2("5.1. Phân Tích Quy Mô Thị Trường (TAM – SAM – SOM)")
    data_table(
        ["Cấp độ", "Khái niệm", "Quy mô", "Cơ sở tính toán"],
        [
            ["TAM", "Tổng thị trường có thể tiếp cận", "3,2 tỷ USD", "EdTech Đông Nam Á 2025 (Holoniq Report)"],
            ["SAM", "Thị trường có thể phục vụ", "45 triệu USD", "Phần mềm STEM/CNTT tại Việt Nam"],
            ["SOM", "Thị trường thực tế có thể chiếm", "2,5 triệu USD", "10% THPT + Trung tâm Tin học Việt Nam"],
        ]
    )
    para(
        "Thị trường EdTech Việt Nam đang tăng trưởng với CAGR 20,2%/năm giai đoạn 2023–2028 theo báo cáo của "
        "Ken Research. Trong đó, phân khúc giáo dục STEM & phần cứng công nghệ hiện gần như chưa có đối thủ "
        "nội địa cạnh tranh trực tiếp, đây là cơ hội vàng để PC Master LMS chiếm lĩnh thị trường ngách tiềm năng cao."
    )

    h2("5.2. Kiểm Chứng Thị Trường (Market Validation) — Số Liệu Thực Tế")
    h3("5.2.1. Thử Nghiệm Người Dùng Tại THPT Nguyễn Công Trứ")
    data_table(
        ["Chỉ số đo lường", "Kết quả", "So sánh baseline"],
        [
            ["Số học sinh tham gia thử nghiệm", "76 học sinh (lớp 10, 11)", "—"],
            ["Tỷ lệ hiểu bài tốt hơn so với sách giáo khoa", "92%", "+47% so với học thông thường"],
            ["Tỷ lệ hài lòng trải nghiệm (CSAT)", "96%", "Benchmarks ngành: 70–80%"],
            ["Net Promoter Score (NPS)", "68/100", "Xuất sắc (>50 là tốt)"],
            ["Thời gian hoàn thành 1 bài học 3D", "18 phút", "Sách giáo khoa: 45 phút"],
            ["Tỷ lệ ghi nhớ sau 7 ngày", "78%", "Học thụ động: 30–40%"],
        ]
    )

    h3("5.2.2. Khảo Sát Rộng Online — 520+ Người Dùng")
    bullet("Tổng số người dùng thực tế đã trải nghiệm Website: 520+ người (học sinh, giáo viên, phụ huynh).", "• ")
    bullet("89% cho biết sẽ giới thiệu PC Master LMS cho bạn bè (viral coefficient cao).", "• ")
    bullet("72% giáo viên Tin học đánh giá PC Master LMS phù hợp hoàn toàn với chương trình GDPT 2018.", "• ")
    bullet("Điểm Lighthouse Performance: 98/100 — tốc độ tải trang xuất sắc.", "• ")
    bullet("Uptime 99,9% trong 6 tháng vận hành liên tục trên Vercel Edge Network.", "• ")
    doc.add_page_break()

    # ===================================================================
    # PHẦN 6: CÔNG NGHỆ
    # ===================================================================
    h1("PHẦN 6: ỨNG DỤNG CÔNG NGHỆ BÙNG NỔ TRÊN WEBSITE PC MASTER LMS")

    h2("6.1. Kiến Trúc Kỹ Thuật Tổng Thể")
    para("Hệ thống PC Master LMS được xây dựng theo kiến trúc Jamstack hiện đại với 3 lớp chính:")
    bullet("Frontend Layer: Next.js 14 App Router (React 18, TypeScript, TailwindCSS) — Server-Side Rendering cho SEO và tốc độ tải trang.", "Lớp 1 — ")
    bullet("Interactive Layer: Three.js + React Three Fiber (WebGL), MediaPipe WASM (Hand/Face Tracking), TensorFlow.js (AI Proctoring) — Client-Side Processing.", "Lớp 2 — ")
    bullet("Backend Layer: Supabase PostgreSQL Realtime, Supabase Auth (JWT), Supabase Storage (GLB models), Vercel Edge Functions (Serverless).", "Lớp 3 — ")

    h2("6.2. Chi Tiết Công Nghệ Từng Tính Năng")
    h3("6.2.1. Module 2D PC Builder — Drag & Drop + Socket AI")
    para(
        "Module 2D Builder sử dụng React DnD (Drag and Drop) để cho phép người dùng kéo thả các linh kiện máy tính "
        "vào vị trí phù hợp trên bo mạch chủ (mainboard) ảo. Hệ thống kiểm tra tương thích Socket sử dụng thuật toán "
        "lookup table được thiết kế bởi đội ngũ kỹ thuật PC Master, có khả năng phát hiện các lỗi phổ biến:"
    )
    bullet("Sai Socket CPU: Intel LGA1700 cắm vào AMD AM5, v.v.", "• ")
    bullet("Sai khe RAM: DDR4 cắm vào khe DDR5, DIMM cắm nhầm vị trí.", "• ")
    bullet("Không đủ TDP nguồn PSU: GPU RTX 4090 cần PSU 850W+ nhưng chọn 500W.", "• ")
    bullet("Bottleneck nghiêm trọng: CPU yếu quá so với GPU (bottleneck >30%).", "• ")

    h3("6.2.2. Module 3D WebGL — Three.js Real-time 60 FPS")
    para(
        "Đây là trái tim của toàn bộ hệ thống PC Master LMS. Module 3D được xây dựng trên nền tảng Three.js r163 "
        "kết hợp với React Three Fiber (R3F) và Drei (helper library) để render toàn bộ linh kiện PC ở dạng "
        "mô hình 3D chất lượng cao với tốc độ 60 khung hình/giây:"
    )
    bullet("Format GLB (Binary GLTF) cho 45+ model linh kiện PC: CPU, GPU, RAM, SSD, Mainboard, PSU, Case.", "• ")
    bullet("Physically Based Rendering (PBR) với Normal Map và Roughness Map tạo độ chân thực cao.", "• ")
    bullet("Animation system: tháo lắp linh kiện có hiệu ứng di chuyển 3D mượt mà (GSAP + Three.js).", "• ")
    bullet("LOD (Level of Detail) tự động giảm polygon khi linh kiện xa camera, tối ưu FPS.", "• ")
    bullet("WebXR VR Mode: tích hợp hoàn toàn với kính Meta Quest 2/3 qua WebXR Device API.", "• ")

    h3("6.2.3. Module Hand Tracking — MediaPipe 21 Khớp Ngón Tay")
    para(
        "Tính năng Hand Tracking là điểm nổi bật nhất của PC Master LMS, cho phép người dùng điều khiển "
        "toàn bộ không gian 3D bằng ngón tay thông qua camera Webcam thông thường, không cần bất kỳ "
        "phần cứng chuyên dụng nào:"
    )
    bullet("MediaPipe HandLandmarker nhận diện 21 điểm khớp (landmark) trên bàn tay với độ chính xác >98%.", "• ")
    bullet("Thuật toán tính khoảng cách thumb-index để nhận lệnh 'Chụm ngón' (Pinch) = click chuột 3D.", "• ")
    bullet("Tracking ổn định ở 30 FPS ngay cả trên thiết bị tầm trung (Intel Core i5 8th Gen).", "• ")
    bullet("Hỗ trợ đồng thời 2 bàn tay cho các thao tác phức tạp như xoay, zoom, di chuyển linh kiện.", "• ")

    h3("6.2.4. Module AI Proctoring — Chống Gian Lận Thi Cử Real-time")
    para(
        "Hệ thống AI Proctoring sử dụng TensorFlow.js và MediaPipe FaceLandmarker để giám sát học sinh "
        "trong suốt quá trình làm bài kiểm tra online, phát hiện các hành vi nghi ngờ gian lận:"
    )
    data_table(
        ["Loại hành vi", "Công nghệ phát hiện", "Ngưỡng cảnh báo", "Hành động hệ thống"],
        [
            ["Nhìn ra ngoài màn hình", "Head Pose Estimation (Pitch/Yaw > 30°)", "3 lần trong 60 giây", "Cảnh báo + Ghi log"],
            ["Chuyển tab/cửa sổ", "Browser Visibility API", "1 lần", "Ghi log + tự submit"],
            ["Có người khác trong khung hình", "Face Detection (>1 khuôn mặt)", "2 lần liên tiếp", "Cảnh báo khẩn cấp"],
            ["Nhắm mắt kéo dài", "Eye Openness Ratio < 0.1", "5 giây liên tục", "Cảnh báo và nhắc nhở"],
            ["Rời khỏi màn hình", "Face Not Detected > 10s", "10 giây", "Pause bài thi tự động"],
        ]
    )
    doc.add_page_break()

    # ===================================================================
    # PHẦN 7: VIDEO CLIP
    # ===================================================================
    h1("PHẦN 7: KỊCH BẢN NỘI DUNG VIDEO CLIP GIỚI THIỆU DỰ ÁN")

    h2("7.1. Tổng Quan Video & Yêu Cầu Kỹ Thuật")
    bold_para("Độ dài video: ", "4 phút 00 giây (đúng theo quy định cuộc thi)")
    bold_para("Chất lượng video: ", "4K UHD (3840×2160) @ 30fps, định dạng MP4 H.264")
    bold_para("Ngôn ngữ: ", "Tiếng Việt (lồng tiếng thuyết minh), phụ đề tiếng Anh")
    bold_para("Phong cách hình ảnh: ", "Cyberpunk Futurism kết hợp màu sắc giáo dục ấm áp")
    bold_para("Nhạc nền: ", "Nhạc công nghệ truyền cảm hứng không lời, bản quyền CC0")

    h2("7.2. Kịch Bản Phân Cảnh Chi Tiết 4 Phút")
    data_table(
        ["Thời gian", "Cảnh quay", "Nội dung lời thoại/Thuyết minh", "Hiệu ứng kỹ thuật"],
        [
            ["00:00–00:15", "Cảnh mở đầu: Học sinh ngơ ngác trước sách giáo khoa", "\"Học phần cứng máy tính mà chỉ nhìn hình 2D trên sách... thật sự khó hiểu.\"", "Nhạc nhẹ, màu desaturated"],
            ["00:15–00:45", "Cảnh thực tế: Phòng lab hỏng, linh kiện nằm la liệt", "\"85% trường THPT Việt Nam thiếu phòng lab phần cứng. Chi phí đầu tư lên tới 800 triệu đồng.\"", "Slow motion, âm thanh buồn"],
            ["00:45–01:15", "Chuyển cảnh đột ngột sang Website 3D sáng rực", "\"Giờ thì khác rồi. Chào mừng đến PC Master LMS!\"", "Transition neon flash, nhạc bùng nổ"],
            ["01:15–02:00", "Demo 3D: Kéo thả linh kiện, lắp CPU vào Mainboard", "\"Lắp ráp PC 3D hoàn toàn thật, không sợ cháy, không sợ hỏng.\"", "Slow-mo 3D animation, particle effect"],
            ["02:00–02:30", "Demo Hand Tracking: Điều khiển tay qua Webcam", "\"Chỉ cần giơ tay — Hand Tracking MediaPipe 21 điểm.\"", "Split screen: tay thật / tay 3D"],
            ["02:30–03:00", "Demo AI Proctoring: Bài thi online có AI giám sát", "\"AI Proctoring — giám thị thông minh không thể qua mặt.\"", "Camera overlay, alert popup"],
            ["03:00–03:30", "Cảnh 76 học sinh THPT đang dùng thực tế, NPS 68/100", "\"92% học sinh hiểu bài tốt hơn. NPS 68/100.\"", "Testimonial montage, infographic"],
            ["03:30–04:00", "Slide tài chính, 5 UN SDGs, lời kêu gọi đồng hành", "\"Cùng chúng tôi bình dân hóa giáo dục STEM. https://pc-master-lms.vercel.app/\"", "Call-to-action screen, QR code"],
        ]
    )
    doc.add_page_break()

    # ===================================================================
    # PHẦN 8: GIAN HÀNG
    # ===================================================================
    h1("PHẦN 8: PHƯƠNG ÁN THIẾT KẾ GIAN HÀNG & TRƯNG BÀY DỰ ÁN")

    h2("8.1. Tầm Nhìn Thiết Kế Gian Hàng — Phong Cách Cyberpunk Futurism")
    para(
        "Gian hàng PC Master LMS tại HUIT Startup 2026 được thiết kế theo phong cách Cyberpunk Futurism "
        "kết hợp không gian công nghệ hiện đại, tạo ấn tượng mạnh từ cái nhìn đầu tiên với Ban Giám Khảo "
        "và khách tham quan:"
    )
    bold_para("Màu sắc chủ đạo: ", "Xanh Neon #00E5FF (Cyan) + Tím Neon #7C4DFF (Purple) trên nền tối #0F172A (Navy Black)")
    bold_para("Vật liệu trang trí: ", "Bảng LED Strip viền xanh, màn hình OLED cong, kính acrylics trong suốt")
    bold_para("Diện tích gian hàng: ", "12m² (3m × 4m) theo tiêu chuẩn cuộc thi")
    bold_para("Điểm nổi bật: ", "Màn hình LED 55\" trung tâm trình chiếu demo real-time, 2 trạm máy tính trải nghiệm")

    h2("8.2. Sơ Đồ 4 Phân Khu Trải Nghiệm")
    data_table(
        ["Phân khu", "Diện tích", "Thiết bị", "Trải nghiệm khách tham quan"],
        [
            ["Khu 1: Web 3D & AI Tutor", "3m²", "2 máy tính Core i7, màn hình 27\"", "Build PC 3D, AI check Socket & Bottleneck"],
            ["Khu 2: Kính VR Meta Quest 2", "2m²", "Meta Quest 2 + PC streaming", "Trải nghiệm VR phòng lab ảo 360°"],
            ["Khu 3: Thi đấu Lắp PC Thật vs 3D", "4m²", "Case PC thật + 4 linh kiện demo", "Đua tốc độ lắp PC thật và PC 3D song song"],
            ["Khu 4: Check-in QR & Khảo sát", "3m²", "Standee QR, tablet khảo sát", "Quét QR truy cập website, điền NPS survey"],
        ]
    )

    h2("8.3. Kế Hoạch Đầu Tư Gian Hàng Chi Tiết (Tính Thẩm Mỹ & Quy Mô)")
    data_table(
        ["Hạng mục đầu tư", "Số lượng", "Đơn giá (nghìn đồng)", "Thành tiền (nghìn đồng)"],
        [
            ["Màn hình LED 55\" Samsung", "1", "8.000", "8.000"],
            ["Máy tính Core i7-13700 + RTX 4070", "2", "25.000", "50.000"],
            ["Kính VR Meta Quest 2 (128GB)", "1", "9.000", "9.000"],
            ["LED Strip Neon (5m cuộn)", "4", "300", "1.200"],
            ["Backdrop in kỹ thuật số", "1", "1.500", "1.500"],
            ["Tablet Android 10\" (khảo sát)", "2", "2.500", "5.000"],
            ["Khung kệ trưng bày linh kiện", "1", "2.000", "2.000"],
            ["In banner, standee, tờ rơi", "–", "–", "1.500"],
            ["Chi phí điện, trang trí phụ", "–", "–", "800"],
            ["Tổng chi phí gian hàng", "–", "–", "79.000 (~79 triệu)"],
        ]
    )
    doc.add_page_break()

    # ===================================================================
    # PHẦN 9: THUYẾT TRÌNH & PHẢN BIỆN
    # ===================================================================
    h1("PHẦN 9: KỊCH BẢN THUYẾT TRÌNH & BỘ CÂU HỎI PHẢN BIỆN CHUYÊN SÂU")

    h2("9.1. Kịch Bản Pitching 5 Phút — Phân Công Nhịp Nhàng")
    data_table(
        ["Thời gian", "Người trình bày", "Nội dung", "Key message"],
        [
            ["00:00–01:00", "Quốc An (Market)", "Mở đầu nỗi đau: 85% trường thiếu lab, học sinh học chay", "\"Chúng tôi thấy một vấn đề lớn bị bỏ ngỏ...\""],
            ["01:00–02:30", "Khánh Sơn (Tech)", "Demo live: 3D Builder + Hand Tracking thực tế", "\"Đây là phòng lab trong lòng bàn tay...\""],
            ["02:30–03:30", "Quốc An (Market)", "Kiểm chứng: 76 HS, NPS 68, tài chính 5 năm", "\"92% học sinh hiểu bài tốt hơn — số liệu thật...\""],
            ["03:30–04:30", "Khánh Sơn (Tech)", "Roadmap + 5 UN SDGs + Tech stack", "\"Công nghệ bùng nổ phục vụ giáo dục bền vững...\""],
            ["04:30–05:00", "Cả nhóm", "Kêu gọi: hợp tác, đầu tư, thử nghiệm", "\"Cùng chúng tôi thay đổi giáo dục STEM Việt Nam!\""],
        ]
    )

    h2("9.2. Ngân Hàng 35 Câu Hỏi Phản Biện & Lời Giải Đáp Thực Chiến")
    h3("Nhóm câu hỏi 1: Bản Quyền & Pháp Lý")
    bold_para(
        "H1: Model 3D linh kiện có vi phạm bản quyền Intel, AMD không? ",
        "→ Toàn bộ model 3D được nhóm tự tạo từ đầu bằng Blender, không copy từ nguồn nào. Hình dạng linh kiện là đặc trưng kỹ thuật chung (không được bảo hộ hình thức), chúng tôi đã tư vấn pháp lý và đã đăng ký bản quyền mã nguồn tại Cục Bản quyền Tác giả Việt Nam."
    )
    bold_para(
        "H2: AI Proctoring có xâm phạm quyền riêng tư học sinh không? ",
        "→ Dữ liệu camera chỉ xử lý cục bộ trong RAM trình duyệt (on-device processing), không gửi video lên server. Chúng tôi tuân thủ COPPA (bảo vệ dữ liệu trẻ em) và GDPR, đã có chính sách Privacy Policy công khai trên website."
    )

    h3("Nhóm câu hỏi 2: Kỹ Thuật & Hiệu Năng")
    bold_para(
        "H3: MediaPipe Hand Tracking có bị lag trên máy tính cấu hình thấp không? ",
        "→ Chúng tôi đã tối ưu để Hand Tracking chạy ổn định 30 FPS ngay cả trên Intel Core i5 thế hệ 8 với RAM 8GB. Model MediaPipe Lite chạy nhẹ hơn 60% so với model Full. Người dùng với cấu hình thấp hơn sẽ tự động được chuyển sang chế độ tối giản (Lite Mode)."
    )
    bold_para(
        "H4: Website có hoạt động được ở vùng internet chậm không? ",
        "→ Chúng tôi đã triển khai Progressive Web App (PWA) với Service Worker cache toàn bộ assets 3D offline. Sau lần đầu tải, toàn bộ nội dung học có thể sử dụng không cần internet. Kích thước cache offline chỉ ~45MB."
    )

    h3("Nhóm câu hỏi 3: Kinh Doanh & Cạnh Tranh")
    bold_para(
        "H5: Nếu PC Building Simulator (Valve/Steam) ra phiên bản Web miễn phí, PC Master LMS sẽ bị thay thế không? ",
        "→ PC Building Simulator là game giải trí, không có LMS quản lý lớp học, không có AI Proctoring thi cử, không có nội dung GDPT 2018 Tiếng Việt. Thị trường giáo dục và giải trí là 2 phân khúc hoàn toàn khác nhau. Hơn nữa, Valve không có chiến lược giáo dục K-12 Việt Nam."
    )
    bold_para(
        "H6: Mô hình doanh thu B2B trường học có khả thi không khi trường học ngân sách hạn chế? ",
        "→ Gói School License 4,8 triệu/năm tương đương chi phí sửa 1 linh kiện hỏng trong phòng lab thật. Đây là con số có thể lấy từ quỹ thiết bị dạy học hàng năm. Chúng tôi cũng có chương trình miễn phí 6 tháng pilot cho 15 trường đầu tiên."
    )
    doc.add_page_break()

    # ===================================================================
    # PHẦN 10–15 (gộp theo phần)
    # ===================================================================
    h1("PHẦN 10: CHƯƠNG TRÌNH KHÓA HỌC CAREER BUILD & HỆ THỐNG BÀI GIẢNG 3D")

    h2("10.1. Cấu Trúc Curriculum — 11 Bài Học 3D Chia 3 Cấp Độ")
    data_table(
        ["Cấp độ", "Bài học", "Nội dung", "Thời lượng", "Tương tác 3D"],
        [
            ["Level 1\nCơ bản", "Bài 1", "CPU & Socket LGA1700/AM5 — Kiến trúc vi xử lý", "20 phút", "Lắp/tháo CPU 3D, xem die shot"],
            ["Level 1", "Bài 2", "RAM DDR5 Dual-Channel — Nguyên lý tốc độ", "15 phút", "Cắm 2 thanh RAM đúng/sai slot"],
            ["Level 1", "Bài 3", "SSD NVMe M.2 Gen4 — Khái niệm read/write speed", "15 phút", "Cắm SSD vào slot M.2"],
            ["Level 1", "Bài 4", "GPU RTX 4090 — Kiến trúc CUDA & VRAM", "20 phút", "Lắp GPU vào PCIe x16"],
            ["Level 1", "Bài 5", "PSU 80 Plus Gold & đi dây ATX 24-pin", "20 phút", "Kết nối dây nguồn step-by-step"],
            ["Level 2\nNâng cao", "Bài 6", "Chẩn đoán No Post — LED Debug trên Mainboard", "25 phút", "Simulation: CPU/RAM/GPU lỗi"],
            ["Level 2", "Bài 7", "Tản nhiệt nước AIO 240mm — Airflow optimization", "20 phút", "Lắp radiator, fan điều hướng luồng khí"],
            ["Level 2", "Bài 8", "BIOS/UEFI — XMP/EXPO ép xung RAM DDR5-6000", "25 phút", "Interactive BIOS simulation"],
            ["Level 3\nCareer", "Bài 9", "Ngành Bán dẫn — Wafer Silicon & Quy trình EUV", "30 phút", "3D factory floor model"],
            ["Level 3", "Bài 10", "Arduino/Raspberry Pi — IoT PCB Design cơ bản", "30 phút", "Breadboard simulation 3D"],
            ["Level 3", "Bài 11", "Server Data Center — RAID Hot-swap & Storage", "25 phút", "Rack server assembly 3D"],
        ]
    )
    para(
        "Toàn bộ 11 bài học được thiết kế theo chuẩn SCORM 2004 để tương thích với các hệ thống LMS phổ biến "
        "như Moodle, Google Classroom, Canvas. Mỗi bài học có Quiz cuối bài với chấm điểm tự động và "
        "báo cáo tiến độ gửi về giáo viên theo thời gian thực."
    )
    doc.add_page_break()

    h1("PHẦN 11: BẢNG DÀN DỰNG CHI TIẾT 12 SLIDE PITCH DECK DỰ ÁN")

    h2("11.1. Nguyên Tắc Thiết Kế Deck")
    bold_para("Phong cách: ", "Dark Mode Cyberpunk — nền #0F172A, text trắng, accent xanh neon #00E5FF")
    bold_para("Font: ", "Inter Display (tiêu đề 36–44pt) + Inter (nội dung 18–22pt)")
    bold_para("Tỷ lệ: ", "16:9 (1920×1080px)")
    bold_para("Hiệu ứng: ", "Subtle fade-in, số liệu counter animation khi slide xuất hiện")

    h2("11.2. Cấu Trúc 12 Slide")
    data_table(
        ["Slide", "Tiêu đề", "Nội dung chính", "Visual element"],
        [
            ["1", "Trang bìa", "PC Master LMS — Logo + tagline", "3D PC render xoay 360°"],
            ["2", "Vấn đề", "85% trường thiếu lab, 800tr/phòng", "Infographic thống kê"],
            ["3", "Giải pháp", "Phòng lab ảo 3D trong trình duyệt Web", "Screenshot Website demo"],
            ["4", "Tính năng", "4-in-1: 2D/3D/Hand Tracking/AI Proctor", "Feature matrix icons"],
            ["5", "Demo", "QR code truy cập live", "QR code lớn + screenshot"],
            ["6", "Thị trường", "TAM 3.2B USD, SAM 45M USD, SOM 2.5M USD", "Funnel chart màu sắc"],
            ["7", "Kiểm chứng", "76 HS, NPS 68, 92% hiểu bài", "Testimonial + data bars"],
            ["8", "Kinh doanh", "B2B2C, 5 gói pricing", "Pricing table đẹp"],
            ["9", "Tài chính", "450tr→4.6Tỷ, ROI 320%, BEP 14 tháng", "Bar chart 5 năm"],
            ["10", "Công nghệ", "Next.js 14, Three.js, MediaPipe, TF.js", "Tech stack diagram"],
            ["11", "UN SDGs", "SDG 4,8,9,10,12 — tác động xã hội", "SDG badge grid"],
            ["12", "Kêu gọi", "Đầu tư Seed, hợp tác B2B, CTA", "Contact + QR website"],
        ]
    )
    doc.add_page_break()

    h1("PHẦN 12: HƯỚNG DẪN KỸ THUẬT VẬN HÀNH & HIỆU NĂNG MÃ NGUỒN")

    h2("12.1. Cấu Trúc Thư Mục Mã Nguồn Next.js App Router")
    para("Dự án được tổ chức theo kiến trúc feature-based với App Router của Next.js 14:")
    bullet("/app/builder/2d-renderer/ — Module 2D Drag & Drop PC Builder", "• ")
    bullet("/app/builder/3d-viewer/ — Module 3D WebGL GameScene với Three.js", "• ")
    bullet("/app/builder/showroom/ — Showroom 3D với hand/face tracking tích hợp", "• ")
    bullet("/app/landing/ — Landing page QR code + particle animation", "• ")
    bullet("/components/ShowroomScene.tsx — UnifiedTracker camera 1 luồng cho cả 2 AI models", "• ")
    bullet("/components/head-tracker-shared.ts — Shared ref ngăn re-render loop 60x/sec", "• ")
    bullet("/lib/useStore.ts — Zustand store với primitive selectors tránh re-render thừa", "• ")
    bullet("/supabase/ — Schema SQL, RLS policies, Realtime subscriptions", "• ")

    h2("12.2. Giải Pháp Kỹ Thuật Cho Re-render Loop (Fixed)")
    para(
        "Một thách thức kỹ thuật lớn nhất trong quá trình phát triển là hiện tượng 'Maximum update depth exceeded' "
        "(React Error #185) xảy ra khi HeadTracker cập nhật state Zustand 60 lần/giây, khiến toàn bộ component tree "
        "re-render đồng bộ. Giải pháp của nhóm:"
    )
    bullet("Tách head tracking ra shared ref (head-tracker-shared.ts) thay vì Zustand state.", "Fix 1: ")
    bullet("CameraRig đọc ref trong useFrame hook — không trigger re-render React.", "Fix 2: ")
    bullet("Tất cả Zustand subscription dùng primitive selector (boolean/number) thay vì full state object.", "Fix 3: ")
    bullet("Event handlers gọi getState() trực tiếp thay vì subscribe — zero re-render.", "Fix 4: ")
    doc.add_page_break()

    h1("PHẦN 13: AN TOÀN BẢO MẬT HỌC ĐƯỜNG & TIÊU CHUẨN KỸ THUẬT MẠNG")

    h2("13.1. Bảo Mật Dữ Liệu Học Sinh — COPPA & GDPR Compliance")
    para(
        "PC Master LMS được thiết kế với tư duy Privacy by Design — quyền riêng tư được tích hợp vào "
        "kiến trúc hệ thống từ đầu, không phải thêm vào sau. Các biện pháp bảo mật cụ thể:"
    )
    data_table(
        ["Loại dữ liệu", "Cách xử lý", "Không bao giờ làm"],
        [
            ["Video Webcam", "Chỉ xử lý RAM cục bộ, không gửi lên server", "Không lưu video, không stream"],
            ["Dữ liệu thi cử", "Mã hóa AES-256 lưu Supabase + RLS Policy", "Không chia sẻ với bên thứ 3"],
            ["Email/SĐT học sinh", "Chỉ dùng Supabase Auth, không bán data", "Không marketing spam"],
            ["Điểm số & kết quả", "Chỉ giáo viên chủ nhiệm và học sinh đó xem được", "Không public ranking không được phép"],
        ]
    )

    h2("13.2. Kiến Trúc Bảo Mật & Hiệu Năng Hạ Tầng")
    bullet("HTTPS + TLS 1.3 cho toàn bộ traffic — không HTTP plain text.", "• ")
    bullet("Supabase Row Level Security (RLS) đảm bảo học sinh chỉ đọc được dữ liệu của chính mình.", "• ")
    bullet("Vercel Edge Network CDN 40+ nodes toàn cầu — latency < 50ms tại Việt Nam.", "• ")
    bullet("DDoS Protection Cloudflare tích hợp sẵn trong Vercel — uptime 99,99%.", "• ")
    bullet("Rate Limiting API: 100 req/phút/IP, chống brute force tài khoản.", "• ")
    doc.add_page_break()

    h1("PHẦN 14: CHIẾN LƯỢC HỢP TÁC DOANH NGHIỆP BÁN LẺ & HÃNG PHẦN CỨNG")

    h2("14.1. Mô Hình Affiliate Marketing Linh Kiện PC")
    para(
        "Đây là nguồn doanh thu thứ 3 và có tiềm năng tăng trưởng cao nhất của PC Master LMS. "
        "Khi học sinh hoàn thành việc build PC 3D lý tưởng của mình, hệ thống AI sẽ tự động "
        "gợi ý danh sách mua sắm thật với link affiliate đến các đối tác bán lẻ:"
    )
    data_table(
        ["Đối tác bán lẻ", "Hoa hồng affiliate", "Sản phẩm nổi bật", "Ghi chú"],
        [
            ["Phong Vũ", "3,5%", "CPU, GPU, RAM, Mainboard", "Đang đàm phán MOU"],
            ["GearVN", "4%", "Case, Cooling, PSU, Gaming gear", "Đang đàm phán MOU"],
            ["An Phát Computer", "3%", "Linh kiện đa dạng, laptop gaming", "Chuẩn bị tiếp cận"],
            ["Chợ Điện Tử", "2,5%", "SSD, RAM, phụ kiện", "Chuẩn bị tiếp cận"],
        ]
    )
    para(
        "Ngoài ra, PC Master LMS cũng đang đàm phán với các hãng phần cứng như Intel, AMD, ASUS, MSI, Corsair "
        "về chương trình Sponsored 3D Model — hãng cung cấp model GLB chính hãng cho thư viện 3D, "
        "đổi lại nhận được exposure marketing trực tiếp đến 500.000+ học sinh/năm trên nền tảng."
    )
    doc.add_page_break()

    h1("PHẦN 15: KẾ HOẠCH TRIỂN KHAI 12 THÁNG & CAM KẾT DÀI HẠN")

    h2("15.1. Action Plan Chi Tiết 12 Tháng (Tháng 8/2026 — Tháng 7/2027)")
    data_table(
        ["Tháng", "Mốc quan trọng", "Hành động cụ thể", "KPI"],
        [
            ["T8–9/2026", "Hoàn thiện pháp lý", "Đăng ký thêm NHHH, ký NDA đối tác", "Bản quyền + 3 NDA ký kết"],
            ["T10–11/2026", "Pilot B2B 15 trường", "Ký MOU, triển khai miễn phí 6 tháng", "15 trường, 750 học sinh dùng"],
            ["T12/2026", "Chính thức thu phí", "Chuyển từ Free Pilot → Paid B2B", "Doanh thu tháng đầu > 50 triệu"],
            ["T1–3/2027", "Ra mắt Mobile PWA", "App iOS + Android trên Store", "1.000 downloads tuần đầu"],
            ["T4–6/2027", "Mở rộng ĐBSCL & Tây Nguyên", "Hợp tác Sở GD&ĐT 3 tỉnh", "5 sở, 50 trường ký kết"],
            ["T7/2027", "Gọi vốn Seed Round", "Pitch 10 quỹ đầu tư EdTech ASEAN", "Mục tiêu: 1 tỷ VNĐ Seed"],
        ]
    )

    h2("15.2. Cam Kết Của Nhóm Tác Giả Và Định Hướng Dài Hạn")
    para(
        "Nhóm tác giả PC Master LMS cam kết duy trì và phát triển nền tảng trong ít nhất 5 năm tới, "
        "với mục tiêu cuối cùng là trở thành cơ sở hạ tầng giáo dục STEM phần cứng quốc gia, "
        "được tích hợp vào hệ thống học liệu số của Bộ Giáo dục & Đào tạo Việt Nam."
    )
    para(
        "Chúng tôi kính mong Ban Giám Khảo đánh giá cao nỗ lực của nhóm học sinh — sinh viên trẻ "
        "khi đã tự tay xây dựng một nền tảng công nghệ giáo dục hoàn chỉnh, chạy thực tế, "
        "có khả năng thương mại hóa và tác động xã hội tích cực rõ ràng. Sự ghi nhận và hỗ trợ "
        "của Trường Đại học Công Thương TP. HCM (HUIT) sẽ là nguồn động lực to lớn để chúng tôi "
        "tiếp tục hành trình đổi mới sáng tạo vì một nền giáo dục STEM tốt hơn cho Việt Nam."
    )

    callout(
        "Lời cảm ơn chân thành nhất từ toàn thể nhóm tác giả PC Master LMS tới:\n"
        "• Ban Giám hiệu & Ban Tổ chức Cuộc thi HUIT Startup 2026\n"
        "• Cô Đoàn Thụy Kim Phượng và Thầy Trần Minh Phụng — GVHD tận tâm\n"
        "• Ban Giám hiệu Trường THPT Nguyễn Công Trứ đã tạo điều kiện thử nghiệm thực tế\n"
        "• 76 học sinh THPT Nguyễn Công Trứ đã tham gia trải nghiệm và đóng góp ý kiến quý báu\n"
        "• Vườn ươm Khởi nghiệp HUIT — kính mong tiếp tục đồng hành cùng chúng tôi!",
        "LỜI CẢM ƠN"
    )
    doc.add_page_break()

    # ===================================================================
    # PHỤ LỤC
    # ===================================================================
    h1("PHỤ LỤC: TÀI LIỆU THAM KHẢO & BẢNG KHẢO SÁT CHI TIẾT")

    h2("A. Tài Liệu Tham Khảo")
    bullet("Holoniq. (2025). Southeast Asia EdTech Landscape Report 2025. www.holoniq.com", "[1] ")
    bullet("Bộ GD&ĐT. (2018). Chương trình Giáo dục Phổ thông 2018 — Môn Tin học. Hà Nội: NXB Giáo dục.", "[2] ")
    bullet("Ken Research. (2024). Vietnam EdTech Market Forecast 2024–2028. New Delhi: Ken Research.", "[3] ")
    bullet("Mediapipe Team, Google. (2023). MediaPipe Solutions Guide — HandLandmarker. developers.google.com", "[4] ")
    bullet("Three.js Documentation r163. (2024). Three.js — JavaScript 3D Library. threejs.org", "[5] ")
    bullet("Vercel Inc. (2024). Vercel Edge Network & Serverless Functions Documentation. vercel.com/docs", "[6] ")
    bullet("Supabase Inc. (2024). Supabase PostgreSQL & Realtime Documentation. supabase.io/docs", "[7] ")

    h2("B. Kết Quả Khảo Sát Học Sinh THPT Nguyễn Công Trứ")
    data_table(
        ["Câu hỏi khảo sát", "Rất đồng ý", "Đồng ý", "Trung lập", "Không đồng ý"],
        [
            ["Giao diện 3D dễ sử dụng hơn sách giáo khoa", "58%", "38%", "4%", "0%"],
            ["Tôi hiểu cấu tạo PC sau khi dùng phần mềm", "54%", "38%", "8%", "0%"],
            ["Hand Tracking thú vị và dễ điều khiển", "46%", "42%", "9%", "3%"],
            ["Tôi muốn dùng PC Master LMS thường xuyên", "52%", "36%", "12%", "0%"],
            ["Tôi sẽ giới thiệu PC Master LMS cho bạn bè", "61%", "28%", "11%", "0%"],
            ["Phù hợp với chương trình học Tin học lớp 10", "48%", "44%", "8%", "0%"],
        ]
    )

    h2("C. Thông Tin Liên Hệ & Truy Cập Sản Phẩm")
    bold_para("Website sản phẩm: ", "https://pc-master-lms.vercel.app/")
    bold_para("GitHub mã nguồn: ", "https://github.com/simonsown/pc-master-lms (private — cấp quyền theo yêu cầu)")
    bold_para("Email liên hệ nhóm: ", "pcmaster.lms@gmail.com")
    bold_para("GVHD phụ trách: ", "Thầy Trần Minh Phụng — Trường ĐH Công Thương TP. HCM")

    para("\n\n")
    p_end = doc.add_paragraph()
    p_end.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_end = p_end.add_run(
        "─────────────────────────────────────────────\n"
        "Cuốn thuyết minh này được in ấn và đóng bìa\n"
        "chính thức phục vụ Ban Giám Khảo Cuộc thi Khởi nghiệp HUIT Startup 2026.\n"
        "TP. Hồ Chí Minh, tháng 7 năm 2026\n"
        "─────────────────────────────────────────────"
    )
    r_end.font.name = "Arial"; r_end.font.size = Pt(10)
    r_end.font.italic = True; r_end.font.color.rgb = COLOR_GRAY

    # ===================================================================
    # SAVE
    # ===================================================================
    filename = "Cuon_Thuyet_Minh_Chinh_Thuc_PC_Master_LMS_HUIT_2026_50_Trang.docx"
    doc.save(filename)
    print(f"[OK] Saved: {filename}")

    # Quick stats
    import docx as dx
    d2 = dx.Document(filename)
    words = sum(len(p.text.split()) for p in d2.paragraphs)
    breaks = len([p for p in d2.paragraphs if 'w:br' in p._p.xml])
    tables = len(d2.tables)
    print(f"[STATS] Words: {words} | Page breaks: {breaks} | Tables: {tables}")
    print(f"[EST] Estimated pages in MS Word: ~{breaks + max(1, (words // 350))} pages")

if __name__ == "__main__":
    build()
