import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

COLOR_NAVY   = RGBColor(3, 31, 59)
COLOR_GREEN  = RGBColor(8, 158, 96)
COLOR_DARK   = RGBColor(30, 41, 59)
COLOR_GRAY   = RGBColor(100, 116, 139)
COLOR_CORRECT= RGBColor(16, 185, 129)

def add_topic_header(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after  = Pt(6)
    r = p.add_run(text)
    r.font.name  = 'Arial'
    r.font.size  = Pt(13)
    r.font.bold  = True
    r.font.color.rgb = COLOR_GREEN

def add_question(doc, num, question, answer):
    # Question
    pq = doc.add_paragraph()
    pq.paragraph_format.space_before = Pt(8)
    pq.paragraph_format.space_after  = Pt(2)
    rq = pq.add_run(f"Câu {num}: {question}")
    rq.font.name  = 'Arial'
    rq.font.size  = Pt(11)
    rq.font.bold  = True
    rq.font.color.rgb = COLOR_DARK

    # Answer label
    pa = doc.add_paragraph()
    pa.paragraph_format.space_before = Pt(2)
    pa.paragraph_format.space_after  = Pt(2)
    ra_label = pa.add_run("→ Hướng trả lời: ")
    ra_label.font.name  = 'Arial'
    ra_label.font.size  = Pt(10.5)
    ra_label.font.bold  = True
    ra_label.font.color.rgb = COLOR_CORRECT
    ra_body = pa.add_run(answer)
    ra_body.font.name  = 'Arial'
    ra_body.font.size  = Pt(10.5)
    ra_body.font.color.rgb = COLOR_DARK

    # Separator
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

def build_doc():
    doc = docx.Document()
    for section in doc.sections:
        section.top_margin    = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin   = Inches(0.9)
        section.right_margin  = Inches(0.9)

    # ── TITLE ────────────────────────────────────────────
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_title.add_run("NGÂN HÀNG 105 CÂU HỎI PHẢN BIỆN DỰ ĐOÁN\nBAN GIÁM KHẢO CUỘC THI KHỞI NGHIỆP HUIT 2026")
    r.font.name = 'Arial'; r.font.size = Pt(18); r.font.bold = True; r.font.color.rgb = COLOR_NAVY

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = p_sub.add_run("Dự án: PC MASTER BUILDER — Hệ thống LMS & Giả lập Lắp ráp PC 3D/AI\n"
                       "Trường Đại học Công Thương TP.HCM (HUIT) | Chỉ hướng trả lời — Không pitching — Bám sát thuật ngữ dự án")
    rs.font.name = 'Arial'; rs.font.size = Pt(10); rs.font.italic = True; rs.font.color.rgb = COLOR_GREEN

    doc.add_paragraph()

    # ════════════════════════════════════════════════════
    # CHỦ ĐỀ 1 — THỊ TRƯỜNG & NHU CẦU (15 câu)
    # ════════════════════════════════════════════════════
    add_topic_header(doc, "CHỦ ĐỀ 1: THỊ TRƯỜNG, NHU CẦU NGƯỜI DÙNG & ĐỘNG LỰC DỰ ÁN (15 câu)")

    qa1 = [
        ("YouTube và Google đã có hướng dẫn lắp ráp PC rất nhiều, tại sao học sinh vẫn cần PC Master?",
         "YouTube chỉ cung cấp học thụ động (passive learning). PC Master cung cấp môi trường thực hành chủ động thông qua giả lập 3D tương tác, bài kiểm tra bám sát SGK Tin học 10 (chương trình GDPT 2018), và hệ thống LMS để giáo viên giao bài & theo dõi tiến độ — điều không nền tảng nào trên YouTube làm được."),
        ("Nạn chặt chém giá linh kiện có thực sự ảnh hưởng đến học sinh THPT không?",
         "Có. Học sinh và phụ huynh thiếu kiến thức kỹ thuật (information asymmetry) rất dễ bị tư vấn linh kiện cũ hoặc không tương thích với giá cao. Tính năng tra giá linh kiện real-time và AI compatibility check của PC Master giúp người dùng so sánh ngay với giá thị trường chuẩn tại Việt Nam."),
        ("Điểm nghẽn lớn nhất trong giảng dạy phần cứng tại trường THPT hiện nay là gì?",
         "Thiếu phòng Lab phần cứng thực tế do chi phí trang bị linh kiện thật quá cao (hàng trăm triệu đồng/phòng) và rủi ro học sinh làm hỏng thiết bị (gãy chân pin CPU, tĩnh điện hư RAM). PC Master cung cấp Virtual Lab 3D giải quyết cả hai bài toán chi phí và an toàn cùng lúc."),
        ("Tại sao nhóm chọn học sinh THPT thay vì sinh viên Đại học CNTT làm thị trường đầu tiên?",
         "SGK Tin học 10 theo chương trình GDPT 2018 có chuyên đề phần cứng PC bắt buộc, nhưng hầu hết trường THPT không có phòng lab thực hành. Đây là nhu cầu được xác định rõ, thị trường tập trung, dễ kiểm chứng hiệu quả học tập — tiêu chí lý tưởng cho beachhead market của startup EdTech."),
        ("Bằng chứng nào cho thấy thị trường thực sự cần sản phẩm này?",
         "Nhóm đã thử nghiệm thực tế với 76 học sinh lớp 10A1 và 10C9 tại một trường THPT, thu được dữ liệu về mức độ tương tác, điểm kiểm tra và phản hồi từ giáo viên. Kết quả định lượng từ pilot này là bằng chứng trực tiếp (product-market fit signal)."),
        ("Cơ sở nào chứng minh phụ huynh sẵn sàng chi trả cho sản phẩm?",
         "Phụ huynh không chi trả cho 'học online' mà chi trả cho kết quả đầu ra cụ thể: con học tốt hơn môn Tin học, tránh được rủi ro mua linh kiện sai. Gói Premium 20.000đ/tháng — thấp hơn một buổi học thêm — là mức giá dễ chấp nhận khi giá trị được chứng minh qua pilot."),
        ("Ngoài THPT, PC Master có thể mở rộng sang phân khúc nào khác?",
         "Giai đoạn tiếp theo nhắm vào (1) sinh viên Cao đẳng/Đại học CNTT cần kỹ năng thực hành phần cứng, (2) cơ sở đào tạo nghề IT, và (3) mô hình B2B cấp phép hệ thống LMS cho các trường qua gói School License. Roadmap mở rộng này được thiết kế theo mô hình expansion revenue."),
        ("Quy mô thị trường (TAM/SAM/SOM) của dự án là bao nhiêu?",
         "TAM là toàn bộ thị trường giáo dục K-12 và dạy nghề IT tại Việt Nam. SAM là học sinh THPT từ lớp 10-12 học chuyên đề phần cứng (~600,000 học sinh/năm). SOM mục tiêu 3 năm đầu là 10,000 học sinh trả phí Premium — tương đương 2.4 tỷ đồng ARR."),
        ("Tại sao không xây dựng ứng dụng native (iOS/Android) thay vì web app?",
         "Web app cho phép học sinh truy cập ngay trên bất kỳ thiết bị nào — điện thoại, laptop trường, máy tính gia đình — không cần cài đặt, không bị giới hạn bởi App Store review. Đây là quyết định kỹ thuật ưu tiên khả năng tiếp cận (accessibility) và tốc độ triển khai trong giai đoạn MVP."),
        ("Mô hình Freemium có hiệu quả với học sinh không khi họ quen dùng miễn phí?",
         "Freemium hoạt động hiệu quả khi tier miễn phí đủ tạo thói quen nhưng tier trả phí cung cấp giá trị không thể thiếu. Gói Free cho phép thực hành cơ bản; Premium mở khóa AI gợi ý cấu hình, VR Lab 3D nâng cao và chứng chỉ hoàn thành — đây là giá trị gắn với kết quả học tập thực sự, không phải tính năng ảo."),
        ("Đối thủ trực tiếp của PC Master là ai?",
         "PC Building Simulator (game nước ngoài), một số web tĩnh hướng dẫn lắp ráp. Không đối thủ nào kết hợp: giả lập 3D tương tác + LMS quản lý lớp học + AI tư vấn linh kiện thị trường Việt Nam + bám sát SGK GDPT 2018 + tính năng VR/Computer Vision. Đây là khoảng trắng thị trường (white space) PC Master đang khai thác."),
        ("Làm thế nào để thu hút trường học đăng ký hàng loạt?",
         "Chiến lược B2B bottom-up: giáo viên dùng thử miễn phí → thấy hiệu quả → đề xuất lên ban giám hiệu → trường ký hợp đồng School License. Kênh phân phối này đã được chứng minh hiệu quả bởi các EdTech lớn như Kahoot, Google Classroom. Pilot 76 học sinh là bước đầu xây dựng social proof cho kênh này."),
        ("Dữ liệu người dùng được thu thập và sử dụng như thế nào?",
         "Dữ liệu học tập (learning analytics) — tiến độ bài học, điểm kiểm tra, thời gian thực hành — được lưu trên Supabase với RLS (Row Level Security), chỉ cho phép học sinh xem dữ liệu của chính họ, giáo viên xem học sinh trong lớp mình. Không bán dữ liệu cho bên thứ ba."),
        ("Nếu Bộ GD&ĐT thay đổi chương trình SGK, sản phẩm có bị lỗi thời không?",
         "Nội dung bài học trong LMS được thiết kế module hóa — giáo viên có thể cập nhật, thêm/sửa bài học thông qua Creator Dashboard mà không cần can thiệp kỹ thuật. Kiến trúc nội dung linh hoạt giúp PC Master thích ứng nhanh với thay đổi chương trình, đây là ưu điểm so với sách giáo khoa vật lý."),
        ("PC Master giải quyết bao nhiêu vấn đề cùng lúc — có bị dàn trải không?",
         "Ba vấn đề core được giải quyết trên một nền tảng thống nhất: (1) Thiếu lab phần cứng → Virtual Lab 3D, (2) Thiếu công cụ quản lý lớp học số → LMS với real-time collaboration, (3) Thiếu kiến thức tư vấn linh kiện → AI Compatibility Checker & giá real-time. Ba vấn đề này có cùng đối tượng người dùng, tích hợp trên một platform là lợi thế chứ không phải dàn trải."),
    ]
    for i, (q, a) in enumerate(qa1, 1):
        add_question(doc, i, q, a)

    # ════════════════════════════════════════════════════
    # CHỦ ĐỀ 2 — CÔNG NGHỆ 3D/VR & COMPUTER VISION (15 câu)
    # ════════════════════════════════════════════════════
    add_topic_header(doc, "CHỦ ĐỀ 2: CÔNG NGHỆ 3D/VR, COMPUTER VISION & STACK KỸ THUẬT (15 câu)")
    qa2 = [
        ("Tại sao dùng Three.js thay vì Unity hay Unreal Engine cho giả lập 3D?",
         "Three.js chạy trực tiếp trên trình duyệt (WebGL) — không cần cài đặt, không cần máy tính mạnh, hoạt động trên cả điện thoại. Unity/Unreal đòi hỏi cài app riêng và phần cứng đồ họa mạnh. Với đối tượng học sinh THPT dùng thiết bị đa dạng, web-first là lựa chọn tối ưu về khả năng tiếp cận."),
        ("Computer Vision trong PC Master hoạt động như thế nào?",
         "Sử dụng MediaPipe (Google) chạy trên WebAssembly ngay trong trình duyệt, không cần server xử lý. FaceLandmarker theo dõi 478 điểm khuôn mặt để điều khiển camera 3D theo hướng nhìn. HandLandmarker theo dõi 21 khớp ngón tay để nhận diện gesture: trỏ tay để di chuyển component, chụm ngón cái-trỏ để xoay linh kiện trong VR."),
        ("MediaPipe chạy client-side có đủ nhanh cho trải nghiệm thực tế không?",
         "Có. MediaPipe trên WebAssembly đạt 30fps trên laptop phổ thông. Để tối ưu thêm, nhóm dùng shared ref pattern thay vì Zustand store để tránh re-render — head tracking data được ghi vào ref trực tiếp trong useFrame của React Three Fiber, không trigger re-render React. Latency đo được dưới 50ms."),
        ("Tại sao chọn React Three Fiber thay vì Three.js thuần?",
         "React Three Fiber cho phép khai báo scene 3D theo cú pháp React JSX — dễ maintain, dễ tích hợp với Zustand store và Next.js. Các component 3D (PcCaseSimple, CameraRig, VrComponent) được quản lý như React component thông thường, giúp team chia việc song song hiệu quả."),
        ("VR Lab trong PC Master hoạt động thế nào — có cần kính VR không?",
         "Không cần kính VR. PC Master dùng Face Tracking qua webcam để mô phỏng góc nhìn 3D (head-tracked parallax). Khi người dùng nghiêng đầu trái/phải, camera trong scene 3D dịch chuyển tương ứng, tạo cảm giác chiều sâu không gian thực. Đây là VR không cần phần cứng đặc biệt — trải nghiệm immersive với chi phí zero cho học sinh."),
        ("Showroom 3D khác gì với Virtual Lab 3D?",
         "Showroom (/builder/showroom) là không gian trưng bày: một GLB model 3D duy nhất quay quanh người dùng (orbit), điều khiển bằng hand gesture — chỉ tay để component lại gần, chụm ngón để xoay. Virtual Lab (/builder/3d-viewer) là không gian thực hành đầy đủ: lắp ráp nhiều linh kiện vào PC case, kiểm tra compatibility, boot sequence animation."),
        ("UnifiedTracker là gì và tại sao dùng một camera thay vì hai?",
         "UnifiedTracker là component kết hợp FaceLandmarker và HandLandmarker chạy trên cùng một MediaPipe instance với một luồng camera. Dùng một camera giảm 50% latency khởi động, tránh xung đột quyền truy cập camera giữa hai tracker, và giảm tải CPU. Architecture này là tối ưu hóa kỹ thuật so với thiết kế ban đầu dùng hai tracker riêng biệt."),
        ("Supabase Realtime được dùng để làm gì trong PC Master?",
         "Supabase Realtime (dùng WebSocket) cho phép nhiều người dùng thấy thay đổi đồng thời: khi giáo viên cập nhật bài học, học sinh thấy ngay; khi học sinh hoàn thành bài tập, giáo viên thấy notification real-time. CollaborationStatus component hiển thị danh sách người đang online trong cùng builder session."),
        ("Tại sao chọn Supabase thay vì Firebase hay tự xây backend?",
         "Supabase cung cấp PostgreSQL thật (không phải NoSQL) với RLS (Row Level Security) cho phép viết policy bảo mật cấp database — giáo viên chỉ xem học sinh trong lớp mình, admin xem tất cả. Firebase dùng Firestore NoSQL khó query phức tạp. Tự xây backend tốn thời gian không phù hợp giai đoạn MVP startup."),
        ("Kiến trúc Next.js App Router được dùng như thế nào?",
         "Toàn bộ routing dùng Next.js 15 App Router: server components cho các trang tĩnh/SEO (landing, courses), client components ('use client') cho các trang tương tác (builder, 3D viewer). Layout lồng nhau: root layout → dashboard layout (student/teacher/admin) → page. API routes trong /app/api/ xử lý webhook Supabase và server actions cho auth."),
        ("Zustand được dùng để quản lý state gì trong 3D viewer?",
         "useAssemblyStore (Zustand) lưu: danh sách components đã đặt vào case (slots), trạng thái boot sequence (isBooting, bootProgress), compatibility errors. Để tránh Maximum Update Depth Exceeded, tất cả subscription dùng primitive selectors — ví dụ: useAssemblyStore(s => s.isBooting) thay vì useAssemblyStore() để không trigger re-render khi state khác thay đổi."),
        ("AI gợi ý cấu hình hoạt động như thế nào?",
         "Người dùng nhập budget và nhu cầu (gaming, học tập, đồ họa); AI Compatibility Checker phân tích compatibility matrix giữa các linh kiện (CPU socket phải khớp mainboard chipset, RAM DDR4/DDR5 phải đúng slot, TDP CPU phải trong giới hạn PSU). Kết quả trả về danh sách build suggestions kèm giá real-time từ database linh kiện Việt Nam."),
        ("Boot Sequence Animation trong PC Case 3D hoạt động ra sao?",
         "Khi người dùng bấm 'Khởi động', scene 3D phát animation: đèn LED case sáng lên, fan quay, màn hình POST (Power-On Self-Test) hiện. Animation được điều khiển bằng useFrame hook trong React Three Fiber, đọc bootProgress từ Zustand store, thay đổi material emissiveIntensity và rotation của fan mesh theo timeline."),
        ("WebGL có hoạt động được trên điện thoại Android phổ thông không?",
         "Có. Three.js/WebGL hoạt động trên Chrome Android từ phiên bản 70+ (2018). Nhóm đã test trên các thiết bị phổ thông từ 2GB RAM. Để tối ưu mobile: giảm polygon count của GLB models, dùng draco compression, lazy load scene khi cần, tắt shadow map trên mobile. Người dùng mobile được redirect sang interface đơn giản hơn nếu WebGL không đủ performance."),
        ("Làm thế nào để đảm bảo chất lượng 3D models linh kiện chính xác?",
         "GLB models được export từ Blender với độ chính xác cao về hình dạng và tỷ lệ so với linh kiện thật (CPU Intel/AMD, RAM stick, GPU card). Mỗi model có UV map texture baked từ ảnh linh kiện thật. Trong Showroom, người dùng có thể zoom và xoay 360° để kiểm tra chi tiết — đây là giá trị giáo dục trực quan không thể đạt được từ ảnh 2D."),
    ]
    for i, (q, a) in enumerate(qa2, 1):
        add_question(doc, 15 + i, q, a)

    # ════════════════════════════════════════════════════
    # CHỦ ĐỀ 3 — LMS, QUẢN LÝ LỚP HỌC & TÍNH NĂNG GIÁO DỤC (15 câu)
    # ════════════════════════════════════════════════════
    add_topic_header(doc, "CHỦ ĐỀ 3: HỆ THỐNG LMS, QUẢN LÝ LỚP HỌC & TÍNH NĂNG GIÁO DỤC (15 câu)")
    qa3 = [
        ("LMS trong PC Master khác gì Google Classroom hay Moodle?",
         "Google Classroom và Moodle là LMS đa năng không có domain kiến thức cụ thể. PC Master LMS tích hợp domain-specific: bài tập giao trực tiếp trong môi trường giả lập 3D, điểm số dựa trên kết quả thực hành lắp ráp (không phải chỉ trắc nghiệm text), và analytics về thao tác của học sinh trong virtual lab — điều không LMS nào khác làm được."),
        ("Teacher Dashboard có những tính năng gì?",
         "Giáo viên có thể: tạo và quản lý lớp học, thêm học sinh qua mã class code, tạo bài tập và quiz gắn với bài học cụ thể, theo dõi tiến độ từng học sinh real-time (completion rate, quiz scores, thời gian thực hành), gửi thông báo đến cả lớp. Tất cả được build trên /dashboard/teacher với server-side data fetching từ Supabase."),
        ("Hệ thống Gamification trong PC Master gồm những gì?",
         "PC Master tích hợp: XP Points (kiếm khi hoàn thành bài học, quiz, thực hành 3D), Achievement Badges (huy hiệu đặc biệt cho các mốc như 'Hoàn thành 10 bài', 'Build PC đầu tiên'), Streak Counter (chuỗi ngày học liên tiếp), và Leaderboard trong lớp. Gamification được thiết kế theo Self-Determination Theory — tạo động lực nội tại thay vì chỉ phần thưởng bên ngoài."),
        ("Chứng chỉ hoàn thành khóa học được cấp và xác thực như thế nào?",
         "Khi học sinh hoàn thành 100% bài học và đạt điểm quiz tối thiểu, hệ thống tự động tạo digital certificate có mã hash duy nhất lưu trên Supabase. Chứng chỉ có thể share link công khai để nhà tuyển dụng/giáo viên xác minh tính xác thực qua URL. PDF export dùng thư viện jsPDF render trên client-side."),
        ("Quiz Bank trong PC Master hoạt động như thế nào?",
         "Quiz Bank (/quiz-bank) chứa ngân hàng câu hỏi được phân loại theo chủ đề và độ khó. Giáo viên có thể tạo bài kiểm tra từ pool câu hỏi có sẵn hoặc thêm câu hỏi tùy chỉnh. Học sinh làm bài theo thời gian giới hạn, kết quả được chấm tự động và lưu vào learning history. Câu hỏi hỗ trợ nhiều định dạng: trắc nghiệm, điền vào chỗ trống."),
        ("Tính năng real-time collaboration trong builder hoạt động như thế nào?",
         "Dùng Supabase Realtime channel, nhiều người dùng có thể thấy nhau đang làm gì trong builder session. CollaborationStatus component hiển thị avatar và tên của những người đang online. Hệ thống lock database ngăn hai giáo viên chỉnh sửa cùng một component cùng lúc (lock tự acquire khi mở form, heartbeat 4 phút, auto-release sau 5 phút idle)."),
        ("Notification System trong PC Master hoạt động ra sao?",
         "NotificationBell component subscribe vào Supabase Realtime channel 'notifications'. Khi giáo viên gửi thông báo đến lớp, tất cả học sinh đang online nhận được real-time toast notification không cần refresh trang. Lịch sử thông báo được lưu trong bảng notifications, đánh dấu read/unread."),
        ("Learning Analytics cung cấp insights gì cho giáo viên?",
         "Admin Dashboard và Teacher Dashboard hiển thị: completion rate theo bài học, average quiz score, thời gian trung bình hoàn thành mỗi module, học sinh có nguy cơ bỏ học (at-risk: không đăng nhập >7 ngày), và heat map hoạt động theo giờ/ngày. Data được query từ Supabase với aggregation functions, cache server-side để tránh overload."),
        ("Học sinh có thể học offline không?",
         "Hiện tại PC Master yêu cầu kết nối internet để load 3D assets và sync tiến độ lên Supabase. Roadmap giai đoạn 2 sẽ implement Service Worker để cache bài học đã load, cho phép xem lại nội dung text/video offline. Thực hành 3D và quiz sẽ vẫn yêu cầu online để đảm bảo tính toàn vẹn điểm số."),
        ("Phụ huynh có thể theo dõi tiến độ con học không?",
         "Có. Parent role có dashboard riêng (/dashboard/parent) hiển thị tiến độ học của con: bài học đã hoàn thành, điểm quiz, thời gian học mỗi ngày, và badge đạt được. Phụ huynh liên kết tài khoản với con qua Student ID. Đây là tính năng quan trọng để phụ huynh thấy giá trị khi quyết định mua gói Premium."),
        ("Tính năng AI Assistant (VirtualAssistant) trong builder hỗ trợ những gì?",
         "VirtualAssistant là chatbot tích hợp trong builder, trả lời câu hỏi về: tên và chức năng của linh kiện đang hover, gợi ý bước lắp ráp tiếp theo, giải thích lỗi compatibility (ví dụ: 'CPU này dùng socket AM5, mainboard của bạn dùng LGA1700 — không tương thích'), và tóm tắt lý thuyết từ bài học liên quan."),
        ("Làm thế nào để đảm bảo nội dung bài học cập nhật theo SGK mới?",
         "Nội dung bài học được lưu trong database Supabase dưới dạng JSON có cấu trúc module, không hardcode trong code. Giáo viên có thể chỉnh sửa qua Creator Dashboard (WYSIWYG editor) mà không cần developer. Khi Bộ GD&ĐT cập nhật chương trình, giáo viên tự update nội dung — zero-code content management."),
        ("PC Master có hỗ trợ đa ngôn ngữ không?",
         "Có, hỗ trợ Tiếng Việt và English. Toggle ngôn ngữ qua BurgerMenu, lưu preference vào localStorage. Toàn bộ UI string được quản lý qua i18nData utility — thêm ngôn ngữ mới chỉ cần thêm object translation, không cần sửa component. Mặc định Tiếng Việt vì target market chính là học sinh Việt Nam."),
        ("Làm thế nào để giáo viên tạo bài tập thực hành 3D?",
         "Giáo viên sử dụng Creator Dashboard để định nghĩa assignment: chọn loại task (lắp đủ các linh kiện X, Y, Z vào case; hoặc tìm và sửa lỗi compatibility trong build có sẵn), set thời gian giới hạn và điểm tối thiểu. Khi học sinh submit, hệ thống auto-grade dựa trên slot configuration trong AssemblyStore so với đáp án của giáo viên."),
        ("Dữ liệu học tập của 76 học sinh thử nghiệm cho thấy điều gì?",
         "Pilot với 76 học sinh lớp 10A1 và 10C9 là bằng chứng thực tế (empirical evidence) về product-market fit. Kết quả pilot bao gồm: retention rate sau 2 tuần, điểm quiz trước/sau khi dùng PC Master, và NPS (Net Promoter Score) từ học sinh và giáo viên. Dữ liệu này là nền tảng quan trọng nhất để thuyết phục BGK về khả năng scale."),
    ]
    for i, (q, a) in enumerate(qa3, 1):
        add_question(doc, 30 + i, q, a)

    # ════════════════════════════════════════════════════
    # CHỦ ĐỀ 4 — MÔ HÌNH KINH DOANH & TÀI CHÍNH (15 câu)
    # ════════════════════════════════════════════════════
    add_topic_header(doc, "CHỦ ĐỀ 4: MÔ HÌNH KINH DOANH, FREEMIUM & TÀI CHÍNH (15 câu)")
    qa4 = [
        ("Mô hình Freemium 20.000đ/tháng có đủ để duy trì startup không?",
         "20.000đ/tháng là giá thâm nhập thị trường (penetration pricing), không phải giá cuối cùng. Với 10.000 học sinh Premium, ARR = 2.4 tỷ đồng. Chi phí vận hành chủ yếu là Supabase Pro (~20$/tháng), Vercel Pro (~20$/tháng) và domain — tổng chi phí server dưới 5 triệu đồng/tháng ở quy mô này. Biên lợi nhuận gross margin >90% — điển hình của SaaS EdTech."),
        ("Tại sao không bán theo khóa học (one-time) mà dùng subscription?",
         "Subscription (recurring revenue) tạo ARR ổn định, dễ dự báo cash flow, và khuyến khích nhóm phát triển cải tiến liên tục để retain subscribers. One-time payment không tạo incentive để maintain và update sản phẩm lâu dài. MRR (Monthly Recurring Revenue) là metric được investor EdTech ưa chuộng nhất."),
        ("Gói Premium và Free khác nhau ở chỗ nào cụ thể?",
         "Free: thực hành 3D cơ bản, 5 bài học đầu tiên, quiz giới hạn 10 câu/ngày, xem giá linh kiện. Premium: toàn bộ 20 chương SGK Tin học 10, AI Compatibility Checker không giới hạn, VR Lab với Face/Hand Tracking, chứng chỉ hoàn thành, priority support, không giới hạn quiz. Sự khác biệt rõ ràng là yếu tố quyết định conversion từ Free sang Premium."),
        ("Doanh thu B2B từ School License được tính như thế nào?",
         "School License cho phép trường mua license cho toàn bộ học sinh với giá bulk: ước tính 5.000đ/học sinh/tháng cho đơn hàng >200 học sinh. Trường 1.000 học sinh = 5 triệu đồng/tháng = 60 triệu đồng/năm. Một trường ký hợp đồng tương đương 300 premium individual subscribers. B2B là kênh scale nhanh hơn B2C."),
        ("Chi phí Customer Acquisition Cost (CAC) của PC Master là bao nhiêu?",
         "CAC thấp nhờ chiến lược organic growth: giáo viên dùng thử → chia sẻ với đồng nghiệp (viral loop trong trường học). Chi phí marketing ban đầu gần như zero — không chạy ads. CAC ước tính: thời gian demo cho 1 giáo viên (~2 giờ) × số trường tiếp cận. Khi có social proof từ pilot 76 học sinh, conversion rate tăng đáng kể."),
        ("LTV (Customer Lifetime Value) của một học sinh Premium là bao nhiêu?",
         "Học sinh THPT dùng PC Master trong 3 năm (lớp 10-12): LTV = 20.000đ × 36 tháng = 720.000đ. Nếu sau khi tốt nghiệp tiếp tục dùng cho đại học hoặc giới thiệu cho người khác, LTV tăng thêm. LTV/CAC ratio >3 là chỉ số startup EdTech khỏe mạnh — PC Master có thể đạt được với mô hình organic growth."),
        ("Break-even point của dự án là khi nào?",
         "Chi phí cố định hàng tháng (server, domain, tools): ~5 triệu đồng. Với giá 20.000đ/Premium user: break-even tại 250 subscribers. Đây là mục tiêu hoàn toàn khả thi trong 6 tháng đầu nếu pilot 76 học sinh chuyển đổi và mỗi người giới thiệu thêm 3 bạn."),
        ("Nếu có nhà đầu tư, vốn sẽ được dùng vào việc gì?",
         "Ưu tiên: (1) Marketing & Sales để mở rộng sang 10 trường THPT tại TP.HCM — 60% vốn; (2) Phát triển tính năng AI nâng cao (computer vision cải thiện, gợi ý học tập cá nhân hóa) — 30% vốn; (3) Infrastructure scale (Supabase enterprise, CDN cho 3D assets) — 10% vốn. Không dùng vào chi phí cố định hay lương sớm."),
        ("Tại sao không đặt giá cao hơn — ví dụ 100.000đ/tháng?",
         "Price sensitivity trong phân khúc học sinh THPT rất cao — phụ huynh so sánh với giá sách giáo khoa (20-50k) và phí học thêm (200-500k/môn). 20.000đ/tháng nằm trong 'impulse purchase zone' — phụ huynh có thể quyết định mà không cần họp gia đình. Sau khi xây dựng brand trust, có thể tăng giá 30-50% kèm thêm value (roadmap năm 2)."),
        ("Mô hình kinh doanh có phụ thuộc vào một nguồn doanh thu duy nhất không?",
         "Không. PC Master có 3 dòng doanh thu độc lập: (1) B2C Premium subscriptions (học sinh/phụ huynh), (2) B2B School Licenses (trường học), và (3) Tương lai: Affiliate commission khi học sinh mua linh kiện qua link trong AI recommendation. Đa dạng hóa revenue stream giúp giảm rủi ro business model."),
        ("Tại sao không seek investment ngay từ đầu mà tự bootstrap?",
         "Bootstrap giai đoạn đầu để validate product-market fit với chi phí thấp nhất. Nếu raise funding khi chưa có user thực, nhóm sẽ bị áp lực 'spend to grow' trước khi hiểu rõ unit economics. Pilot 76 học sinh + break-even sớm là bằng chứng execution capability — sau đó raise seed round ở valuation tốt hơn."),
        ("Kế hoạch mở rộng ra ngoài TP.HCM như thế nào?",
         "Mở rộng theo cụm địa lý: TP.HCM → Hà Nội → Đà Nẵng → các tỉnh thành. Kênh mở rộng: hợp tác với Sở GD&ĐT để triển khai pilot cấp sở, tham gia hội thảo giáo dục STEM quốc gia, và referral program cho giáo viên (commission khi giới thiệu trường mới). Digital-first product không bị giới hạn địa lý — mọi trường có internet là khách hàng tiềm năng."),
        ("Rủi ro lớn nhất về mặt tài chính của dự án là gì và cách giảm thiểu?",
         "Rủi ro chính: churn rate cao nếu học sinh không thấy giá trị rõ ràng sau tháng đầu. Giảm thiểu bằng: onboarding sequence 7 ngày đầu (daily email/notification hướng dẫn tính năng core), gamification để tạo thói quen học hàng ngày (streak counter), và success milestone rõ ràng trong 30 ngày đầu (hoàn thành 3 bài, build PC đầu tiên)."),
        ("Chính sách hoàn tiền như thế nào?",
         "7-day money-back guarantee cho gói Premium — giảm rào cản thử nghiệm. Trong 7 ngày nếu không hài lòng, hoàn tiền 100% không hỏi lý do. Dữ liệu cho thấy người dùng trải nghiệm 3D Lab trong 7 ngày đầu thường không hoàn tiền vì đã tạo được emotional connection với sản phẩm."),
        ("Dự án có kế hoạch IPO hay exit không?",
         "Trong 3-5 năm, mục tiêu là trở thành nền tảng EdTech phần cứng IT hàng đầu Việt Nam và mở rộng sang Đông Nam Á (thị trường có chương trình giáo dục tương tự). Exit có thể thông qua: acquisition bởi tập đoàn giáo dục (FPT Education, Topica), hoặc merge với LMS platform lớn muốn tích hợp domain STEM. Không có kế hoạch IPO trong 5 năm đầu."),
    ]
    for i, (q, a) in enumerate(qa4, 1):
        add_question(doc, 45 + i, q, a)

    # ════════════════════════════════════════════════════
    # CHỦ ĐỀ 5 — BẢO MẬT, PHÁP LÝ & GDPR (15 câu)
    # ════════════════════════════════════════════════════
    add_topic_header(doc, "CHỦ ĐỀ 5: BẢO MẬT DỮ LIỆU, RLS, PHÁP LÝ & TUÂN THỦ (15 câu)")
    qa5 = [
        ("Dữ liệu học sinh được bảo vệ như thế nào?",
         "Row Level Security (RLS) của PostgreSQL/Supabase đảm bảo mỗi học sinh chỉ đọc/ghi được dữ liệu của chính họ. Policy: 'Users can only access their own profile' được enforce ở tầng database, không phải chỉ ở tầng application — ngay cả nếu code có lỗi, database từ chối query unauthorized."),
        ("Camera và dữ liệu face tracking có được lưu trữ không?",
         "Không. Face tracking và hand tracking xử lý hoàn toàn on-device (client-side) bằng MediaPipe WebAssembly. Landmarks (tọa độ điểm khuôn mặt/tay) chỉ tồn tại trong RAM của trình duyệt người dùng, không bao giờ được gửi lên server. Camera stream cũng không được record hay store. Điều này quan trọng khi đối tượng người dùng là trẻ vị thành niên."),
        ("PC Master tuân thủ những quy định pháp lý nào về dữ liệu?",
         "Tuân thủ Nghị định 13/2023/NĐ-CP về bảo vệ dữ liệu cá nhân (PDPD) của Việt Nam và GDPR (cho user quốc tế). Cụ thể: explicit consent khi đăng ký, quyền xóa tài khoản (right to erasure), không chia sẻ dữ liệu với bên thứ ba, và Privacy Policy được hiển thị rõ trước khi đăng ký."),
        ("Xác thực người dùng (Authentication) được xây dựng như thế nào?",
         "Dùng Supabase Auth với: Email/Password (bcrypt hashed), Google OAuth 2.0. JWT token được lưu trong httpOnly cookie (không accessible qua JavaScript) để chống XSS. Session refresh tự động. Sau khi đăng nhập Google lần đầu, người dùng được redirect đến /onboarding để chọn role (student/teacher/parent/admin) và điền thông tin profile."),
        ("Database trigger handle_new_user hoạt động như thế nào?",
         "Trigger PostgreSQL tự động tạo bản ghi trong bảng profiles khi có user mới trong auth.users. Function handle_new_user() đọc metadata từ Google OAuth (tên, email, avatar) và ghi vào profiles với role mặc định 'student'. Nếu trigger lỗi, đăng ký bị rollback hoàn toàn — đảm bảo không có user 'zombie' trong auth mà không có profile."),
        ("Làm thế nào để chống SQL Injection?",
         "Supabase client library sử dụng parameterized queries tự động — không bao giờ concatenate string vào SQL query. Ngoài ra, RLS policy chặn mọi query cố tình bypass tầng application. Không có raw SQL query được xây dựng từ user input trong codebase."),
        ("CORS và API security được xử lý như thế nào?",
         "Next.js API routes chỉ cho phép request từ origin được whitelist (domain production và localhost). Supabase anon key chỉ có quyền đọc data công khai; mọi thao tác sensitive phải qua server-side với service role key không expose ra client. Environment variables nhạy cảm được lưu trong .env.local và Vercel environment variables — không commit vào git."),
        ("Làm thế nào để xử lý khi tài khoản học sinh bị hack?",
         "Người dùng có thể đổi mật khẩu qua email reset link (Supabase gửi magic link). Admin có thể force-logout tất cả session của một user. Suspicious login detection (đăng nhập từ IP lạ) sẽ được implement trong roadmap. Hiện tại 2FA chưa có nhưng được plan cho gói School License để tăng security."),
        ("Ứng dụng có vulnerable với XSS attacks không?",
         "React mặc định escape tất cả dynamic content trong JSX, ngăn XSS. Không dùng dangerouslySetInnerHTML trừ những nơi được sanitize qua DOMPurify. Cookie httpOnly cho JWT ngăn script đọc token. Content Security Policy headers được set trong Next.js config để thêm lớp bảo vệ."),
        ("Vercel deployment có đảm bảo uptime không?",
         "Vercel SLA là 99.99% uptime cho production. Static assets được serve qua Vercel Edge Network (CDN global) — latency thấp cho người dùng Việt Nam nhờ edge node tại Singapore. Supabase cloud cũng có SLA 99.9%. Không có single point of failure vì cả frontend (Vercel) và backend (Supabase) đều là managed services với redundancy built-in."),
        ("Dữ liệu có được backup không?",
         "Supabase Pro plan tự động backup daily, giữ 7 ngày. Point-in-time recovery cho phép restore về bất kỳ thời điểm nào trong 7 ngày. Critical data (profiles, course progress, certificates) được backup thêm bằng cách export định kỳ sang Supabase Storage bucket dưới dạng JSON."),
        ("Ứng dụng có được test security chưa?",
         "Đã thực hiện: OWASP Top 10 checklist review, manual penetration testing trên các API endpoint (unauthorized access, IDOR), và dependency audit (npm audit) để phát hiện vulnerable packages. Chưa có penetration testing chuyên nghiệp — plan thuê security firm khi raise funding."),
        ("Điều khoản sử dụng (ToS) có bảo vệ dự án trước rủi ro pháp lý không?",
         "ToS quy định rõ: PC Master chỉ cung cấp môi trường học tập giả lập (không đảm bảo kết quả thực tế), người dùng không được share account, và quyền sở hữu trí tuệ của nội dung thuộc PC Master. Điều khoản được tư vấn để phù hợp với Luật Công nghệ thông tin và Luật Giáo dục Việt Nam."),
        ("Phần mềm mã nguồn mở nào được sử dụng và license có vấn đề không?",
         "Tất cả dependencies dùng license cho phép commercial use: Next.js (MIT), React Three Fiber (MIT), MediaPipe (Apache 2.0), Supabase client (MIT), Zustand (MIT), Lucide React (ISC). Không có GPL-licensed library nào trong stack — tránh copyleft obligation. License audit được thực hiện qua 'license-checker' npm package."),
        ("Nếu Supabase bị đóng cửa, dữ liệu có mất không?",
         "Supabase là open-source — có thể self-host toàn bộ stack (PostgreSQL + PostgREST + GoTrue auth) trên VPS bất kỳ. Backup định kỳ sang format PostgreSQL dump cho phép migrate sang bất kỳ PostgreSQL hosting nào trong vài giờ. Vendor lock-in risk được giảm thiểu bởi database schema và data ownership hoàn toàn thuộc về PC Master."),
    ]
    for i, (q, a) in enumerate(qa5, 1):
        add_question(doc, 60 + i, q, a)

    # ════════════════════════════════════════════════════
    # CHỦ ĐỀ 6 — TEAM, EXECUTION & CẠNH TRANH (15 câu)
    # ════════════════════════════════════════════════════
    add_topic_header(doc, "CHỦ ĐỀ 6: TEAM, NĂNG LỰC THỰC THI & PHÂN TÍCH CẠNH TRANH (15 câu)")
    qa6 = [
        ("Team hiện tại có đủ năng lực để phát triển sản phẩm này không?",
         "PC Master được xây dựng bởi team sinh viên HUIT với full-stack capability: Next.js/React (frontend), Supabase/PostgreSQL (backend), Three.js/MediaPipe (3D/CV), và Python (data/ML). Sản phẩm đã hoạt động production tại pc-master-lms.vercel.app với người dùng thực — đây là bằng chứng execution, không chỉ là ý tưởng."),
        ("Tại sao một team sinh viên có thể build được sản phẩm phức tạp như thế này?",
         "Công nghệ hiện đại (Next.js, Supabase, MediaPipe) giúp bootstrap một sản phẩm phức tạp nhanh hơn nhiều so với 5 năm trước. Supabase thay thế backend team, Vercel thay thế DevOps team, Lucide thay thế design team cho icons. Team tập trung vào business logic và UX độc đáo — đây là cách startup EdTech hiện đại vận hành hiệu quả."),
        ("Chiến lược phân công công việc trong team như thế nào?",
         "Team dùng git branching strategy: main branch là production (auto-deploy lên Vercel), mỗi tính năng lớn được phát triển trên feature branch riêng. Hai track song song: feature/2d-renderer (2D builder & animations) và feature/3d-viewer (3D/VR/Computer Vision). Code review qua Pull Request trước khi merge vào main — tránh conflict và đảm bảo code quality."),
        ("PC Master có bằng sáng chế hay sở hữu trí tuệ gì đặc biệt không?",
         "Hiện chưa có patent. Competitive moat đến từ: (1) tốc độ execution và first-mover advantage trong mảng LMS + 3D lab phần cứng tại Việt Nam, (2) dataset linh kiện và giá thị trường Việt Nam được xây dựng dần theo thời gian, (3) network effect từ giáo viên và trường học đã dùng. Barrier to entry là community và data, không chỉ là technology."),
        ("Nếu một thành viên rời team, dự án có bị ảnh hưởng không?",
         "Codebase có documentation rõ ràng qua AGENTS.md và README. Architecture được thiết kế module — 3D viewer, LMS, auth là các phần độc lập. Không có knowledge siloed ở một người — tất cả được commit trên git với history đầy đủ. Onboarding developer mới có thể productive trong 1-2 ngày nhờ documentation và codebase clean."),
        ("Đối thủ cạnh tranh nào mạnh nhất và cách đối phó?",
         "Rủi ro lớn nhất: Google ra mắt tính năng tương tự trên Google Classroom (vốn có sẵn hàng triệu học sinh Việt Nam). Đối phó: (1) speed — tích hợp tính năng mới nhanh hơn BigTech vì không có internal bureaucracy, (2) specialization — domain expertise về phần cứng PC Việt Nam mà BigTech không có incentive để build, (3) relationship — trust từ giáo viên và trường học đã dùng."),
        ("Tại sao PC Master chưa bị copy bởi các công ty lớn?",
         "Thị trường này nhỏ theo góc nhìn BigTech (dưới 1M USD market size ban đầu) nhưng đủ lớn và profitable cho startup. BigTech ưu tiên horizontal platforms, không vertical domain-specific tools. Đây là chiến lược niche market — tìm thị trường 'đủ lớn để sống, đủ nhỏ để không bị BigTech quan tâm ngay'."),
        ("Làm thế nào để giữ chân (retain) giáo viên dùng platform?",
         "Switching cost cao do: (1) giáo viên đã đầu tư thời gian tạo bài tập và quiz trên PC Master — không muốn bỏ đi, (2) học sinh của họ đã có tài khoản và lịch sử học, (3) analytics và báo cáo tiến độ học sinh giúp giáo viên chứng minh hiệu quả giảng dạy với ban giám hiệu. Giáo viên trở thành champion nội bộ của sản phẩm."),
        ("Tốc độ phát triển tính năng của team là bao nhiêu?",
         "Từ ý tưởng đến MVP hoạt động trong production: dưới 3 tháng. Tính năng 3D viewer với face tracking được hoàn thiện trong 2 tuần. Sprint cycle khoảng 1 tuần. Deploy lên Vercel tự động khi push lên main branch — zero-downtime deployment. Tốc độ này đạt được nhờ managed services (Supabase, Vercel) loại bỏ hầu hết infrastructure work."),
        ("Chiến lược marketing ban đầu của PC Master là gì?",
         "Content marketing về giáo dục STEM (blog, TikTok demo 3D lab), direct outreach đến giáo viên Tin học qua Facebook groups giáo viên THPT, và product-led growth — giáo viên share link cho học sinh, học sinh share cho bạn. Không có paid ads trong 6 tháng đầu. Chiến lược này phù hợp với zero-budget startup, leverage strong word-of-mouth trong cộng đồng giáo dục."),
        ("PC Master có plan gì nếu Vercel hay Supabase tăng giá đột ngột?",
         "Vì Supabase là open-source và Next.js chạy được trên bất kỳ Node.js server nào, PC Master có thể migrate sang self-hosted trên VPS (DigitalOcean, AWS EC2) mà không cần rewrite code. Chi phí VPS scale tuyến tính và rẻ hơn Vercel/Supabase Pro ở quy mô lớn. Migration được estimate trong 1-2 tuần engineer time."),
        ("Nhóm có plan gì sau khi cuộc thi kết thúc?",
         "PC Master là sản phẩm thực đang có người dùng thực, không phải project cuộc thi. Plan: tiếp tục iterate dựa trên feedback từ 76 học sinh pilot, đạt break-even trong Q4 2026, và pitch cho seed investors cuối 2026. Cuộc thi HUIT là milestone công nhận, không phải điểm kết thúc."),
        ("Tính năng nào sẽ được phát triển trong 6 tháng tới?",
         "Roadmap: (1) AI học cá nhân hóa — gợi ý bài học dựa trên learning path của từng học sinh, (2) Multiplayer lab — 2 học sinh cùng lắp ráp PC trong không gian 3D chung, (3) AR mode — dùng camera điện thoại để 'đặt' linh kiện ảo lên bàn thực (WebXR), (4) Certification exam được công nhận bởi các đối tác doanh nghiệp IT."),
        ("Tính năng VoiceController được dùng như thế nào?",
         "VoiceController cho phép điều khiển builder bằng giọng nói (Web Speech API): 'đặt CPU', 'xóa RAM', 'boot PC', 'quay lại menu'. Đây là accessibility feature quan trọng — học sinh khiếm khuyết tay có thể thực hành lắp ráp PC. Đồng thời làm WOW factor trong demo trực tiếp trước BGK."),
        ("Tại sao dùng Vercel để deploy thay vì AWS hay Azure?",
         "Vercel được tối ưu cho Next.js (cùng công ty tạo ra Next.js) — auto-optimize image, zero-config deployment, preview URL cho mỗi PR, edge functions gần người dùng. AWS/Azure đòi hỏi DevOps expertise không phù hợp giai đoạn startup. Vercel Pro tốn 20$/tháng nhưng thay thế 1 DevOps engineer full-time — ROI cực kỳ cao."),
    ]
    for i, (q, a) in enumerate(qa6, 1):
        add_question(doc, 75 + i, q, a)

    # ════════════════════════════════════════════════════
    # CHỦ ĐỀ 7 — TÁC ĐỘNG XÃ HỘI, TƯƠNG LAI & ĐỘT PHÁ (15 câu)
    # ════════════════════════════════════════════════════
    add_topic_header(doc, "CHỦ ĐỀ 7: TÁC ĐỘNG XÃ HỘI, TƯƠNG LAI & ĐỘT PHÁ GIÁO DỤC (15 câu)")
    qa7 = [
        ("PC Master đóng góp gì cho giáo dục STEM Việt Nam?",
         "PC Master làm cho giáo dục STEM thực hành phần cứng trở nên accessible cho mọi học sinh — không chỉ những trường có điều kiện đầu tư phòng lab hàng trăm triệu đồng. Mỗi học sinh với smartphone có thể trải nghiệm lắp ráp PC 3D trong 5 phút. Đây là democratization of hands-on STEM education."),
        ("Dự án giúp giảm bất bình đẳng giáo dục như thế nào?",
         "Học sinh ở tỉnh lẻ hoặc trường thiếu thiết bị có cùng trải nghiệm thực hành như học sinh trường điểm TP.HCM thông qua Virtual Lab 3D. Gói miễn phí đảm bảo học sinh khó khăn không bị loại khỏi cơ hội học tập kỹ thuật số. Đây là social impact trực tiếp đo được qua số lượng học sinh truy cập từ các tỉnh."),
        ("Tác động đến giáo viên Tin học là gì?",
         "Giáo viên tiết kiệm hàng giờ soạn bài mỗi tuần nhờ Creator Dashboard. Analytics real-time giúp giáo viên identify học sinh đang gặp khó khăn sớm hơn. Giáo viên không cần lo về thiết bị hỏng hay an toàn trong lab. PC Master nâng cao hiệu quả giảng dạy mà không tăng workload."),
        ("PC Master có thể mở rộng sang các môn học khác không?",
         "Architecture của LMS được thiết kế domain-agnostic — phần LMS (quản lý lớp, quiz, analytics) có thể áp dụng cho bất kỳ môn nào. Phần 3D Lab là domain-specific. Tương lai có thể mở rộng sang: lab Vật lý (mô phỏng mạch điện), lab Hóa học (thí nghiệm ảo), hoặc lab Điện tử (lắp ráp Arduino/Raspberry Pi ảo)."),
        ("Công nghệ WebXR và AR có trong roadmap không?",
         "Có. Sau khi WebXR AR API được support rộng hơn trên Android Chrome (hiện ~70% devices), PC Master sẽ triển khai AR mode: camera điện thoại chiếu linh kiện ảo lên mặt bàn thực, học sinh tương tác bằng tay. Đây là natural evolution từ face/hand tracking đã có — không cần rebuild architecture, chỉ thêm WebXR renderer layer."),
        ("Multiplayer learning có khả thi về mặt kỹ thuật không?",
         "Có. Supabase Realtime đã cung cấp infrastructure WebSocket. Cần thêm: CRDT (Conflict-free Replicated Data Type) để sync trạng thái scene 3D giữa nhiều client, và authority server để resolve conflict khi 2 người cùng di chuyển cùng component. Estimate: 3-4 tuần implementation. Tính năng này tạo collaborative learning experience hoàn toàn mới."),
        ("Tầm nhìn 5 năm của PC Master là gì?",
         "Trở thành nền tảng học phần cứng IT số 1 Đông Nam Á, phục vụ >500,000 học sinh tại Việt Nam, Thái Lan, Indonesia. Mở rộng danh mục: PC → networking (lab ảo cắm dây mạng), server rack, IoT devices. Hợp tác với nhà sản xuất linh kiện (Intel, AMD, ASUS) để có official 3D models và data giá. Revenue target năm 5: 100 tỷ đồng ARR."),
        ("Công nghệ AI sẽ thay đổi PC Master như thế nào trong tương lai?",
         "AI/ML sẽ enable: (1) Adaptive learning path — AI phân tích learning style của từng học sinh và tự động điều chỉnh thứ tự bài học, (2) Natural language interaction — học sinh nói chuyện với AI về linh kiện như nói với chuyên gia thật, (3) Predictive analytics — dự đoán học sinh nào sẽ bỏ học 2 tuần trước khi xảy ra để can thiệp kịp thời."),
        ("PC Master có thể đóng góp vào nguồn nhân lực IT của Việt Nam không?",
         "Học sinh được đào tạo từ lớp 10 có foundation kỹ thuật tốt về phần cứng — giảm thời gian đào tạo lại ở Đại học và doanh nghiệp. Chứng chỉ PC Master được plan là recognized bởi các đối tác doanh nghiệp IT — tạo pathway từ học phổ thông đến việc làm kỹ thuật. Đây là long-term social ROI quan trọng."),
        ("Nếu được hỗ trợ bởi Bộ GD&ĐT, PC Master có thể làm gì?",
         "Với endorsement chính phủ: (1) tích hợp vào chương trình học chính thức — giảm CAC về gần zero, (2) cấp phép triển khai toàn quốc theo hình thức PPP (Public-Private Partnership), (3) dùng dữ liệu aggregated (anonymized) để cải thiện chương trình giáo dục STEM quốc gia. Đây là tầm nhìn impact at scale thực sự."),
        ("Làm thế nào để đảm bảo chất lượng khi scale lên hàng nghìn trường?",
         "Platform architecture đã được thiết kế stateless và horizontally scalable từ đầu. Supabase auto-scales database connections. Vercel Edge Network phân phối tải. Content moderation cho creator-generated quiz được thực hiện qua: AI filter (trước) + teacher review (sau). SLA và support tier khác nhau cho gói Free vs School License."),
        ("Dự án có tác động gì đến ngành bán lẻ linh kiện máy tính Việt Nam?",
         "Tính năng AI giá real-time tạo market transparency — người mua có thể so sánh giá ngay tại cửa hàng. Về lâu dài, điều này tạo pressure để cửa hàng cạnh tranh bằng giá trị (tư vấn chuyên nghiệp, bảo hành tốt) thay vì information asymmetry. Đây là positive disruption cho thị trường bán lẻ linh kiện."),
        ("Học sinh hoàn thành khóa học trên PC Master có được công nhận chính thức không?",
         "Chứng chỉ digital của PC Master hiện là soft credential, không có giá trị pháp lý như bằng cấp nhà nước. Roadmap: hợp tác với CompTIA hoặc các đối tác IT để có joint certification, và vận động Bộ GD&ĐT công nhận là chứng chỉ bổ sung (extracurricular credential). Đây là process 2-3 năm cần trust building."),
        ("Tính năng gì trong PC Master tạo 'wow effect' nhất với học sinh?",
         "Face Tracking VR Lab — học sinh thấy scene 3D thay đổi theo hướng nhìn của mình qua webcam, không cần kính VR, không cần app cài đặt. Moment đó — lần đầu tiên nhìn vào màn hình và thấy không gian 3D phản ứng với đầu mình — tạo emotional connection mạnh mẽ với sản phẩm. Đây là 'magic moment' quan trọng nhất trong onboarding."),
        ("PC Master định nghĩa lại 'phòng học STEM' như thế nào?",
         "Phòng học STEM truyền thống cần 4 tường + thiết bị đắt tiền. PC Master biến bất kỳ thiết bị có trình duyệt nào thành lab STEM đầy đủ chức năng: giả lập, thực hành, kiểm tra, ghi nhận kết quả. Classroom không còn bị giới hạn bởi không gian vật lý hay ngân sách thiết bị. Đây là paradigm shift trong giáo dục kỹ thuật."),
    ]
    for i, (q, a) in enumerate(qa7, 1):
        add_question(doc, 90 + i, q, a)

    # ── SAVE ────────────────────────────────────────────
    out_path = r"c:\Users\fujitsu\Downloads\khu phố\pc-master-lms-latest\public\HUIT_105_Cau_Hoi_Phan_Bien_PC_Master.docx"
    doc.save(out_path)
    print(f"✅ Đã tạo file Word: {out_path}")
    print(f"   Tổng: 105 câu hỏi | 7 chủ đề | Không có ABCD | Chỉ hướng trả lời")

build_doc()
