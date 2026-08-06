import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def create_document():
    doc = docx.Document()

    # Set Margins
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # Styles & Colors
    COLOR_NAVY = RGBColor(3, 31, 59)      # #031F3B
    COLOR_PRIMARY = RGBColor(8, 158, 96)   # #089E60
    COLOR_DARK = RGBColor(30, 41, 59)     # #1E293B
    COLOR_GRAY = RGBColor(100, 116, 139)  # #64748B
    COLOR_CORRECT = RGBColor(16, 185, 129)# #10B981
    COLOR_RED = RGBColor(239, 68, 68)     # #EF4444

    # Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("NGÂN HÀNG 105 CÂU HỎI PHẢN BIỆN DỰ ĐOÁN\nBAN GIÁM KHẢO CUỘC THI KHỞI NGHIỆP HUIT 2026")
    run_title.font.name = 'Arial'
    run_title.font.size = Pt(20)
    run_title.font.bold = True
    run_title.font.color.rgb = COLOR_NAVY

    # Subtitle
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = p_sub.add_run("DỰ ÁN: PC MASTER BUILDER - HỆ THỐNG LMS & GIẢ LẬP LẮP RÁP PC 3D/AI\nTrường Đại học Công Thương TP.HCM (HUIT) | Phân tích Đáp án & Hướng trả lời thuyết phục BGK")
    run_sub.font.name = 'Arial'
    run_sub.font.size = Pt(11)
    run_sub.font.italic = True
    run_sub.font.color.rgb = COLOR_PRIMARY

    doc.add_paragraph() # spacing

    # Overview Box
    table_intro = doc.add_table(rows=1, cols=1)
    table_intro.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table_intro.cell(0, 0)
    cell.width = Inches(6.8)
    
    # Set shading / background color
    shading_elm = parse_xml(r'<w:shd {} w:fill="F1F5F9"/>'.format(nsdecls('w')))
    cell._tc.get_or_add_tcPr().append(shading_elm)

    p_intro = cell.paragraphs[0]
    p_intro.paragraph_format.space_before = Pt(6)
    p_intro.paragraph_format.space_after = Pt(6)
    r_intro = p_intro.add_run(
        "📌 HƯỚNG DẪN SỬ DỤNG CHO ĐỘI THI PC MASTER BUILDER:\n"
        "Bộ câu hỏi này được tổng hợp và thiết kế từ góc nhìn của các Chuyên gia Phản biện & Ban Giám khảo Cuộc thi Khởi nghiệp HUIT. "
        "Mỗi câu hỏi kèm theo 4 phương án lựa chọn, ĐÁP ÁN CHUẨN THUYẾT PHỤC BGK NHẤT và PHẦN GIẢI THÍCH CHI TIẾT nêu rõ lý do tại sao phương án đó tối ưu, cũng như những bẫy/lỗi sai thường gặp ở 3 phương án còn lại.\n"
        "Các chủ đề bao gồm: Vấn đề thị trường, Giải pháp công nghệ 3D/AI, Mô hình kinh doanh 20k/tháng, Thử nghiệm 76 học sinh 10A1/10C9, Bảo mật Supabase, Tài chính & Kế hoạch mở rộng."
    )
    r_intro.font.name = 'Arial'
    r_intro.font.size = Pt(10)
    r_intro.font.color.rgb = COLOR_DARK

    doc.add_paragraph()

    # Question generator script logic
    # We will write out all 105 structured questions across 7 topics
    
    topics = [
        {
            "id": 1,
            "name": "CHỦ ĐỀ 1: VẤN ĐỀ THỊ TRƯỜNG, ĐỘNG LỰC DỰ ÁN & NHU CẦU NGUỜI DÙNG (15 CÂU)",
            "questions": [
                {
                    "q": "Câu 1: Giám khảo HUIT hỏi: 'Hiện nay có rất nhiều trang web hướng dẫn lắp ráp máy tính trên YouTube và Google, vậy tại sao học sinh vẫn cần dự án PC Master LMS của các bạn?'",
                    "options": [
                        "A. Vì YouTube và Google chứa quá nhiều quảng cáo rác.",
                        "B. Vì các video trên mạng chỉ mang tính thụ động, học sinh xem xong không được tự tay thực hành giả lập 3D, không có bài kiểm tra đánh giá theo chương trình SGK Tin học 10.",
                        "C. Vì dự án PC Master LMS miễn phí 100% tất cả tính năng cho mọi người dùng.",
                        "D. Vì ứng dụng của chúng tôi chạy mượt hơn YouTube trên điện thoại cùi."
                    ],
                    "correct": "B. Vì các video trên mạng chỉ mang tính thụ động, học sinh xem xong không được tự tay thực hành giả lập 3D, không có bài kiểm tra đánh giá theo chương trình SGK Tin học 10.",
                    "explain": "ĐÁP ÁN ĐÚNG LÀ B.\n- Lý do đúng: BGK Startup luôn đánh giá cao tính tương tác (Interactive Learning) và sự gắn kết với khung chương trình giáo dục phổ thông (SGK Tin học 10 mới). Việc học qua video YouTube chỉ là tiếp thu một chiều (Passive learning), người học dễ quên và không thể tự thao tác kiểm tra độ tương thích linh kiện như mô phỏng 3D của PC Master.\n- Phân tích đáp án sai:\n + A sai vì quảng cáo không phải là nỗi đau (pain point) cốt lõi của việc giáo dục phần cứng.\n + C sai vì dự án có mô hình Freemium (tính phí Premium 20k cho AI/VR), nếu trả lời miễn phí 100% sẽ tự mâu thuẫn với Mô hình kinh doanh.\n + D sai vì mang tính cảm tính, không dựa trên bằng chứng kỹ thuật hay giá trị giáo dục."
                },
                {
                    "q": "Câu 2: Giám khảo HUIT hỏi: 'Vấn đề nạn chặt chém giá linh kiện máy tính mà các bạn đề cập có thực sự cấp thiết đối với học sinh cấp 3 hay không?'",
                    "options": [
                        "A. Không cấp thiết vì học sinh không có tiền mua máy tính.",
                        "B. Rất cấp thiết vì đa số phụ huynh và học sinh thiếu kiến thức kỹ thuật, dễ bị cửa hàng tư vấn linh kiện cũ giá cao hoặc không tương thích, gây lãng phí hàng triệu đồng.",
                        "C. Chỉ cấp thiết với các thợ sửa chữa máy tính chuyên nghiệp.",
                        "D. Không quan trọng vì các cửa hàng lớn như FPT hay Thế Giới Di Động đã niêm yết giá cố định."
                    ],
                    "correct": "B. Rất cấp thiết vì đa số phụ huynh và học sinh thiếu kiến thức kỹ thuật, dễ bị cửa hàng tư vấn linh kiện cũ giá cao hoặc không tương thích, gây lãng phí hàng triệu đồng.",
                    "explain": "ĐÁP ÁN ĐÚNG LÀ B.\n- Lý do đúng: Học sinh và phụ huynh khi tự mua/nâng cấp PC phục vụ học tập rất dễ gặp rủi ro 'bất đồng thông tin' (Information asymmetry). Việc PC Master tích hợp tính năng Tra giá linh kiện real-time và AI chống chặt chém giúp người dùng kiểm tra ngay giá thị trường chuẩn và độ tương thích linh kiện.\n- Phân tích đáp án sai:\n + A & C sai vì đánh giá thấp nhu cầu thực tế của thị trường EduTech & tư vấn tiêu dùng.\n + D sai vì các cửa hàng nhỏ lẻ hoặc mua bán linh kiện cũ/sang tay chiếm thị phần lớn và là nơi xảy ra nạn chênh giá nhiều nhất."
                },
                {
                    "q": "Câu 3: Giám khảo HUIT hỏi: 'Điểm nghẽn lớn nhất trong việc giảng dạy môn Tin học phần cứng tại các trường THPT hiện nay là gì?'",
                    "options": [
                        "A. Học sinh không thích học môn Tin học.",
                        "B. Thiếu phòng Lab phần cứng thực tế do chi phí trang bị linh kiện thật quá cao và rủi ro học sinh làm hỏng hóc thiết bị trong quá trình thực hành.",
                        "C. Giáo viên không có đủ trình độ để dạy lắp ráp máy tính.",
                        "D. Chương trình SGK không yêu cầu học sinh học về phần cứng."
                    ],
                    "correct": "B. Thiếu phòng Lab phần cứng thực tế do chi phí trang bị linh kiện thật quá cao và rủi ro học sinh làm hỏng hóc thiết bị trong quá trình thực hành.",
                    "explain": "ĐÁP ÁN ĐÚNG LÀ B.\n- Lý do đúng: Đây chính là Nỗi đau cốt lõi (Core Pain Point) của ngành giáo dục STEM/Tin học. Mua 40 bộ máy tính cho 1 phòng Lab tốn hàng trăm triệu đồng, chưa kể linh kiện dể gãy chân pin CPU, tĩnh điện hư RAM khi học sinh thao tác. Giải pháp giả lập 3D Lab của PC Master giải quyết triệt để bài toán chi phí và an toàn thiết bị.\n- Phân tích đáp án sai:\n + A sai vì học sinh rất thích học thực hành nếu có công cụ trực quan hấp dẫn.\n + C sai vì chạm tới định kiến không tốt về giáo viên, thiếu tính thuyết phục sư phạm.\n + D sai vì SGK Tin học 10 mới (Bộ Kết nối tri thức/Cánh diều) đều có chuyên đề phần cứng PC."
                },
                {
                    "q": "Câu 4: Giám khảo HUIT hỏi: 'Tại sao dự án lại chọn đối tượng thử nghiệm đầu tiên là học sinh THPT thay vì sinh viên CNTT Đại học?'",
                    "options": [
                        "A. Vì sinh viên Đại học không cần học lắp ráp máy tính.",
                        "B. Vì học sinh THPT bắt đầu tiếp xúc SGK Tin học 10 theo chương trình GDPT 2018, nhu cầu học trực quan cao nhưng thiếu phòng lab phần cứng, giúp dự án dễ kiểm chứng hiệu quả giáo dục nhất.",
                        "C. Vì học sinh THPT dễ bị thuyết phục mua gói trả phí hơn.",
                        "D. Vì nhóm dự án chỉ quen biết các trường THPT."
                    ],
                    "correct": "B. Vì học sinh THPT bắt đầu tiếp xúc SGK Tin học 10 theo chương trình GDPT 2018, nhu cầu học trực quan cao nhưng thiếu phòng lab phần cứng, giúp dự án dễ kiểm chứng hiệu quả giáo dục nhất.",
                    "explain": "ĐÁP ÁN ĐÚNG LÀ B.\n- Lý do đúng: Lựa chọn thị trường mục tiêu ban đầu (Beachhead Market) dựa trên sự thay đổi về chính sách/chương trình giáo dục GDPT 2018. Học sinh THPT là tập người dùng tập trung, rõ ràng về chương trình học.\n- Phân tích đáp án sai:\n + A & D sai vì đưa ra lý do thiếu tính nghiên cứu thị trường chuyên nghiệp.\n + C sai vì học sinh THPT thường phụ thuộc tài chính vào phụ huynh, không phải đối tượng quyết định chi tiền dễ dàng nhất nếu không có giá trị thật."
                },
                {
                    "q": "Câu 5: Giám khảo HUIT hỏi: 'Nếu thị trường có xuất hiện một phần mềm tương tự của nước ngoài, yếu tố nào giúp PC Master giữ vững lợi thế?'",
                    "options": [
                        "A. Chúng tôi bán giá rẻ hơn 10 lần.",
                        "B. PC Master bám sát SGK Tin học Việt Nam, hỗ trợ tiếng Việt 100%, tích hợp AI tra giá linh kiện thị trường Việt Nam và có hệ thống LMS quản lý lớp học dành cho giáo viên Việt Nam.",
                        "C. Phần mềm của chúng tôi không bao giờ gặp lỗi hệ thống.",
                        "D. Chúng tôi sẽ kiện các đối thủ vi phạm bản quyền."
                    ],
                    "correct": "B. PC Master bám sát SGK Tin học Việt Nam, hỗ trợ tiếng Việt 100%, tích hợp AI tra giá linh kiện thị trường Việt Nam và có hệ thống LMS quản lý lớp học dành cho giáo viên Việt Nam.",
                    "explain": "ĐÁP ÁN ĐÚNG LÀ B.\n- Lý do đúng: Đây là Hào khí cạnh tranh (Competitive Moat) về tính bản địa hóa (Localization) và hệ sinh thái toàn diện. Các phần mềm nước ngoài như PC Building Simulator không có tính năng LMS cho giáo viên giao bài tập, không có dữ liệu giá linh kiện tại Việt Nam và không bám chuẩn SGK 2018.\n- Phân tích đáp án sai:\n + A sai vì cuộc chiến về giá (Price war) không phải là chiến lược bền vững cho Startup.\n + C sai vì tuyên bố phi thực tế trong kỹ thuật phần mềm.\n + D sai vì không tập trung vào năng lực nội tại của sản phẩm."
                }
            ]
        }
    ]

    # Additional Topics Data Generator
    # We will generate comprehensive 105 Q&As dynamically in Python
    return doc

print("Python generator template ready.")
