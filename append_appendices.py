# append_appendices.py — Append extra appendix pages to existing dossier
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

COLOR_NAVY = RGBColor(15, 23, 42)
COLOR_BLUE = RGBColor(14, 116, 144)
COLOR_GREEN = RGBColor(16, 185, 129)
COLOR_DARK = RGBColor(51, 65, 85)

doc = docx.Document('Cuon_Thuyet_Minh_Chinh_Thuc_PC_Master_LMS_HUIT_2026_50_Trang.docx')

def set_bg(cell, hex_color):
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shd)

def add_heading(text, size=12, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.font.name = 'Arial'
    r.font.size = Pt(size)
    r.font.bold = True
    r.font.color.rgb = color if color else COLOR_BLUE

def add_para(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.2
    r = p.add_run(text)
    r.font.name = 'Arial'
    r.font.size = Pt(10.5)
    r.font.color.rgb = COLOR_DARK

def add_bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.2
    r = p.add_run(text)
    r.font.name = 'Arial'
    r.font.size = Pt(10.5)
    r.font.color.rgb = COLOR_DARK

def add_table(headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = 'Table Grid'
    for i, h in enumerate(headers):
        c = t.cell(0, i)
        set_bg(c, '0F172A')
        pr = c.paragraphs[0]
        pr.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rr = pr.add_run(h)
        rr.font.name = 'Arial'; rr.font.size = Pt(9.5); rr.font.bold = True
        rr.font.color.rgb = RGBColor(255, 255, 255)
    for ri, row in enumerate(rows):
        bg = 'F8FAFC' if ri % 2 == 0 else 'EFF6FF'
        for ci, val in enumerate(row):
            c = t.cell(ri + 1, ci)
            set_bg(c, bg)
            pp = c.paragraphs[0]
            pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            rr = pp.add_run(str(val))
            rr.font.name = 'Arial'; rr.font.size = Pt(9.5)
            rr.font.color.rgb = COLOR_DARK
    doc.add_paragraph()

# ===== PHU LUC D: SWOT =====
doc.add_page_break()
add_heading('PHU LUC D: PHAN TICH SWOT CHI TIET DU AN PC MASTER LMS', 14, COLOR_NAVY)

add_heading('D.1. Diem Manh (Strengths) — Loi The Noi Tai')
strengths = [
    'Cong nghe doc dao va tien phong: PC Master LMS la nen tang duy nhat tai Viet Nam ket hop dong thoi 4 cong nghe bung no (3D WebGL, Hand Tracking, AI Proctoring, VR Mode) trong mot san pham giao duc hoan chinh chay 100% tren trinh duyet Web.',
    'Chi phi van hanh cuc thap nho kien truc Client-Side: Toan bo tinh toan nang (WebGL 3D, MediaPipe AI, TensorFlow.js) thuc thi truc tiep tren thiet bi nguoi dung. Bien loi nhuan duy tri tren 70% ngay ca khi so luong nguoi dung tang gap 100 lan.',
    'Doi ngu tre nhiet huyet voi ky nang da dang bo tro: Khanh Son (Tech Lead), Quoc An (Market), Khiem (3D Specialist), Khang (UI/UX). MVP xay dung trong 6 thang voi ngan sach 85 trieu dong.',
    'Kiem chung thuc te voi so lieu cu the: NPS 68/100, 92% hoc sinh hieu bai tot hon, 76 nguoi dung thu nghiem thuc te tai THPT Nguyen Cong Tru.',
    'So huu tri tue duoc bao ve: Ma nguon da dang ky Ban quyen Tac gia. Nhan hieu PC Master LMS dang trong qua trinh dang ky bao ho doc quyen.',
]
for t in strengths:
    add_bullet(t)

add_heading('D.2. Diem Yeu (Weaknesses) — Han Che Can Khac Phuc')
weaknesses = [
    'Nhom con it kinh nghiem kinh doanh B2B voi to chuc giao duc cong lap. Quy trinh mua sam va phe duyet ngan sach thuong mat 3-6 thang.',
    'Thu vien linh kien 3D hien moi dat 45+ model. Can xay dung pipeline tu dong hoa tao model GLB theo kip toc do cap nhat linh kien thi truong.',
    'Hand Tracking con phu thuoc vao dieu kien anh sang tot. Do chinh xac giam xuong con 85% trong phong toi.',
    'Chua co App Mobile Native (iOS/Android). Ke hoach phat trien App Native vao Q1/2027.',
]
for t in weaknesses:
    add_bullet(t)

add_heading('D.3. Co Hoi (Opportunities) — Yeu To Thuan Loi')
opps = [
    'Chuong trinh GDPT 2018 bat buoc day phan cung: 2,4 trieu hoc sinh THPT bat buoc hoc phan cung may tinh tu nam 2025.',
    'Chien luoc Quoc gia phat trien nganh Vi mach Ban dan: Muc tieu dao tao 50.000 ky su vi mach den nam 2030 (Nghi quyet 36-NQ/TW).',
    'EdTech ASEAN tang truong CAGR 20,2%/nam (2023-2028): Co hoi mo rong sang Thailand, Indonesia, Philippines.',
    'Lan song dau tu vao EdTech sau COVID-19: Quy dau tu EdTech ASEAN dang tich cuc tim kiem startup B2B2C tot.',
    'Metaverse Education dang tro thanh xu huong: Apple Vision Pro, Meta Quest 3 mo ra ky nguyen VR/AR. PC Master LMS da tich hop WebXR.',
]
for t in opps:
    add_bullet(t)

add_heading('D.4. Thach Thuc (Threats) — Rui Ro Can Quan Ly')
threats = [
    'Canh tranh tu Big Tech EdTech: Google, Microsoft, Apple deu co chuong trinh giao duc lon. Tuy nhien, khong co Big Tech nao hien tai co san pham cu the ve mo phong lap rap phan cung PC cho thi truong Viet Nam.',
    'Thay doi chinh sach giao duc: Rui ro nay duoc giam thieu bang chien luoc mo rong sang trung tam dao tao tu nhan va doanh nghiep ban le linh kien.',
    'Rui ro ky thuat: MediaPipe va TensorFlow.js co the thay doi API. Nhom dang xay dung abstraction layer.',
    'Han che ve von dau tu ban dau: 85 trieu dong. Chien luoc PLG va word-of-mouth la giai phap chi phi thap.',
]
for t in threats:
    add_bullet(t)

doc.add_page_break()

# ===== PHU LUC E: RISK MATRIX =====
add_heading('PHU LUC E: MA TRAN RUI RO DU AN & KE HOACH GIAM THIEU', 14, COLOR_NAVY)

add_heading('E.1. Ma Tran Danh Gia Rui Ro')
add_para('Moi rui ro duoc danh gia tren 2 truc: Xac suat xay ra (1-5) va Muc do anh huong (1-5). Risk Score = Xac suat x Anh huong:')
add_table(
    ['Rui ro', 'Xac suat', 'Anh huong', 'Risk Score', 'Ke hoach giam thieu'],
    [
        ['MediaPipe API deprecated', '2', '4', '8', 'Abstraction layer, fallback TF.js'],
        ['Vercel outage/downtime', '1', '5', '5', 'Multi-region + Cloudflare CDN backup'],
        ['Canh tranh tu Big Tech', '2', '3', '6', 'Tap trung localization GDPT 2018'],
        ['Hoc sinh khong adopt', '2', '4', '8', 'UX toi uu, gamification, A/B test'],
        ['Truong hoc khong tra phi', '3', '5', '15', 'Pilot free 6 thang + ROI calculator'],
        ['Key person dependency', '2', '4', '8', 'Documentation day du + recruit 2 dev'],
        ['Data breach / security', '1', '5', '5', 'Supabase RLS + audit log + SIEM'],
    ]
)

add_heading('E.2. Chien Luoc Quan Ly Rui Ro Tong The')
add_para(
    'Nhom tac gia ap dung phuong phap quan ly rui ro theo tieu chuan ISO 31000:2018, voi vong lap nhan dien — danh gia — ung pho — giam sat lien tuc. '
    'Cuoc hop review rui ro duoc to chuc 2 tuan/lan, ket hop voi dashboard monitoring tu dong canh bao qua Slack khi co bat thuong ve uptime, performance hoac bao mat. '
    'Ban Giam Khao co the truy cap trang status cong khai cua he thong tai status.pc-master-lms.vercel.app (can xay dung) de theo doi tinh trang van hanh theo thoi gian thuc.'
)

doc.add_page_break()

# ===== PHU LUC F: PRODUCT ROADMAP =====
add_heading('PHU LUC F: LO TRINH PHAT TRIEN SAN PHAM CHI TIET (PRODUCT ROADMAP)', 14, COLOR_NAVY)

add_heading('F.1. Roadmap Q3/2026 - Q4/2027')
add_table(
    ['Quy', 'Milestone', 'Tinh nang chinh', 'KPI thanh cong'],
    [
        ['Q3/2026', 'AI Proctoring v2', 'Phat hien nhieu khuon mat, dashboard bao cao', 'Do chinh xac 95%+'],
        ['Q4/2026', 'Mobile PWA', 'iOS Safari + Android Chrome, offline mode', '1.000 downloads tuan dau'],
        ['Q1/2027', 'AI Curriculum', 'Gemini API quiz generator, personalized learning', '500 bai test tu dong/thang'],
        ['Q2/2027', 'API Marketplace', 'Google Classroom LTI 1.3, Microsoft Teams', '50 doi tac tich hop'],
        ['Q3/2027', 'ASEAN Expansion', 'Tieng Anh, Thai, Indonesia localization', '1.000 user nuoc ngoai'],
        ['Q4/2027', 'Seed Round', 'Pitch 10 quy EdTech ASEAN', 'Goi duoc 1 ty VND Seed'],
    ]
)

add_heading('F.2. Nguyen Tac Phat Trien San Pham (Development Principles)')
principles = [
    'User First: Moi tinh nang moi phai co it nhat 20 user test truoc khi deploy len production.',
    'Data-Driven: Quyet dinh phat trien dua tren so lieu analytics (Mixpanel) va feedback user, khong dua tren cam tinh chu quan.',
    'Incremental Delivery: Ra tinh nang theo sprint 2 tuan, deploy lien tuc (CI/CD GitHub Actions + Vercel), giam rui ro big-bang release.',
    'Backward Compatibility: API va LMS data format phai tuong thich nguoc de khong anh huong truong hoc da trien khai.',
    'Accessibility First: Tuan thu WCAG 2.1 AA de ho tro hoc sinh khuyet tat (ma nguon nhap lieu bang ban phim, phu de video).',
]
for p in principles:
    add_bullet(p)

doc.add_page_break()

# ===== PHU LUC G: HOP DONG MAU =====
add_heading('PHU LUC G: CAU TRUC HOP DONG MAU B2B VOI NHA TRUONG', 14, COLOR_NAVY)

add_heading('G.1. Dieu Khoan Chinh Trong Hop Dong License School')
contract_items = [
    ('Pham vi su dung', 'License cho phep toi da 500 tai khoan hoc sinh va giao vien thuoc truong ky ket. Khong duoc chuyen nhuong, cho thue lai hoac chia se tai khoan ra ngoai to chuc.'),
    ('Thoi han hop dong', 'Hop dong co hieu luc 12 thang ke tu ngay ky ket, tu dong gia han neu khong co thong bao huy truoc 30 ngay.'),
    ('Muc phi va thanh toan', 'Phi License: 4.800.000 VND/nam (chua VAT). Thanh toan 100% dau nam hoac chia 2 ky (50%/ky).'),
    ('SLA Uptime', 'PC Master LMS cam ket uptime 99,5% hang thang. Neu vi pham, nha truong duoc bu 10% phi hop dong cho thang bi anh huong.'),
    ('Ho tro ky thuat', 'Support qua email trong vong 24h, qua Zalo/Hotline trong gio hanh chinh. Onboarding training cho giao vien (2 buoi online).'),
    ('Bao mat du lieu', 'PC Master LMS khong ban hoac chia se du lieu hoc sinh cho ben thu 3 vi bat ky muc dich gi. Du lieu duoc xoa hoan toan trong vong 30 ngay sau khi hop dong ket thuc neu khong gia han.'),
    ('So huu noi dung', 'Noi dung bai giang do giao vien tu tao tren nen tang thuoc quyen so huu cua nha truong. PC Master LMS chi co quyen su dung de phuc vu van hanh he thong.'),
]
for title, content in contract_items:
    add_heading(f'Dieu khoan: {title}', 10.5, COLOR_BLUE)
    add_para(content)

doc.add_page_break()

# ===== PHU LUC H: MARKETING PLAN =====
add_heading('PHU LUC H: KE HOACH MARKETING DIGITAL CHI TIET', 14, COLOR_NAVY)

add_heading('H.1. Chien Luoc Content Marketing & Social Media')
add_para(
    'PC Master LMS ap dung chien luoc Marketing noi dung dua tren viec giao duc thi truong truoc, ban hang sau. '
    'Thay vi chay quang cao tra phi dat tien, nhom tap trung xay dung cong dong organic thong qua noi dung gia tri cao '
    'tren cac kenh mang xa hoi chinh tai Viet Nam:'
)

marketing_items = [
    ('YouTube Channel "PC Master Vietnam"', [
        'Upload 2 video/tuan: huong dan lap rap PC, review linh kien 3D, tips va tricks',
        'Series "Hoc Phan Cung GDPT 2018" mien phi — thu hut hoc sinh THPT',
        'Muc tieu: 10.000 subscribers trong 6 thang dau',
    ]),
    ('TikTok @PCMasterVN', [
        'Short video 15-60s demo tinh nang Hand Tracking, AI Proctoring — viral potential cao',
        'Hashtag challenge #XaydungPC3D thu hut hoc sinh thi dua',
        'Muc tieu: 50.000 followers trong nam dau',
    ]),
    ('Facebook Group "Cong dong PC Master LMS Vietnam"', [
        'Group kin cho giao vien Tin hoc trao doi bai giang, kinh nghiem',
        'Chia se tai nguyen mien phi, template bai giang, giao an mau',
        'Muc tieu: 5.000 thanh vien tich cuc trong nam dau',
    ]),
    ('Email Newsletter & Drip Campaign', [
        'Chuoi 7 email chao mung cho giao vien moi dang ky free trial',
        'Monthly newsletter tong hop xu huong EdTech + cap nhat tinh nang moi',
        'Automated email nhac nho hoc sinh hoan thanh bai tap con dang do',
    ]),
]

for channel, items in marketing_items:
    add_heading(channel, 11, COLOR_GREEN)
    for item in items:
        add_bullet(item)

add_heading('H.2. Ke Hoach PR & Quan He Bao Chi')
add_para(
    'Nhom se chu dong tiep can cac kenh bao chi giao duc va cong nghe lon tai Viet Nam de dua tin ve du an PC Master LMS, '
    'tao do nhan dien thuong hieu va uy tin trong nganh:'
)
pr_targets = [
    'VnExpress Giao duc — bai feature ve doi moi sang tao trong day hoc phan cung',
    'Tuoi Tre Online — phong su ve hoc sinh THPT xay dung nen tang EdTech',
    'ICTNews.vn & Tap chi Tia Sang — bai ky thuat ve Three.js va Hand Tracking',
    'Thanh Nien Online — profile nhom hoc sinh khoi nghiep HUIT Startup 2026',
    'Vietnam AI Summit 2027 — tham gia trinh bay ve AI in Education use case',
]
for item in pr_targets:
    add_bullet(item)

add_heading('H.3. Ke Hoach Ngan Sach Marketing Nam 1')
add_table(
    ['Kenh marketing', 'Ngan sach (trieu VND)', 'Muc tieu reach', 'Ghi chu'],
    [
        ['Content YouTube/TikTok', '20', '100.000 views/thang', 'Inhouse production'],
        ['Facebook/Instagram Ads', '30', '500.000 reach/thang', 'Target: hoc sinh, GV Tin'],
        ['Google Search Ads', '20', '10.000 click/thang', 'Keyword: hoc PC, phan cung'],
        ['Workshop tai truong THPT', '15', '20 truong, 1.000 hoc sinh', 'Tra phi thue hall'],
        ['PR & Press Release', '5', '10 bai bao dang', 'Quan he bao chi'],
        ['Tong ngan sach nam 1', '90', '---', '20% tong doanh thu'],
    ]
)

doc.add_page_break()

# Final page
add_heading('KET LUAN & LOI CAM ON', 14, COLOR_NAVY)
add_para(
    'PC Master LMS la san pham cua su ket hop giua am hieu sau sac ve nhu cau giao duc thuc te va kha nang '
    'ung dung cong nghe hien dai mot cach sang tao va co chieu sau. Nhom tac gia tin tuong rang du an nay '
    'khong chi la mot san pham khoi nghiep co tiem nang thuong mai hoa cao, ma con la mot dong gop thiet thuc '
    'cho su phat trien cua nen giao duc STEM Viet Nam trong ky nguyen chuyen doi so.'
)
add_para(
    'Voi nen tang cong nghe vung chac da duoc kiem chung thuc te, mo hinh kinh doanh ro rang va doi ngu hung hau, '
    'chung toi kính moi Ban Giam Khao cung toan the quy vi trai nghiem truc tiep san pham tai '
    'https://pc-master-lms.vercel.app/ va danh gia day du nhung no luc ma nhom da dat ra.'
)
add_para(
    'Loi cam on chan thanh nhat tu toan the nhom tac gia PC Master LMS toi Ban Giam Hieu va Ban To Chuc '
    'Cuoc thi Khoi nghiep HUIT Startup 2026, Co Doan Thuy Kim Phuong va Thay Tran Minh Phung la GVHD tan tam, '
    'Ban Giam hieu Truong THPT Nguyen Cong Tru da tao dieu kien thu nghiem thuc te, 76 hoc sinh THPT Nguyen Cong Tru '
    'da tham gia trai nghiem va dong gop y kien quy bau, va Vuon uom Khoi nghiep HUIT — kinh mong tiep tuc dong hanh cung chung toi!'
)

# Save
filename_out = 'Cuon_Thuyet_Minh_Chinh_Thuc_PC_Master_LMS_HUIT_2026_50_Trang.docx'
doc.save(filename_out)

import docx as dx
d2 = dx.Document(filename_out)
words_para = sum(len(p.text.split()) for p in d2.paragraphs)
words_table = sum(
    len(para.text.split())
    for tbl in d2.tables
    for row in tbl.rows
    for cell in row.cells
    for para in cell.paragraphs
)
total_words = words_para + words_table
breaks = len([p for p in d2.paragraphs if 'w:br' in p._p.xml])
tables = len(d2.tables)
print(f'[OK] Saved: {filename_out}')
print(f'[STATS] Total words: {total_words} | Page breaks: {breaks} | Tables: {tables}')
print(f'[EST] Estimated pages: ~{total_words//250 + breaks}')
