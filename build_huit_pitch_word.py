import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def create_winning_pitch_docx():
    doc = docx.Document()

    # Page Margins
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    COLOR_NAVY = RGBColor(3, 31, 59)       # #031F3B
    COLOR_PRIMARY = RGBColor(8, 158, 96)   # #089E60
    COLOR_DARK = RGBColor(30, 41, 59)      # #1E293B
    COLOR_BLUE = RGBColor(14, 116, 144)    # #0E7490
    COLOR_ACCENT = RGBColor(217, 119, 6)   # #D97706

    def add_title(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.font.name = 'Arial'
        r.font.size = Pt(18)
        r.font.bold = True
        r.font.color.rgb = COLOR_NAVY
        return p

    def add_subtitle(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.font.name = 'Arial'
        r.font.size = Pt(11)
        r.font.italic = True
        r.font.color.rgb = COLOR_PRIMARY
        return p

    def add_h1(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run(text)
        r.font.name = 'Arial'
        r.font.size = Pt(14)
        r.font.bold = True
        r.font.color.rgb = COLOR_NAVY
        return p

    def add_h2(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(text)
        r.font.name = 'Arial'
        r.font.size = Pt(12)
        r.font.bold = True
        r.font.color.rgb = COLOR_BLUE
        return p

    def add_p(text, bold_prefix="", italic=False):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        if bold_prefix:
            r_b = p.add_run(bold_prefix)
            r_b.font.name = 'Arial'
            r_b.font.size = Pt(10)
            r_b.font.bold = True
            r_b.font.color.rgb = COLOR_DARK
        r = p.add_run(text)
        r.font.name = 'Arial'
        r.font.size = Pt(10)
        r.font.italic = italic
        r.font.color.rgb = COLOR_DARK
        return p

    def add_callout(text, title="📌 LƯU Ý QUAN TRỌNG:"):
        tbl = doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = tbl.cell(0, 0)
        cell.width = Inches(6.8)
        shd = parse_xml(r'<w:shd {} w:fill="F8FAFC"/>'.format(nsdecls('w')))
        cell._tc.get_or_add_tcPr().append(shd)

        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        
        r_t = p.add_run(f"{title}\n")
        r_t.font.name = 'Arial'
        r_t.font.size = Pt(10)
        r_t.font.bold = True
        r_t.font.color.rgb = COLOR_PRIMARY

        r_v = p.add_run(text)
        r_v.font.name = 'Arial'
        r_v.font.size = Pt(9.5)
        r_v.font.color.rgb = COLOR_DARK
        doc.add_paragraph()

    # Document Header
    add_title("BỘ TÀI LIỆU PITCHING 5 PHÚT CHUẨN HUIT 2026\n& BẢN HƯỚNG DẪN ĐẠT ĐIỂM TỐI ĐA 100/100 ĐIỂM")
    add_subtitle("DỰ ÁN: PC MASTER BUILDER (PC MASTER LMS)\nĐơn vị: Trường THPT Nguyễn Công Trứ & HUIT | Tác giả: Nguyễn Phúc Khánh Sơn, Đặng Quốc An, Nguyễn Phạm Gia Khiêm, Ngô Minh Khang\nGVHD: Cô Đoàn Thụy Kim Phượng & Thầy Trần Minh Phụng")

    doc.add_paragraph()

    # Introduction Box
    add_callout(
        "Tài liệu này được biên soạn bám sát 100% BẢNG TIÊU CHÍ CHẤM ĐIỂM CHÍNH THỨC PHỤ LỤC 4 (BÁN KẾT & CHUNG KẾT) của Cuộc thi Khởi nghiệp HUIT 2026.\n"
        "Bao gồm: Phân tích tối ưu 100 điểm, Kịch bản Pitching 5 phút chính xác 300 giây, Kịch bản phản biện 10 SDGs, Hướng dẫn quay Video giới thiệu 5đ và Thiết kế gian hàng trưng bày 20đ.",
        "🎯 MỤC TIÊU: ĐẠT ĐIỂM TỐI ĐA (100/100 ĐIỂM) TẠI VÒNG BÁN KẾT & CHUNG KẾT HUIT 2026"
    )

    # SECTION 1
    add_h1("PHẦN 1: BẢN PHÂN TÍCH ĐỐI CHIẾU TIÊU CHÍ CHẤM ĐIỂM 100 ĐIỂM (PHỤ LỤC 4 HUIT)")
    add_p("Bảng tiêu chí chính thức của HUIT chia thành 3 Hạng mục chính (Tổng 100 điểm). Dự án PC Master LMS đáp ứng từng tiêu chí như sau:")

    # Table 1: Rubric Mapping
    table_rubric = doc.add_table(rows=1, cols=4)
    table_rubric.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table_rubric.rows[0].cells
    headers = ["Hạng mục tiêu chí", "Điểm tối đa", "Nội dung kiểm tra của BGK", "Chiến lược PC Master lấy điểm tuyệt đối"]
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        shd = parse_xml(r'<w:shd {} w:fill="031F3B"/>'.format(nsdecls('w')))
        hdr_cells[i]._tc.get_or_add_tcPr().append(shd)
        p = hdr_cells[i].paragraphs[0]
        p.runs[0].font.name = 'Arial'
        p.runs[0].font.bold = True
        p.runs[0].font.color.rgb = RGBColor(255, 255, 255)
        p.runs[0].font.size = Pt(9.5)

    rubric_data = [
        ("I.1 Tính thẩm mỹ/sáng tạo gian hàng", "5 điểm", "Gian hàng đẹp, độc đáo, thu hút thị giác.", "Setup góc 3D Cyberpunk, màn hình lớn demo Hand-Tracking 3D và QR code trải nghiệm."),
        ("I.2 Tính quy mô/đầu tư gian hàng", "5 điểm", "Sự chuẩn bị chỉn chu, bài bản của đội thi.", "Trang bị Tablet/Laptop chạy mượt Web pc-master-lms.vercel.app, banner standee in ấn sắc nét."),
        ("I.3 Nổi bật sản phẩm trưng bày", "10 điểm", "Khách tham quan & BGK được trải nghiệm trực tiếp.", "Cho BGK tự dùng cử chỉ tay (Hand-tracking) bốc RAM/CPU 3D ngay tại gian hàng."),
        ("II.1 Kỹ năng thuyết trình gian hàng", "10 điểm", "Diễn đạt lưu loát, tự tin, phân công rõ ràng.", "Phân công Khánh Sơn (Trưởng nhóm - Tech) & Quốc An (Market/Sales) trình bày nhịp nhàng."),
        ("II.2 Kỹ năng phản biện ấn tượng", "10 điểm", "Trả lời sắc bén, tự tin, có số liệu thực tế.", "Dùng số liệu thực nghiệm 76 học sinh 10A1/10C9, +28% điểm số, 20k/tháng, 10 SDGs."),
        ("III.1 Tính sáng tạo độc đáo", "15 điểm", "Sản phẩm mới, công nghệ mới, khác biệt rõ.", "LMS 3D đầu tiên tại VN kết hợp WebGL 3D, MediaPipe Hand-Tracking & AI Chống chặt chém."),
        ("III.2 Năng lực tổ chức thực hiện", "15 điểm", "Kế hoạch SX (2đ), MKT (5đ), Tài chính (5đ), Nhân sự (3đ).", "Lộ trình Agile/Scrum, Mô hình Freemium PLG, Break-even 250 users, 4 founder tự chủ tech 100%."),
        ("III.3 Hiệu quả KT & Tác động xã hội", "15 điểm", "Giải quyết nỗi đau, đáp ứng 17 SDGs UN, khả thi.", "Tiết kiệm 600M/phòng lab, phụng sự 10 UN SDGs (SDG 4, 9, 12, 8, 10...), triệt tiêu E-waste."),
        ("III.4 Thị trường tiềm năng", "5 điểm", "Nhu cầu lớn, phân khúc rõ, đối thủ & số liệu kiểm chứng.", "Thị trường 3.000+ trường THPT (1M học sinh/năm), thử nghiệm 76 học sinh 10A1/10C9 (+28% điểm)."),
        ("III.5 Ứng dụng công nghệ", "5 điểm", "Công nghệ hiện đại, đột phá so với đối thủ.", "Next.js 16, Supabase RLS, MediaPipe Vision Edge AI, WebGL Three.js (đối thủ chưa có)."),
        ("III.6 Video clip giới thiệu dự án", "5 điểm", "Đủ thông tin nhóm, quá trình, giá trị cốt lõi, âm thanh đẹp.", "Video 3 phút quay thực tế học sinh trải nghiệm, đồ họa 3D, âm thanh lồng tiếng chuyên nghiệp.")
    ]

    for cat, pts, check, strat in rubric_data:
        row_cells = table_rubric.add_row().cells
        row_cells[0].text = cat
        row_cells[1].text = pts
        row_cells[2].text = check
        row_cells[3].text = strat
        for i in range(4):
            p = row_cells[i].paragraphs[0]
            p.runs[0].font.name = 'Arial'
            p.runs[0].font.size = Pt(9)
            p.runs[0].font.color.rgb = COLOR_DARK

    doc.add_paragraph()

    # SECTION 2
    add_h1("PHẦN 2: KỊCH BẢN THUYẾT TRÌNH PITCHING 5 PHÚT (CHÍNH XÁC 300 GIÂY)")
    add_callout(
        "Kịch bản được thiết kế khớp từng giây với lời thoại của 2 thí sinh Khánh Sơn và Quốc An.\n"
        "Đảm bảo phân bổ đầy đủ 6 Tiêu chí thuyết minh: Sáng tạo (15đ) - Năng lực thực hiện (15đ) - SDGs & Xã hội (15đ) - Thị trường (5đ) - Công nghệ (5đ) - Video (5đ).",
        "⏱️ QUY ĐỊNH THỜI GIAN: ĐÚNG 5 PHÚT (300 GIÂY) — KHÔNG THỪA KHÔNG THIẾU 1 GIÂY"
    )

    pitch_slides = [
        ("SLIDE 1: BÌA DỰ ÁN & GIỚI THIỆU THÀNH VIÊN", "00:00 - 00:30 (30 giây)",
         "• Tên dự án: PC MASTER BUILDER (PC Master LMS)\n• Slogan: Nền tảng Giả lập Lắp ráp PC 3D & Hệ thống LMS Thông minh\n• Đơn vị: THPT Nguyễn Công Trứ & HUIT | GVHD: Cô Đoàn Thụy Kim Phượng & Thầy Trần Minh Phụng\n• Thuyết trình: Nguyễn Phúc Khánh Sơn (Leader - Tech) & Đặng Quốc An (Market/Sales)",
         "Sơn: 'Kính chào Ban Giám Khảo và toàn thể hội thi! Em là Khánh Sơn và đây là Quốc An, đại diện cho nhóm tác giả dự án PC Master Builder - Nền tảng học tập & mô phỏng phần cứng máy tính 3D tích hợp AI đầu tiên cho học sinh phổ thông tại Việt Nam!'",
         "Tạo ấn tượng chuyên nghiệp, tự tin, quét mã QR Web ngay trên Slide bàng thiết bị di động."),

        ("SLIDE 2: THỰC TRẠNG & NỖI ĐẠO THỊ TRƯỜNG (PROBLEM)", "00:30 - 01:10 (40 giây)",
         "• Nỗi đau 1: 85% trường THPT thiếu phòng lab thực hành phần cứng (GDPT 2018 bắt buộc Tin học 10).\n• Nỗi đau 2: Chi phí phòng lab thật quá đắt (600 - 800 triệu VNĐ/phòng 40 máy), rủi ro gãy socket CPU/hỏng RAM rất cao.\n• Nỗi đau 3: Học sinh bị 'học chay' thụ động; phụ huynh khi mua PC bị cửa hàng 'chặt chém' chênh giá.",
         "Sơn: 'Thưa quý vị, Chương trình GDPT 2018 bắt buộc 100% học sinh Lớp 10 phải học phần cứng PC. Tuy nhiên, hơn 85% trường THPT hiện nay không có phòng lab phần cứng do chi phí trang bị lên tới 800 triệu đồng. Rủi ro hỏng linh kiện và rác thải điện tử cực kỳ lớn khiến học sinh vẫn phải 'học chay'. Ngoài ra, khi tự đi mua máy tính, học sinh và phụ huynh thường chịu sự bất đồng thông tin và bị 'chặt chém' giá rất nặng!'",
         "Đánh trúng tiêu chí Thị trường (5đ) & Nỗi đau xã hội thực tế (15đ)."),

        ("SLIDE 3: GIẢI PHÁP & ĐỘT PHÁ CÔNG NGHỆ (SOLUTION & TECH)", "01:10 - 02:00 (50 giây)",
         "• Giải pháp: Phòng Lab phần cứng ảo 3D 100% trên Web, 0 đồng thiết bị cứng.\n• Đột phá 1: Thao tác 3D tương tác 360 độ & Điều khiển Cử chỉ tay (Hand-tracking qua Webcam bằng MediaPipe).\n• Đột phá 2: AI Tutor Chống chặt chém & Tra giá real-time từ API các đại lý lớn (Phong Vũ, GearVN).\n• Đột phá 3: Next.js 16 + Supabase RLS bảo mật phân quyền LMS 3 Role (Học sinh, Giáo viên, Phụ huynh).",
         "Sơn: 'PC Master Builder giải quyết triệt để vấn đề này với Phòng Lab Ảo 3D ngay trên Web! Điểm độc đáo nhất là công nghệ Hand-tracking qua Webcam bằng MediaPipe: Học sinh có thể co nắm ngón tay để bốc, xoay, cắm RAM, CPU 3D như thật mà KHÔNG CẦN KÍNH VR ĐẮT TIỀN. Kết hợp cùng AI Chống Chặt Chém tự động check tương thích Socket, tính TDP nguồn và cảnh báo chênh giá thị trường real-time!'",
         "Lấy trọn 15đ Tính Sáng Tạo + 5đ Ứng Dụng Công Nghệ WebGL/MediaPipe."),

        ("SLIDE 4: THỰC NGHIỆM VALIDATION & ĐÓNG GÓP 10 SDGS UN", "02:00 - 02:50 (50 giây)",
         "• Thực nghiệm thực tế: 76 học sinh Lớp 10A1 & 10C9 (THPT Nguyễn Công Trứ) trải nghiệm.\n• Số liệu kiểm chứng: Điểm kiểm tra phần cứng tăng +28%, Tỷ lệ hoàn thành bài học 94%, 92% hài lòng.\n• Đóng góp 10 SDGs UN: Nổi bật với SDG 4 (Giáo dục chất lượng 3D), SDG 9 (Công nghệ đổi mới AI), SDG 12 (Giảm 100% E-waste rác thải điện tử), SDG 8 (Định hướng Vi mạch - Bán dẫn), SDG 10 (Bình đẳng số).",
         "An: 'Sản phẩm của chúng em đã chạy thực tế tại pc-master-lms.vercel.app và được kiểm chứng trên 76 học sinh THPT. Kết quả: Điểm thi thực hành tăng 28%, tỷ lệ hoàn thành bài học đạt 94%! Đặc biệt, PC Master tự hào phụng sự 10 Mục tiêu Phát triển Bền vững (UN SDGs): Nổi bật là SDG 4 giáo dục chất lượng, SDG 9 công nghệ AI sáng tạo, SDG 12 triệt tiêu 100% rác thải điện tử E-waste, SDG 8 định hướng nhân lực Vi mạch - Bán dẫn và SDG 10 bình đẳng công nghệ số vùng xa!'",
         "Lấy trọn 15đ Hiệu Quả Kinh Tế & Tác Động Xã Hội (10 SDGs)."),

        ("SLIDE 5: MÔ HÌNH KINH DOANH, TÀI CHÍNH & KHẢ NĂNG NHÂN RỘNG", "02:50 - 03:40 (50 giây)",
         "• 3 Nguồn thu bền vững: (1) B2C Premium 20k/tháng (bằng 1 ly trà sữa); (2) B2B LMS Trường học (15-25 triệu/năm); (3) Affiliate Marketing bán linh kiện (hoa hồng 2-5%).\n• Tài chính & Vận hành: Chi phí Cloud Serverless lean chỉ ~2M/tháng, Break-even tại 250 Pro users (dưới 4 tháng), Biên lợi nhuận ròng SaaS >75%.\n• Phân bổ 50M gọi vốn/tiền thưởng: 40% Hạ tầng Cloud/AI API, 40% Marketing B2B 10 trường THPT, 20% Đăng ký Bản quyền IP.",
         "An: 'Về Mô hình Kinh doanh, PC Master tạo ra 3 nguồn thu bền vững: Gói B2C Pro chỉ 20.000đ/tháng - bằng giá 1 ly trà sữa; Gói B2B LMS cho nhà trường tiết kiệm hàng trăm triệu tiền phòng lab; và Hoa hồng Affiliate 2-5% từ đại lý bán lẻ. Nhờ hạ tầng Serverless tinh gọn, dự án đạt điểm huề vốn chỉ với 250 học sinh trả phí và biên lợi nhuận ròng trên 75%. Nếu giành 50 triệu tiền thưởng, nhóm sẽ dành 40% nâng cấp Cloud API, 40% phủ sóng B2B 10 trường THPT và 20% đăng ký Bản quyền Tác giả!'",
         "Lấy trọn 15đ Năng Lực Thực Hiện (Kế hoạch MKT 5đ, Tài chính 5đ, Nhân sự 3đ, SX 2đ)."),

        ("SLIDE 6: ĐỘI NGŨ FOUNDER, BẢN QUYỀN & THÔNG ĐỆP KẾT MẠNH MẼ", "03:40 - 04:30 (50 giây)",
         "• Đội ngũ tự chủ 100%: Khánh Sơn (Fullstack/Cloud), Quốc An (AI Vision), Gia Khiêm (3D WebGL), Minh Khang (UI/UX & Content SGK). Zero Outsource spend.\n• Bảo chứng học thuật: GVHD Cô Đoàn Thụy Kim Phượng & Thầy Trần Minh Phụng.\n• Pháp lý & IP: Đăng ký Bản quyền Tác giả mã nguồn phần mềm & Nhãn hiệu độc quyền.",
         "Sơn: 'Đội ngũ founder HUIT của chúng em tự chủ kỹ thuật 100% từ 3D WebGL, AI Computer Vision đến Fullstack Cloud mà không tốn 1 đồng thuê ngoài. Dự án được bảo chứng học thuật bởi Cô Kim Phượng và Thầy Minh Phụng, đồng thời đã hoàn tất đăng ký Bản quyền Tác giả phần mềm. PC Master LMS xứng đáng giành chiến thắng vì đây là sản phẩm Make-in-HUIT tự chủ công nghệ, giải đúng nỗi đau giáo dục bằng công nghệ 3D/AI và mang lại giá trị nhân văn to lớn cho cộng đồng!'",
         "Ấn tượng đanh thép, khẳng định năng lực Founder và tính tự chủ Make-in-HUIT.")
    ]

    for title, timing, content, speech, note in pitch_slides:
        add_h2(f"🎬 {title} [{timing}]")
        add_p(content, bold_prefix="📌 Nội dung hiển thị Slide & Ý chính: ")
        add_p(speech, bold_prefix="🗣️ Lời thoại Pitching thực tế: ", italic=True)
        add_p(note, bold_prefix="💡 Bí quyết ghi điểm BGK: ")
        doc.add_paragraph()

    # SECTION 3
    add_h1("PHẦN 3: KỊCH BẢN PHẢN BIỆN TẠI GIAN HÀNG & HỘI ĐỒNG (BỘ CÂU HỎI TRỌNG YẾU)")
    add_p("BGK Bán kết & Chung kết HUIT sẽ chấm 10 điểm Kỹ năng phản biện ấn tượng (Tiêu chí II.2). Dưới đây là 5 câu hỏi trọng yếu chắc chắn bị hỏi và câu trả lời mẫu:")

    qa_list = [
        ("BGK hỏi: 'Tại sao nhóm định giá gói Pro chỉ 20.000 VNĐ/tháng? Liệu có đủ bù đắp chi phí vận hành?'",
         "Kính thưa Ban Giám khảo, 20.000đ/tháng là chiến lược Định giá tâm lý (Psychological Pricing) - bằng đúng 1 ly trà sữa bình dân. Mức giá này giúp tối đa hóa tỷ lệ chuyển đổi (Conversion Rate). Nhờ hạ tầng Serverless Pay-as-you-go trên Vercel và Supabase, chi phí server cho 1.000 active users chỉ tốn ~1.5 - 2 triệu VNĐ/tháng, nên chỉ cần 250 user Pro là hệ thống đã huề vốn và đạt biên lợi nhuận >75%."),

        ("BGK hỏi: 'Dự án đóng góp cho 10 Mục tiêu SDGs của Liên Hợp Quốc có bị ôm đồm quá không?'",
         "Kính thưa Ban Giám khảo, 10 SDGs này là kết quả tự nhiên của mô hình ảo hóa 3D trên Cloud: Khi ảo hóa phòng lab 3D -> Đạt SDG 4 (Giáo dục 3D), SDG 9 (Hạ tầng AI/Web), SDG 10 (Bình đẳng vùng xa). Khi không dùng linh kiện thật thử nghiệm sai -> Đạt SDG 3 (An toàn), SDG 12 (Giảm 100% E-waste), SDG 13 (Khí hậu). Khi định hướng nghề nghiệp và liên kết đại lý -> Đạt SDG 5, 8, 11, 17. Đây là tác động cộng hưởng bền vững của sản phẩm EdTech số hóa!"),

        ("BGK hỏi: 'Nếu các trường THPT máy tính cấu hình quá yếu thì ứng dụng 3D có chạy nổi không?'",
         "Kính thưa Ban Giám khảo, PC Master sử dụng thuật toán nén Google Draco 3D giúp giảm 85% dung lượng model (từ 50MB xuống 3MB), kết hợp thuật toán Dynamic LOD tự động điều chỉnh số lượng Polygon theo cấu hình máy. Hệ thống chạy mượt 60 FPS ngay cả trên máy văn phòng Core i3 cũ RAM 4GB hoặc trên điện thoại smartphone thông qua trình duyệt Web!"),

        ("BGK hỏi: 'Làm sao nhóm ngăn chặn đối thủ lớn copy mô hình sản phẩm này?'",
         "Kính thưa Ban Giám khảo, nhóm tạo ra 3 lớp rào cản vững chắc (Moat): (1) Kho tài sản mô hình 3D tối ưu độc quyền; (2) Thuật toán AI Proctoring & Tra giá đã đăng ký Bản quyền Tác giả; (3) Tốc độ thực thi tinh gọn (Agile Execution) của 4 founder tự chủ 100% kỹ thuật, có thể tung tính năng mới trong 48h."),

        ("BGK hỏi: 'Trong 1 câu duy nhất, tại sao PC Master xứng đáng giành vị trí Quán quân HUIT 2026?'",
         "Kính thưa Ban Giám khảo, PC Master xứng đáng giành vị trí Quán quân vì đây là sản phẩm Make-in-HUIT 100% tự chủ công nghệ, giải quyết triệt để nỗi đau giáo dục phổ thông bằng công nghệ 3D/AI đột phá, sở hữu mô hình kinh doanh tài chính siêu bền vững và mang lại giá trị nhân văn to lớn cho 10 mục tiêu phát triển bền vững của Liên Hợp Quốc!")
    ]

    for q, a in qa_list:
        add_p(q, bold_prefix="❓ ", italic=False)
        add_p(a, bold_prefix="🎯 Trả lời mẫu: ", italic=True)
        doc.add_paragraph()

    # SECTION 4
    add_h1("PHẦN 4: HƯỚNG DẪN DỰNG VIDEO CLIP GIỚI THIỆU DỰ ÁN 3 PHÚT (ĐẠT 5/5 ĐIỂM TIÊU CHÍ III.6)")
    add_p("Tiêu chí III.6 dành 5 điểm cho Video clip giới thiệu dự án. Đội thi cần dựng Video theo cấu trúc 3 phút như sau:")
    add_p("• 00:00 - 00:30 (30s): Giới thiệu nhóm tác giả (Sơn, An, Khiêm, Khang), đơn vị THPT Nguyễn Công Trứ & HUIT, logo PC Master.")
    add_p("• 00:30 - 01:15 (45s): Quá trình hình thành dự án & Nỗi đau phòng lab phần cứng 800 triệu.")
    add_p("• 01:15 - 02:15 (60s): Mô tả sản phẩm thực tế: Quay màn hình tính năng Hand-Tracking 3D bốc RAM/CPU, AI Tra giá & LMS chấm điểm.")
    add_p("• 02:15 - 02:45 (30s): Giá trị cốt lõi & Đóng góp 10 UN SDGs (Giáo dục chất lượng, giảm rác thải E-waste).")
    add_p("• 02:45 - 03:00 (15s): Kết video: Lời kêu gọi hợp tác, nhạc nền công nghệ sôi động, hiển thị QR Code và link pc-master-lms.vercel.app.")

    doc.add_paragraph()

    # SECTION 5
    add_h1("PHẦN 5: CHUẨN BỊ GIAN HÀNG & TRƯNG BÀY DEMO (ĐẠT 20/20 ĐIỂM TIÊU CHÍ I & II)")
    add_p("Tiêu chí I & II dành tới 40 điểm cho Gian hàng và Thuyết trình tại gian hàng. Đội thi chuẩn bị như sau:")
    add_p("1. Setup Gian hàng (20 điểm):", bold_prefix="• ")
    add_p("   - Đặt 1 Laptop/PC màn hình lớn chạy sẵn trang Builder 3D `pc-master-lms.vercel.app`.")
    add_p("   - In Standee khổ A0 màu Cyberpunk Dark Mode hiển thị rõ Slogan, 10 SDGs và QR Code.")
    add_p("   - Chuẩn bị 1 Webcam nét để khách tham quan & BGK tự trải nghiệm bốc linh kiện 3D bằng Hand-tracking.")
    add_p("2. Thuyết trình & Trực gian hàng (20 điểm):", bold_prefix="• ")
    add_p("   - Phân công Khánh Sơn & Quốc An túc trực 100% thời gian tại gian hàng.")
    add_p("   - Mặc đồng phục chỉn chu, đeo thẻ dự thi HUIT, chủ động chào đón BGK với nụ cười tự tin.")

    filename = "Kich_Ban_Pitching_5_Phut_Dat_Diem_Toi_Da_Startup_HUIT_2026.docx"
    doc.save(filename)
    print(f"Winning pitch document created successfully: {filename}")

if __name__ == '__main__':
    create_winning_pitch_docx()
