import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def create_element(name):
    return OxmlElement(name)

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_table_borders(table, color="D3D3D3", sz="4", val="single"):
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>\n'
        f'  <w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>\n'
        f'  <w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>\n'
        f'  <w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>\n'
        f'  <w:insideV w:val="none"/>\n'
        f'  <w:left w:val="none"/>\n'
        f'  <w:right w:val="none"/>\n'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

def make_callout(doc, text, title=""):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    
    cell = tbl.cell(0, 0)
    cell.width = Inches(6.5)
    set_cell_background(cell, "F0F4F8")
    
    # Left border blue accent
    tcPr = cell._tc.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>\n'
        f'  <w:top w:val="none"/>\n'
        f'  <w:left w:val="single" w:sz="24" w:space="0" w:color="1E40AF"/>\n'
        f'  <w:bottom w:val="none"/>\n'
        f'  <w:right w:val="none"/>\n'
        f'</w:tcBorders>'
    )
    tcPr.append(borders)
    set_cell_margins(cell, top=140, bottom=140, left=200, right=200)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    if title:
        run_title = p.add_run(f"📌 {title}\n")
        run_title.font.name = "Arial"
        run_title.font.size = Pt(11)
        run_title.font.bold = True
        run_title.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)
        
    run_text = p.add_run(text)
    run_text.font.name = "Arial"
    run_text.font.size = Pt(10)
    run_text.font.color.rgb = RGBColor(0x33, 0x41, 0x55)
    
    # spacing after callout table
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

def build_document():
    doc = Document()
    
    # Page setup - Margins 1 inch
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Styles config
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Arial'
    normal_style.font.size = Pt(10.5)
    normal_style.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
    
    # Header Title Section
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(4)
    r_title = title_p.add_run("TÀI LIỆU GIỚI THIỆU DỰ ÁN & HƯỚNG DẪN THÀNH VIÊN MỚI")
    r_title.font.name = "Arial"
    r_title.font.size = Pt(18)
    r_title.font.bold = True
    r_title.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A) # Primary Dark Blue
    
    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_p.paragraph_format.space_after = Pt(18)
    r_sub = sub_p.add_run("DỰ ÁN: PC MASTER LMS (PC MASTER BUILDER)\nNền Tảng Giáo Dục & Mô Phỏng Lắp Ráp Máy Tính Tương Tác Tích Hợp AI & Hand Tracking")
    r_sub.font.name = "Arial"
    r_sub.font.size = Pt(12)
    r_sub.font.italic = True
    r_sub.font.color.rgb = RGBColor(0x47, 0x55, 0x69)

    make_callout(
        doc,
        "Tài liệu này được biên soạn dành cho thành viên mới gia nhập nhóm dự án PC Master LMS. "
        "Tài liệu bao gồm toàn bộ thông tin về mục tiêu dự án, đối tượng người dùng, tính năng hệ thống, "
        "công nghệ sử dụng, nguồn dữ liệu/tài nguyên và quy trình phối hợp làm việc nhóm.",
        "THÔNG IN ONBOARDING DÀNH CHO THÀNH VIÊN MỚI"
    )

    # Helper function for Headings
    def add_h1(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.keep_with_next = True
        r = p.add_run(text)
        r.font.name = "Arial"
        r.font.size = Pt(14)
        r.font.bold = True
        r.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF) # Blue Accent
        return p

    def add_h2(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
        r = p.add_run(text)
        r.font.name = "Arial"
        r.font.size = Pt(12)
        r.font.bold = True
        r.font.color.rgb = RGBColor(0x0D, 0x94, 0x88) # Teal Accent
        return p

    def add_bullet(text, bold_prefix="", level=0):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(4)
        if bold_prefix:
            r_b = p.add_run(bold_prefix)
            r_b.font.bold = True
            r_b.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
        r_t = p.add_run(text)
        r_t.font.color.rgb = RGBColor(0x33, 0x41, 0x55)
        return p

    # -------------------------------------------------------------
    # PHẦN 1: TỔNG QUAN DỰ ÁN (WEB LÀ GÌ?)
    # -------------------------------------------------------------
    add_h1("PHẦN 1: TỔNG QUAN DỰ ÁN (WEB LÀ GÌ?)")
    
    add_h2("1.1. Giới thiệu chung")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.add_run(
        "PC Master LMS (PC Master Builder) là một nền tảng giáo dục trực tuyến đột phá (EdTech), "
        "kết hợp giữa hệ thống quản lý học tập (LMS) và trình mô phỏng thực hành lắp ráp máy tính 2D/3D tương tác real-time. "
        "Web giúp học sinh học kiến thức phần cứng máy tính một cách trực quan, sinh động mà không cần chuẩn bị linh kiện thật đắt đỏ "
        "hoặc đối mặt với rủi ro hỏng hóc thiết bị."
    )
    
    add_bullet(" PC Master Builder / PC Master LMS", "Tên chính thức: ")
    add_bullet(" https://pc-master-lms.vercel.app", "URL triển khai (Production): ")
    add_bullet(" Next.js 16 (App Router), React 19, Supabase, Tailwind CSS, Google Gemini AI, MediaPipe Hand Tracking, Three.js.", "Nền tảng cốt lõi: ")

    add_h2("1.2. Bối cảnh ra đời & Sứ mệnh")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.add_run(
        "Hiện nay, việc giảng dạy phần cứng máy tính trong chương trình Tin học THPT (đặc biệt là bộ sách Kết Nối Tri Thức và Cánh Diều) "
        "thường gặp nhiều khó khăn: nhà trường thiếu kinh phí trang bị phòng lab thực hành phần cứng, linh kiện thực tế dễ bị hỏng hóc "
        "khi học sinh thao tác sai, và lý thuyết trên sách vở thường khô khan, khó hình dung. "
        "PC Master ra đời để giải quyết triệt để bài toán này bằng công nghệ mô phỏng web 100% không cần cài đặt."
    )

    add_h2("1.3. Mục tiêu chiến lược")
    add_bullet("Thay thế phương pháp đọc-chép khô khan bằng mô phỏng kéo-thả và tương tác 3D trực quan.", "1. Trực quan hóa giáo dục: ")
    add_bullet("Cung cấp các công cụ kiểm tra tương thích socket, tính toán công suất TDP, cảnh báo nghẽn cổ chai.", "2. Chuẩn hóa kiến thức phần cứng: ")
    add_bullet("Trợ lý AI Gemini đóng vai trò trợ giảng 24/7, tự động sinh đề kiểm tra và gợi ý cấu hình phù hợp.", "3. Ứng dụng Trí tuệ Nhân tạo (AI): ")
    add_bullet("Cho phép học sinh điều khiển gắp/xoay linh kiện máy tính thông qua webcam mà không cần chuột/bàn phím.", "4. Đột phá công nghệ Hand Tracking: ")
    add_bullet("Liên kết 4 đối tượng: Admin, Giáo viên, Học sinh và Phụ huynh trên cùng một nền tảng.", "5. Số hóa toàn diện quản lý học tập: ")

    # -------------------------------------------------------------
    # PHẦN 2: ĐỐI TƯỢNG HƯỚNG ĐẾN (TARGET AUDIENCE)
    # -------------------------------------------------------------
    add_h1("PHẦN 2: ĐỐI TƯỢNG HƯỚNG ĐẾN (TARGET AUDIENCE)")
    
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.add_run("Hệ thống được thiết kế với phân quyền chặt chẽ (RBAC) phục vụ 4 nhóm người dùng chính:")

    # Table for roles
    table_role = doc.add_table(rows=5, cols=3)
    table_role.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table_role)
    
    headers = ["Vai Trò (Role)", "Đối Tượng Mẫu", "Chức Năng & Giá Trị Mang Lại"]
    widths = [Inches(1.5), Inches(1.8), Inches(3.2)]
    
    for i, h in enumerate(headers):
        cell = table_role.cell(0, i)
        cell.width = widths[i]
        set_cell_background(cell, "1E3A8A")
        set_cell_margins(cell, top=120, bottom=120, left=120, right=120)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(h)
        r.font.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r.font.size = Pt(10)

    roles_data = [
        ("Học Sinh\n(Student)", "Học sinh THPT (Lớp 10, 11, 12), người yêu thích công nghệ", "Học bài giảng tương tác, thực hành lắp ráp PC 2D/3D, điều khiển cử chỉ tay, làm bài kiểm tra, thi đấu 2 người, tích điểm XP, nhận chứng chỉ số."),
        ("Giáo Viên\n(Teacher)", "Giáo viên môn Tin học các trường THPT, giảng viên", "Tạo và quản lý lớp học (mã join class), soạn bài giảng Markdown/Video, tạo bộ câu hỏi kiểm tra, theo dõi tiến độ học sinh, cấp chứng chỉ."),
        ("Phụ Huynh\n(Parent)", "Cha mẹ học sinh muốn theo dõi việc học của con", "Theo dõi thời gian thực con đang học, xem báo cáo số bài đã hoàn thành, thời gian tự học, điểm kiểm tra trung bình, liên kết nhiều tài khoản con."),
        ("Quản Trị Viên\n(Admin)", "Ban quản trị nhà trường, kĩ thuật viên hệ thống", "Quản lý toàn bộ người dùng, quản lý trường học, cấu hình hệ thống, theo dõi biểu đồ phân tích (Analytics), giám sát sức khỏe DB (Health Monitoring).")
    ]

    for row_idx, data in enumerate(roles_data, start=1):
        for col_idx, text in enumerate(data):
            cell = table_role.cell(row_idx, col_idx)
            cell.width = widths[col_idx]
            if row_idx % 2 == 0:
                set_cell_background(cell, "F8FAFC")
            set_cell_margins(cell, top=100, bottom=100, left=120, right=120)
            p = cell.paragraphs[0]
            r = p.add_run(text)
            r.font.size = Pt(9.5)
            if col_idx == 0:
                r.font.bold = True
                r.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # -------------------------------------------------------------
    # PHẦN 3: CÁC TÍNH NĂNG NỔI BẬT (WEB HAS WHAT FEATURES?)
    # -------------------------------------------------------------
    add_h1("PHẦN 3: CÁC TÍNH NĂNG NỔI BẬT CỦA WEBSITE")

    add_h2("3.1. Trình Lắp Ráp PC Mô Phỏng (Builder Core Engine)")
    add_bullet("Mô phỏng chân thực các linh kiện CPU, RAM, Mainboard, GPU, PSU, SSD/HDD, Cooler, Case với thao tác snap chính xác.", "Chế độ Lắp Ráp 2D Drag & Drop: ")
    add_bullet("Mô hình linh kiện 3D GLB xoay 360 độ, cho phép phóng to/thu nhỏ và xem chi tiết cấu trúc phần cứng.", "Chế độ 3D Showroom & Viewer: ")
    add_bullet("Tính tổng công suất tiêu thụ điện (TDP) tức thì và khuyến nghị công suất nguồn (PSU) phù hợp.", "Tính toán TDP thời gian thực: ")
    add_bullet("Tự động phát hiện lỗi sai socket (ví dụ: cắm CPU Intel vào mainboard AMD), sai chuẩn RAM (DDR4 vs DDR5) hoặc kích thước case nhỏ hơn GPU.", "Kiểm tra Tương thích (Compatibility Check): ")
    add_bullet("Mô phỏng màn hình BIOS và Windows 11 khởi động sau khi học sinh hoàn thành lắp ráp thành công.", "Mô phỏng Boot Windows 11: ")

    add_h2("3.2. Điều Khiển Bằng Cử Chỉ Tay (Hand Tracking qua Webcam)")
    add_bullet("Sử dụng Google MediaPipe Tasks Vision để nhận diện 21 điểm khớp trên bàn tay người dùng.", "Công nghệ Computer Vision: ")
    add_bullet("Cho phép người dùng đưa tay trước camera để di chuyển con trỏ chuột, chụm ngón tay (Pinch) để gắp/thả linh kiện, xoay linh kiện 3D mà không cần chạm chuột.", "Thao tác không chạm (Touchless): ")
    add_bullet("Hỗ trợ 9 cử chỉ tay khác nhau như Pinch, Open Palm, Fist, Point Up, Victory...", "Đa dạng cử chỉ: ")

    add_h2("3.3. Trợ Lý Trí Tuệ Nhân Tạo (AI Guru - Google Gemini)")
    add_bullet("Trò chuyện trực tiếp trong builder, đưa ra gợi ý khi học sinh gặp khó khăn hoặc thắc mắc về linh kiện.", "Trợ giảng 24/7: ")
    add_bullet("Tự động gợi ý cấu hình PC tối ưu dựa trên ngân sách ảo và mục tiêu sử dụng (Gaming, Render 3D, Học tập).", "Gợi ý cấu hình thông minh: ")
    add_bullet("Hỗ trợ giáo viên tạo nhanh bộ câu hỏi trắc nghiệm từ nội dung bài học và tóm tắt bài giảng tự động.", "Công cụ hỗ trợ Giáo viên: ")

    add_h2("3.4. Hệ Thống Kiểm Tra & Giám Sát Thi (Proctored Quiz/Exam)")
    add_bullet("Hỗ trợ 5 dạng câu hỏi phong phú: Trắc nghiệm đơn, Trắc nghiệm nhiều đáp án, Đúng/Sai, Điền từ vào chỗ trống, Sắp xếp thứ tự.", "5 Dạng câu hỏi linh hoạt: ")
    add_bullet("Tích hợp tính năng Proctored giám sát camera và phát hiện chuyển tab để đảm bảo tính minh bạch trong các kỳ thi.", "Giám sát thi trực tuyến: ")
    add_bullet("Hệ thống tự động chấm điểm bài làm và lưu trữ lịch sử thi chi tiết cho từng học sinh.", "Tự động chấm điểm: ")

    add_h2("3.5. Hệ Thống Chứng Chỉ Số (PDF & QR Code Verification)")
    add_bullet("Khi hoàn thành khóa học hoặc vượt qua kỳ thi, hệ thống tự động sinh chứng chỉ PDF chất lượng cao.", "Cấp chứng chỉ tự động: ")
    add_bullet("Mỗi chứng chỉ có mã Hash/UUID duy nhất kèm mã QR. Bất kỳ ai cũng có thể quét QR hoặc truy cập `/verify/[code]` để kiểm tra tính hợp lệ.", "Xác thực QR Code: ")

    add_h2("3.6. Các Chế Độ Game Hóa & Mở Rộng")
    add_bullet("Giả lập mua sắm linh kiện với 6 cửa hàng ảo, giúp học sinh luyện kỹ năng quản lý ngân sách.", "Chợ Máy Tính (Virtual Market): ")
    add_bullet("Hai học sinh cùng thi đấu xem ai lắp ráp PC đúng và nhanh hơn.", "Chế độ 2 Người Chơi (Multiplayer): ")
    add_bullet("Tích hợp điểm kinh nghiệm (XP), cấp độ (Level), chuỗi ngày học (Streak), huy hiệu thành tích và pháo hoa ăn mừng (Confetti).", "Gamification: ")

    # -------------------------------------------------------------
    # PHẦN 4: CÔNG NGHỆ SỬ DỤNG VÀ NGUỒN GỐC TÀI NGUYÊN (TECH STACK & SOURCES)
    # -------------------------------------------------------------
    add_h1("PHẦN 4: CÔNG NGHỆ SỬ DỤNG & NGUỒN GỐC TÀI NGUYÊN")

    add_h2("4.1. Bảng Tổng Hợp Công Nghệ (Tech Stack)")
    
    table_tech = doc.add_table(rows=10, cols=3)
    table_tech.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table_tech)

    headers = ["Lĩnh Vực", "Công Nghệ / Thư Viện", "Vai Trò & Mục Đích Sử Dụng"]
    widths = [Inches(1.8), Inches(2.2), Inches(2.5)]
    
    for i, h in enumerate(headers):
        cell = table_tech.cell(0, i)
        cell.width = widths[i]
        set_cell_background(cell, "1E3A8A")
        set_cell_margins(cell, top=120, bottom=120, left=120, right=120)
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.font.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r.font.size = Pt(10)

    tech_data = [
        ("Frontend Core", "Next.js 16.1.6 (App Router), React 19.2.3, TypeScript", "Framework chính cho web application, Server Components & Client rendering tối ưu SEO và hiệu năng."),
        ("Backend & Data", "Supabase (PostgreSQL, Auth, Realtime, Storage), @supabase/ssr", "Xác thực người dùng (Email + Google OAuth), lưu trữ dữ liệu lớp học/bài thi, real-time messaging."),
        ("Styling & UI", "Tailwind CSS, Framer Motion, Lucide React Icons", "Thiết kế giao diện hiện đại, mượt mà, hiệu ứng chuyển trang và biểu tượng chuẩn hóa."),
        ("AI Assistant", "Google Gemini AI SDK (@google/generative-ai)", "Trợ lý AI Guru tư vấn lắp ráp, gợi ý cấu hình PC và sinh đề bài kiểm tra tự động."),
        ("Hand Tracking", "Google MediaPipe Tasks Vision (@mediapipe/tasks-vision)", "Nhận diện cử chỉ tay 21 điểm qua webcam để điều khiển thao tác lắp ráp 2D/3D không cần chuột."),
        ("3D Graphics", "Three.js, @react-three/fiber, @react-three/drei", "Render mô hình 3D linh kiện máy tính, showroom xoay linh kiện và hiệu ứng hạt trong không gian 3D."),
        ("Drag & Drop 2D", "@dnd-kit/core, @dnd-kit/sortable, @dnd-kit/utilities", "Xử lý thao tác kéo-thả linh kiện mượt mà trong trình lắp ráp 2D Builder."),
        ("PDF & QR Code", "@react-pdf/renderer, jspdf, html2canvas, qrcode.react", "Xuất file PDF chứng chỉ học tập, tạo mã QR xác thực chứng chỉ trực tuyến."),
        ("Analytics & State", "TanStack React Query, Recharts, Zustand", "Quản lý state ứng dụng, vẽ biểu đồ thống kê học tập cho Admin, Teacher, Student.")
    ]

    for row_idx, data in enumerate(tech_data, start=1):
        for col_idx, text in enumerate(data):
            cell = table_tech.cell(row_idx, col_idx)
            cell.width = widths[col_idx]
            if row_idx % 2 == 0:
                set_cell_background(cell, "F8FAFC")
            set_cell_margins(cell, top=100, bottom=100, left=120, right=120)
            p = cell.paragraphs[0]
            r = p.add_run(text)
            r.font.size = Pt(9.5)
            if col_idx == 0:
                r.font.bold = True
                r.font.color.rgb = RGBColor(0x0D, 0x94, 0x88)

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    add_h2("4.2. Nguồn Gốc Dữ Liệu & Tài Nguyên (Lấy Từ Đâu?)")

    make_callout(
        doc,
        "Thành viên mới thường thắc mắc các dữ liệu linh kiện, bài giảng, mô hình 3D và AI model được lấy từ nguồn nào. "
        "Dưới đây là thống kê chi tiết nguồn gốc tài nguyên của dự án:",
        "XUẤT XỨ TÀI NGUYÊN & DỮ LIỆU"
    )

    add_bullet("Nội dung bài học, khái niệm phần cứng, bài kiểm tra được biên soạn dựa trên chương trình Sách Giáo Khoa Tin học THPT hiện hành (Bộ sách Kết Nối Tri Thức Với Cuộc Sống & Bộ sách Cánh Diều - Lớp 10, 11, 12).", "1. Nội dung Giáo trình & Bài học: ")
    add_bullet("Thông số kỹ thuật của CPU (socket, TDP, core/thread), Mainboard (chipset, form factor), GPU, RAM (bus, DDR4/DDR5), PSU (wattage) được tổng hợp từ thông số chính thức của các hãng phần cứng lớn như Intel, AMD, Nvidia, ASUS, MSI, Gigabyte, Corsair, Kingston.", "2. Dữ liệu Linh kiện PC: ")
    add_bullet("Các file 3D GLB/gLTF của linh kiện (Case, GPU, RAM, CPU...) được thu thập từ các kho tài nguyên 3D mã nguồn mở và miễn phí bản quyền như Sketchfab, Poly Pizza, Kenney.nl, Itch.io, OpenGameArt với giấy phép CC0 hoặc Free Commercial License.", "3. Mô hình 3D (3D Assets): ")
    add_bullet("Sử dụng Pre-trained Hand Landmarker Model do Google nghiên cứu và phát hành công khai qua thư viện MediaPipe, cho phép chạy trực tiếp trên trình duyệt WebAssembly mà không cần server xử lý nặng.", "4. AI Model Hand Tracking: ")
    add_bullet("Sử dụng Google Gemini API với System Prompt được thiết kế chuyên biệt (Custom System Instructions) để định hình cá tính 'AI Guru' am hiểu phần cứng máy tính và sư phạm.", "5. AI Knowledge Base: ")

    # -------------------------------------------------------------
    # PHẦN 5: CẤU TRÚC DỰ ÁN & HƯỚNG DẪN DÀNH CHO THÀNH VIÊN MỚI
    # -------------------------------------------------------------
    add_h1("PHẦN 5: CẤU TRÚC DỰ ÁN & QUY TRÌNH LÀM VIỆC NHÓM")

    add_h2("5.1. Cấu Trúc Thư Mục Dự Án (Source Code Layout)")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.add_run("Dự án được tổ chức theo chuẩn Next.js App Router:")

    folders_info = [
        ("app/", "Chứa các Route & Pages của ứng dụng (app/builder, app/student, app/teacher, app/parent, app/admin, app/quiz...)"),
        ("components/", "Chứa các UI Components tái sử dụng. Bao gồm components 2D builder, 3D viewer (GameScene, ShowroomScene), Hand tracker, Quiz, Profile..."),
        ("lib/", "Chứa Server Actions, hàm tiện ích, cấu hình Supabase Auth, RBAC middleware và Grading Engine chấm điểm."),
        ("hooks/", "Custom React Hooks (ví dụ: useStore, useAuth, useHandTracker)."),
        ("database/", "Chứa SQL scripts, migrations schema Supabase."),
        ("public/", "Chứa static assets: hình ảnh linh kiện, âm thanh hiệu ứng, file 3D (.glb)."),
        ("types/", "Chứa các kiểu dữ liệu TypeScript (User, Component, Class, Quiz, Certificate...).")
    ]
    for f, desc in folders_info:
        add_bullet(desc, bold_prefix=f"{f} — ")

    add_h2("5.2. Quy Trình Phối Hợp Git & Branching Strategy")
    
    make_callout(
        doc,
        "Để tránh xung đột code (merge conflict) giữa các thành viên, nhóm áp dụng quy tắc phân chia nhánh Git như sau:\n"
        "• main: Branch Production, tự động deploy lên https://pc-master-lms.vercel.app\n"
        "• feature/2d-renderer: Nhánh phát triển module 2D Builder & Animation\n"
        "• feature/3d-viewer: Nhánh phát triển 3D Viewer & MediaPipe Gesture Control\n\n"
        "QUY TRÌNH COMMIT & DEPLOY:\n"
        "1. git checkout -b feature/<tên-tính-năng>\n"
        "2. Chạy `npm run build` ở local trước khi push để đảm bảo không lỗi TypeScript/Lint!\n"
        "3. git add . && git commit -m 'Mô tả thay đổi rõ ràng'\n"
        "4. git push origin feature/<tên-tính-năng>\n"
        "5. Tạo Pull Request -> Review -> Merge vào main.",
        "QUY TẮC GIT & DEPLOY (BẮT BUỘC ĐỌC)"
    )

    add_h2("5.3. Hướng Dẫn Chạy Dự Án Local (Getting Started)")
    add_bullet("Tải và cài đặt Node.js phiên bản 18 trở lên (khuyên dùng v20 LTS).", "Bước 1: ")
    add_bullet("Mở terminal tại thư mục dự án và chạy câu lệnh `npm install` để cài đặt dependencies.", "Bước 2: ")
    add_bullet("Tạo file `.env.local` từ file `.env.example` và điền các thông số Supabase URL, Anon Key, Gemini API Key.", "Bước 3: ")
    add_bullet("Chạy câu lệnh `npm run dev` và truy cập `http://localhost:3000` trên trình duyệt.", "Bước 4: ")

    # Footer note
    doc.add_paragraph().paragraph_format.space_before = Pt(20)
    p_end = doc.add_paragraph()
    p_end.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_end = p_end.add_run("--- CHÚC BẠN CÓ TRẢI NGHIỆM LÀM VIỆC TUYỆT VỜI CÙNG NHÓM PC MASTER LMS! ---")
    r_end.font.bold = True
    r_end.font.size = Pt(11)
    r_end.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)

    output_path = "Gioi_Thieu_Du_An_PC_Master_LMS.docx"
    doc.save(output_path)
    print("Document created successfully!")

if __name__ == "__main__":
    build_document()
