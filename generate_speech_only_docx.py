import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def create_speech_docx():
    doc = docx.Document()

    # Page Margins
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    COLOR_NAVY = RGBColor(3, 31, 59)       # #031F3B
    COLOR_PRIMARY = RGBColor(8, 158, 96)   # #089E60
    COLOR_DARK = RGBColor(30, 41, 59)      # #1E293B
    COLOR_BLUE = RGBColor(14, 116, 144)    # #0E7490
    COLOR_ACCENT = RGBColor(217, 119, 6)   # #D97706

    # Header
    p_t = doc.add_paragraph()
    p_t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_t = p_t.add_run("KỊCH BẢN NÓI THUYẾT TRÌNH 5 PHÚT ĐỈNH CAO (CHÍNH XÁC 300 GIÂY)\nCHUYÊN DÀNH CHO ĐỘI THI PC MASTER BUILDER - STARTUP HUIT 2026")
    r_t.font.name = 'Arial'
    r_t.font.size = Pt(16)
    r_t.font.bold = True
    r_t.font.color.rgb = COLOR_NAVY

    p_s = doc.add_paragraph()
    p_s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_s = p_s.add_run("Dự án: PC MASTER BUILDER (PC Master LMS) | Thể hiện: Khánh Sơn & Quốc An\nPhân đoạn chi tiết từng giây - Lời thoại cuốn hút - Thuật ngữ đỉnh cao - Phụng sự 10 UN SDGs")
    r_s.font.name = 'Arial'
    r_s.font.size = Pt(10.5)
    r_s.font.italic = True
    r_s.font.color.rgb = COLOR_PRIMARY

    doc.add_paragraph()

    # Box guide
    tbl_in = doc.add_table(rows=1, cols=1)
    tbl_in.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl_in.cell(0, 0)
    cell.width = Inches(6.8)
    shd = parse_xml(r'<w:shd {} w:fill="F8FAFC"/>'.format(nsdecls('w')))
    cell._tc.get_or_add_tcPr().append(shd)

    p_in = cell.paragraphs[0]
    p_in.paragraph_format.space_before = Pt(6)
    p_in.paragraph_format.space_after = Pt(6)
    r_in = p_in.add_run(
        "📌 HƯỚNG DẪN LUYỆN TẬP CHO 2 THÍ SÍNH SƠN & AN:\n"
        "• Kịch bản được phân chia chi tiết thành 6 Đoạn (tương ứng 6 Slide chính), phân vai rõ ràng từng câu thoại.\n"
        "• Đã lồng ghép mượt mà 10 Mục tiêu Phát triển Bền vững (UN SDGs: SDG 4, 9, 12, 8, 10, 13, 5, 17, 3, 11) cùng toàn bộ thuật ngữ chuyên môn đắt giá (Edge AI, WebGL 3D, Hand-Tracking, Freemium PLG, LTV/CAC 36x, Break-even 250 users...).\n"
        "• Chú ý giọng đọc truyền cảm, nhấn giọng ở các từ in đậm và giữ đúng nhịp thời gian quy định!"
    )
    r_in.font.name = 'Arial'
    r_in.font.size = Pt(9.5)
    r_in.font.color.rgb = COLOR_DARK

    doc.add_paragraph()

    sections_data = [
        {
            "part": "ĐOẠN 1: BẮT ĐẦU & MỞ MÀN ẤN TƯỢNG (SLIDE 1)",
            "time": "⏱️ Thời lượng: 00:00 - 00:30 (Đúng 30 giây)",
            "speaker": "🎤 Người nói: Khánh Sơn (Giọng tự tin, dõng dạc, mỉm cười chào BGK)",
            "action": "🎬 Thao tác Slide: Bật Slide 1 (Bìa dự án Neon Cyberpunk), Sơn giơ điện thoại quét mã QR Web pc-master-lms.vercel.app trực tiếp trên màn hình.",
            "script": [
                ("Sơn:", "Kính chào Ban Giám khảo và toàn thể hội thi! Em là Khánh Sơn - Developer chính, và đây là Quốc An - phụ trách Thị trường, đại diện cho nhóm tác giả trường THPT Nguyễn Công Trứ và HUIT đến với cuộc thi Khởi nghiệp HUIT 2026!"),
                ("Sơn:", "Hôm nay, chúng em tự hào mang tới giải pháp: PC Master Builder - Nền tảng Giáo dục Thực hành Lắp ráp PC 3D & Hệ thống LMS Thông minh tích hợp AI đầu tiên tại Việt Nam, sẵn sàng giải quyết triệt để bài toán giáo dục phần cứng cho thế hệ Gen Z!")
            ],
            "keywords": "📌 Keyword & SDGs xuất hiện: EdTech Make-in-Vietnam, Interactive 3D Simulation, QR-Code Realtime Access."
        },
        {
            "part": "ĐOẠN 2: THỰC TRẠNG & NỖI ĐẠO THỊ TRƯỜNG BẤT ĐỒNG THÔNG TIN (SLIDE 2)",
            "time": "⏱️ Thời lượng: 00:30 - 01:10 (Đúng 40 giây)",
            "speaker": "🎤 Người nói: Khánh Sơn (Giọng trăn trở, nhấn mạnh vào nỗi đau thực tế)",
            "action": "🎬 Thao tác Slide: Chuyển Slide 2 (Con số 85% trường thiếu phòng lab đỏ rực & icon linh kiện hỏng/chặt chém giá).",
            "script": [
                ("Sơn:", "Thưa Ban Giám khảo, Chương trình GDPT 2018 bắt buộc 100% học sinh Lớp 10 phải học phần cứng PC. Tuy nhiên, một thực trạng đáng buồn là hơn 85% các trường THPT hiện nay KHÔNG CÓ PHÒNG LAB THỰC HÀNH!"),
                ("Sơn:", "Trang bị một phòng lab thật cho 40 học sinh tiêu tốn tới 800 TRIỆU ĐỒNG, chưa kể rủi ro cong chân socket CPU hay cháy linh kiện do tĩnh điện (ESD) cực kỳ cao. Học sinh bắt buộc phải 'học chay' qua hình vẽ 2D khô khan."),
                ("Sơn:", "Song song đó, khi tự mua máy tính ngoài đời, phụ huynh và học sinh chịu sự Bất Đồng Thông Tin (Information Asymmetry) nghiêm trọng và thường xuyên bị các cửa hàng 'chặt chém' chênh giá tồn kho!")
            ],
            "keywords": "📌 Keyword & SDGs xuất hiện: Information Asymmetry, ESD Component Risk, Cost Barrier 800M VND, Passive 2D Learning."
        },
        {
            "part": "ĐOẠN 3: GIẢI PHÁP ĐỘT PHÁ CÔNG NGHỆ 3D/AI & CỬ CHỈ TAY (SLIDE 3)",
            "time": "⏱️ Thời lượng: 01:10 - 02:00 (Đúng 50 giây)",
            "speaker": "🎤 Người nói: Khánh Sơn (Giọng hào hứng, tự hào thể hiện sức mạnh công nghệ)",
            "action": "🎬 Thao tác Slide: Chuyển Slide 3 (Video mượt 5s cử chỉ tay Hand-Tracking bốc RAM 3D & AI Chống chặt chém).",
            "script": [
                ("Sơn:", "PC Master Builder giải quyết triệt để nỗi đau đó với Phòng Lab Ảo 3D 100% trên Web - Chi phí thiết bị bằng 0 đồng!"),
                ("Sơn:", "Đột phá lớn nhất của dự án nằm ở Công nghệ Điều khiển Cử chỉ tay (Hand-Tracking) bằng Edge AI MediaPipe Vision: Học sinh chỉ cần giơ bàn tay trước Webcam laptop là có thể co, nắm ngón tay để bốc, xoay 360 độ và gắn RAM, CPU, VGA vào Mainboard như thực tế - KHÔNG CẦN KÍNH VR ĐẮT TIỀN!"),
                ("Sơn:", "Đi kèm là Trợ lý AI Chống Chặt Chém tự động quét API giá thị trường real-time, kiểm tra tương thích Socket và tính tổng TDP công suất nguồn, đảm bảo an toàn tuyệt đối cho người học!")
            ],
            "keywords": "📌 Keyword & SDGs xuất hiện: WebGL Three.js, MediaPipe 21 Hand Landmarks (Edge AI), Zero-VR-Friction, Real-time Price API."
        },
        {
            "part": "ĐOẠN 4: THỰC NGHIỆM VALIDATION & PHỤNG SỰ 10 UN SDGS (SLIDE 4)",
            "time": "⏱️ Thời lượng: 02:00 - 02:50 (Đúng 50 giây)",
            "speaker": "🎤 Người nói: Đặng Quốc An (Giọng mạnh mẽ, thuyết phục bằng số liệu thực tế)",
            "action": "🎬 Thao tác Slide: Chuyển Slide 4 (Biểu đồ +28% điểm số, 76 học sinh 10A1/10C9 & Vòng tròn 10 UN SDGs).",
            "script": [
                ("An:", "Sản phẩm của chúng em không nằm trên giấy! PC Master đã vận hành thực tế tại pc-master-lms.vercel.app và được kiểm chứng trên 76 học sinh Lớp 10A1 và 10C9 THPT Nguyễn Công Trứ."),
                ("An:", "Kết quả thực nghiệm A/B Testing chứng minh: Điểm thi thực hành phần cứng TĂNG 28%, Tỷ lệ hoàn thành bài học đạt 94% và 92% học sinh vô cùng yêu thích!"),
                ("An:", "Đặc biệt thưa BGK, PC Master tự hào phụng sự 10 MỤC TIÊU PHÁT TRIỂN BỀN VỮNG (UN SDGs): Nổi bật là SDG 4 (Giáo dục chất lượng 3D), SDG 9 (Hạ tầng công nghệ AI/WebGL), SDG 12 (Cắt giảm 100% rác thải điện tử E-waste), SDG 8 (Định hướng nhân lực Vi mạch - Bán dẫn), SDG 10 (Bình đẳng công nghệ số vùng xa), SDG 13 (Hành động khí hậu giảm dấu chân Carbon), SDG 5 (Bình đẳng giới STEM cho nữ sinh), SDG 3 (An toàn điện tuyệt đối), SDG 11 (Trường học thông minh) và SDG 17 (Hợp tác đa bên)!")
            ],
            "keywords": "📌 Keyword & SDGs xuất hiện: Empirical A/B Testing (+28% Score), 10 UN SDGs (SDG 4, 9, 12, 8, 10, 13, 5, 3, 11, 17), E-waste Elimination."
        },
        {
            "part": "ĐOẠN 5: MÔ HÌNH KINH DOANH, TÀI CHÍNH LEAN & PHÂN BỔ VỐN (SLIDE 5)",
            "time": "⏱️ Thời lượng: 02:50 - 03:40 (Đúng 50 giây)",
            "speaker": "🎤 Người nói: Đặng Quốc An (Giọng sắc bén, thể hiện tư duy tài chính chuyên nghiệp)",
            "action": "🎬 Thao tác Slide: Chuyển Slide 5 (3 Hộp dòng tiền B2C/B2B/Affiliate & Biểu đồ phân bổ 50 triệu gọi vốn).",
            "script": [
                ("An:", "Về Mô hình Kinh doanh, PC Master xây dựng 3 Nguồn thu bền vững: Gói B2C Pro chỉ 20.000đ/tháng - bằng giá 1 ly trà sữa bình dân; Gói B2B LMS cho Nhà trường (15-25 triệu/năm) tiết kiệm hàng trăm triệu tiền phòng lab; và Hoa hồng Affiliate 2-5% từ các đại lý bán sỉ linh kiện lớn!"),
                ("An:", "Nhờ kiến trúc Serverless Cloud tinh gọn, chi phí duy trì chỉ ~2 triệu/tháng. Dự án đạt ĐIỂM HUỀ VỐN (Break-even) chỉ với 250 người dùng trả phí và đạt biên lợi nhuận ròng SaaS >75% cùng tỷ lệ LTV/CAC ấn tượng 36 lần!"),
                ("An:", "Nếu giành 50 TRIỆU ĐỒNG tiền thưởng giải Nhất HUIT, nhóm sẽ phân bổ chuẩn xác: 40% Nâng cấp Cloud Server & Gemini AI API; 40% Marketing B2B phủ sóng 10 trường THPT; và 20% Đăng ký Bảo hộ Bản quyền Tác giả!")
            ],
            "keywords": "📌 Keyword & SDGs xuất hiện: Freemium PLG Model, Psychological Pricing (20k/mo), Break-even 250 users, LTV/CAC 36x, Net Margin >75%."
        },
        {
            "part": "ĐOẠN 6: ĐỘI NGŨ FOUNDER MAKE-IN-HUIT & THÔNG ĐỆP KẾT ĐANH THÉP (SLIDE 6)",
            "time": "⏱️ Thời lượng: 03:40 - 04:30 (Đúng 50 giây)",
            "speaker": "🎤 Người nói: Khánh Sơn (Giọng hào hùng, kết thúc bùng nổ cảm xúc)",
            "action": "🎬 Thao tác Slide: Chuyển Slide 6 (Hình ảnh 4 Founder, Giấy chứng nhận Bản quyền IP & Logo HUIT).",
            "script": [
                ("Sơn:", "Đội ngũ 4 Founder HUIT của chúng em TỰ CHỦ KỸ THUẬT 100% từ Fullstack, 3D WebGL đến AI Computer Vision mà không tốn 1 đồng thuê ngoài! Dự án được bảo chứng học thuật bởi GVHD Cô Kim Phượng và Thầy Minh Phụng, đồng thời đã hoàn tất Đăng ký Bản quyền Tác giả phần mềm."),
                ("Sơn:", "Kính thưa Ban Giám khảo, PC Master Builder xứng đáng giành vị trí QUÁN QUÂN HUIT 2026 vì đây là một sản phẩm EdTech Make-in-HUIT 100% tự chủ công nghệ, giải đúng nỗi đau giáo dục phổ thông bằng công nghệ 3D/AI đột phá, phụng sự 10 mục tiêu phát triển bền vững và mang lại giá trị nhân văn to lớn cho nền giáo dục Việt Nam!"),
                ("Sơn & An cùng cúi chào:", "Em xin chân thành cảm ơn Ban Giám khảo đã chú ý lắng nghe!")
            ],
            "keywords": "📌 Keyword & SDGs xuất hiện: 100% Technical Autonomy, Make-in-HUIT Pride, Copyright IP Safeguard, Champion Impact Vision."
        }
    ]

    for sec in sections_data:
        p_p = doc.add_paragraph()
        p_p.paragraph_format.space_before = Pt(14)
        p_p.paragraph_format.space_after = Pt(4)
        r_p = p_p.add_run(sec["part"])
        r_p.font.name = 'Arial'
        r_p.font.size = Pt(12)
        r_p.font.bold = True
        r_p.font.color.rgb = COLOR_NAVY

        p_m = doc.add_paragraph()
        p_m.paragraph_format.space_before = Pt(2)
        p_m.paragraph_format.space_after = Pt(2)
        r_tm = p_m.add_run(f"{sec['time']}  |  {sec['speaker']}\n")
        r_tm.font.name = 'Arial'
        r_tm.font.size = Pt(9.5)
        r_tm.font.bold = True
        r_tm.font.color.rgb = COLOR_BLUE

        r_ac = p_m.add_run(f"{sec['action']}")
        r_ac.font.name = 'Arial'
        r_ac.font.size = Pt(9.5)
        r_ac.font.italic = True
        r_ac.font.color.rgb = COLOR_ACCENT

        # Speech table
        tbl_s = doc.add_table(rows=0, cols=2)
        tbl_s.alignment = WD_TABLE_ALIGNMENT.CENTER
        tbl_s.autofit = False

        for spk, text in sec["script"]:
            row = tbl_s.add_row()
            c0 = row.cells[0]
            c1 = row.cells[1]
            c0.width = Inches(1.1)
            c1.width = Inches(5.7)

            shd = parse_xml(r'<w:shd {} w:fill="F1F5F9"/>'.format(nsdecls('w')))
            c0._tc.get_or_add_tcPr().append(shd)

            p0 = c0.paragraphs[0]
            p0.paragraph_format.space_before = Pt(4)
            p0.paragraph_format.space_after = Pt(4)
            r0 = p0.add_run(spk)
            r0.font.name = 'Arial'
            r0.font.size = Pt(9.5)
            r0.font.bold = True
            r0.font.color.rgb = COLOR_NAVY

            p1 = c1.paragraphs[0]
            p1.paragraph_format.space_before = Pt(4)
            p1.paragraph_format.space_after = Pt(4)
            r1 = p1.add_run(text)
            r1.font.name = 'Arial'
            r1.font.size = Pt(9.5)
            r1.font.color.rgb = COLOR_DARK

        p_kw = doc.add_paragraph()
        p_kw.paragraph_format.space_before = Pt(4)
        p_kw.paragraph_format.space_after = Pt(8)
        r_kw = p_kw.add_run(sec["keywords"])
        r_kw.font.name = 'Arial'
        r_kw.font.size = Pt(9)
        r_kw.font.bold = True
        r_kw.font.color.rgb = COLOR_PRIMARY

        doc.add_paragraph()

    filename = "Kich_Ban_Noi_Thuyet_Trinh_5_Phut_Dinh_Cao.docx"
    doc.save(filename)
    print(f"Speech document created successfully: {filename}")

if __name__ == '__main__':
    create_speech_docx()
