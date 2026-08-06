import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def build_speech_word():
    doc = docx.Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    COLOR_NAVY  = RGBColor(3, 31, 59)
    COLOR_GREEN = RGBColor(8, 158, 96)
    COLOR_DARK  = RGBColor(30, 41, 59)
    COLOR_BLUE  = RGBColor(14, 116, 144)

    # ---- HEADER ----
    p_h = doc.add_paragraph()
    p_h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_h = p_h.add_run("BÀI THUYẾT TRÌNH CHÍNH THỨC")
    r_h.font.name = "Arial"
    r_h.font.size = Pt(20)
    r_h.font.bold = True
    r_h.font.color.rgb = COLOR_NAVY

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_sub = p_sub.add_run("DỰ ÁN: PC MASTER BUILDER (PC Master LMS)")
    r_sub.font.name = "Arial"
    r_sub.font.size = Pt(13)
    r_sub.font.bold = True
    r_sub.font.color.rgb = COLOR_GREEN

    p_sub2 = doc.add_paragraph()
    p_sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_sub2 = p_sub2.add_run("Cuộc thi Khởi nghiệp & Sáng tạo HUIT 2026 | Bảng Học Sinh THPT")
    r_sub2.font.name = "Arial"
    r_sub2.font.size = Pt(11)
    r_sub2.font.italic = True
    r_sub2.font.color.rgb = COLOR_DARK

    doc.add_paragraph()

    def add_line(text, bold=False, size=11, color=None, center=False, italic=False, space_before=3, space_after=3):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.line_spacing = 1.4
        if center:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.font.name = "Arial"
        r.font.size = Pt(size)
        r.font.bold = bold
        r.font.italic = italic
        r.font.color.rgb = color or COLOR_DARK
        return p

    def add_divider():
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run("─" * 60)
        r.font.name = "Arial"
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(200, 210, 220)

    def add_section_label(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(f"【 {text} 】")
        r.font.name = "Arial"
        r.font.size = Pt(10)
        r.font.bold = True
        r.font.color.rgb = COLOR_BLUE

    # ---- OPENING ----
    add_section_label("MỞ ĐẦU")
    add_line("Kính thưa Ban Giám khảo và quý thầy cô!", bold=True, size=11.5)
    doc.add_paragraph()
    add_line(
        "Em tên là Khánh Sơn, cùng bạn Quốc An đại diện cho nhóm học sinh trường THPT Nguyễn Công Trứ, "
        "thành phố Hồ Chí Minh, xin phép được trình bày dự án PC Master Builder — nền tảng phòng lab mô phỏng "
        "lắp ráp máy tính 3D và AI đầu tiên tại Việt Nam dành riêng cho học sinh phổ thông."
    )

    add_divider()

    # ---- PROBLEM ----
    add_section_label("VẤN ĐỀ & NỖI ĐAU THỰC TẾ")
    add_line(
        "Thưa quý vị, chúng em muốn bắt đầu bằng một câu hỏi rất thực tế: Học sinh chúng em học phần cứng máy tính như thế nào?",
        italic=True
    )
    doc.add_paragraph()
    add_line(
        "Theo Chương trình GDPT 2018 mới, môn Tin học Lớp 10 bắt buộc học sinh phải nắm vững cấu tạo và "
        "quy trình lắp ráp máy tính. Nhưng thực tế, hơn 85% trường THPT trên cả nước không có phòng thực "
        "hành phần cứng, bởi vì chi phí trang bị một phòng lab với 40 bộ máy thật lên tới hơn 800 triệu đồng "
        "— một con số quá lớn so với ngân sách của hầu hết các trường. Chưa kể, mỗi lần học sinh thực hành "
        "tháo lắp thật, một chiếc socket CPU bị gãy chân là thiệt hại ngay 3 đến 5 triệu đồng, chưa kể rủi "
        "ro cháy nổ do tĩnh điện."
    )
    doc.add_paragraph()
    add_line(
        "Kết quả là, chúng em — những học sinh lớp 10 — chỉ được \"học chay\" qua những hình vẽ 2D khô khan trong "
        "sách giáo khoa. Ra ngoài tự mua máy tính học tập, chúng em và phụ huynh hoàn toàn bị bất đồng thông tin "
        "và rất dễ bị nhân viên cửa hàng đẩy giá linh kiện tồn kho."
    )
    doc.add_paragraph()
    add_line("Đó là ba nỗi đau thực tế mà PC Master Builder ra đời để giải quyết triệt để.", bold=True)

    add_divider()

    # ---- SOLUTION ----
    add_section_label("GIẢI PHÁP & CÔNG NGHỆ ĐỘT PHÁ")
    add_line(
        "Vì thế, nhóm chúng em đã tự tay xây dựng PC Master Builder — một phòng lab thực hành ảo hoàn toàn "
        "trên nền Web, miễn phí 100%, không cần cài đặt, không cần mua bất kỳ thiết bị nào."
    )
    doc.add_paragraph()
    add_line(
        "Điểm đặc biệt và độc đáo nhất mà nhóm em tự làm được là công nghệ Hand-Tracking — nhận diện cử chỉ tay "
        "qua Webcam. Bạn học sinh chỉ cần giơ bàn tay trước camera laptop là có thể co ngón tay để bốc cây RAM, "
        "xoay GPU 360 độ và gắn từng linh kiện vào bo mạch chủ như ngoài đời thực — mà không cần mua kính VR đắt "
        "tiền. Đây là ứng dụng thư viện MediaPipe của Google, xử lý AI hoàn toàn trên trình duyệt, không gửi video "
        "về máy chủ, bảo mật tuyệt đối và chạy mượt ngay cả trên máy tính văn phòng cũ kỹ của các trường."
    )
    doc.add_paragraph()
    add_line(
        "Bên cạnh đó, chúng em còn tích hợp Trợ lý AI Guru hoạt động như một giáo viên phụ tá ảo, tự động kiểm "
        "tra tính tương thích socket CPU với Mainboard, tính tổng công suất TDP của cả bộ máy và cảnh báo ngay lập "
        "tức khi học sinh lắp sai. Đặc biệt, chức năng Chống Chặt Chém Giá kết nối API các đại lý lớn như Phong Vũ, "
        "GearVN để hiển thị giá thị trường thật, bảo vệ phụ huynh khi đi mua linh kiện cho con."
    )
    doc.add_paragraph()
    add_line(
        "Toàn bộ hệ thống còn là một LMS đa vai trò hoàn chỉnh: Học sinh học theo lộ trình cá nhân hóa và nhận "
        "chứng chỉ xác thực bằng mã QR. Giáo viên có Dashboard quản lý lớp, tự tạo bài giảng và giao bài tập. "
        "Phụ huynh theo dõi thời gian thực tiến độ học của con. Tất cả bảo mật phân quyền tuyệt đối bằng Supabase "
        "Row Level Security chuẩn GDPR."
    )

    add_divider()

    # ---- VALIDATION + SDGs ----
    add_section_label("KẾT QUẢ THỰC NGHIỆM & 10 MỤC TIÊU PHÁT TRIỂN BỀN VỮNG")
    add_line(
        "Đây không phải dự án lý thuyết suông. Chúng em đã đưa ứng dụng chạy thực tế tại pc-master-lms.vercel.app "
        "và tổ chức thử nghiệm trực tiếp trên 76 học sinh lớp 10A1 và 10C9 tại trường trong 4 tuần liên tục."
    )
    doc.add_paragraph()
    add_line(
        "Kết quả đo lường định lượng thực sự thuyết phục: Điểm thi thực hành của các bạn tăng trung bình 28%. "
        "Tỷ lệ hoàn thành toàn bộ bài học đạt 94%. Và 92% học sinh phản hồi rằng học với PC Master Builder dễ "
        "hiểu và thú vị hơn hẳn phương pháp truyền thống.",
        bold=True
    )
    doc.add_paragraph()
    add_line("Đặc biệt, dự án của nhóm học sinh chúng em đóng góp thiết thực vào 10 Mục tiêu Phát triển Bền vững của Liên Hợp Quốc:")
    doc.add_paragraph()

    sdgs = [
        ("SDG 4 — Giáo dục chất lượng", "Chúng em đưa bài học phần cứng 3D chuẩn SGK đến tay mọi học sinh dù trường có phòng lab hay không."),
        ("SDG 9 — Công nghiệp và Đổi mới sáng tạo", "Chúng em tiên phong ứng dụng WebGL, Computer Vision và Cloud AI vào giáo dục phổ thông Việt Nam."),
        ("SDG 12 — Tiêu dùng có trách nhiệm", "Không còn linh kiện nào bị hỏng vì thực hành sai — chúng em xóa bỏ 100% rác thải điện tử E-waste."),
        ("SDG 8 — Việc làm tốt", "Module Career Build định hướng học sinh vào ngành Vi mạch và Bán dẫn đang thiếu nhân lực trầm trọng."),
        ("SDG 10 — Giảm bất bình đẳng", "Học sinh vùng sâu vùng xa không có phòng lab vẫn học được như học sinh thành phố."),
        ("SDG 13 — Hành động vì khí hậu", "Giảm dấu chân carbon khi không cần sản xuất, vận chuyển linh kiện thực hành."),
        ("SDG 5 — Bình đẳng giới", "Khuyến khích nữ sinh tự tin tham gia học kỹ thuật phần cứng trong ngành STEM."),
        ("SDG 3 — Sức khỏe & An toàn", "Loại bỏ hoàn toàn nguy cơ giật điện, cháy nổ và tai nạn do kim loại sắc nhọn trong thực hành thật."),
        ("SDG 11 — Thành phố bền vững", "Đóng góp vào mô hình Trường học thông minh và Cộng đồng tri thức số tại địa phương."),
        ("SDG 17 — Hợp tác vì mục tiêu chung", "Liên kết 4 bên: Trường ĐH HUIT + Trường THPT + Đại lý Bán lẻ + Nền tảng EdTech."),
    ]

    for title, desc in sdgs:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.3
        r1 = p.add_run(title + ": ")
        r1.font.name = "Arial"
        r1.font.size = Pt(10)
        r1.font.bold = True
        r1.font.color.rgb = COLOR_NAVY
        r2 = p.add_run(desc)
        r2.font.name = "Arial"
        r2.font.size = Pt(10)
        r2.font.color.rgb = COLOR_DARK

    add_divider()

    # ---- BUSINESS MODEL ----
    add_section_label("MÔ HÌNH KINH DOANH & TÀI CHÍNH")
    add_line(
        "Về mô hình phát triển bền vững, là học sinh nên chúng em hiểu rõ tâm lý các bạn. Chúng em chọn mức "
        "giá Gói Pro chỉ 20.000 VNĐ mỗi tháng — đúng bằng giá một ly trà sữa bình dân — để bạn nào cũng có "
        "thể tự đăng ký. Song song đó, chúng em phát triển Gói B2B LMS cho nhà trường với giá 15 đến 25 triệu "
        "đồng mỗi năm, giúp nhà trường tiết kiệm hàng trăm triệu tiền mua máy thật trong khi học sinh vẫn được "
        "học đầy đủ."
    )
    doc.add_paragraph()
    add_line(
        "Nhờ hạ tầng Cloud Serverless tinh gọn, chi phí duy trì hệ thống chỉ khoảng 2 triệu đồng mỗi tháng. "
        "Dự án đạt điểm hòa vốn chỉ với 250 học sinh trả phí và đạt biên lợi nhuận ròng trên 75% — hoàn toàn "
        "khả thi để tự vận hành lâu dài mà không phụ thuộc tài trợ bên ngoài.",
        bold=True
    )
    doc.add_paragraph()
    add_line(
        "Nếu may mắn nhận được 50 triệu tiền thưởng từ cuộc thi HUIT, chúng em sẽ dành 40% nâng cấp server AI "
        "và mô hình 3D, 40% giới thiệu sản phẩm đến 10 trường THPT bạn tại TP.HCM, và 20% đăng ký bản quyền "
        "tác giả chính thức để bảo hộ sản phẩm."
    )

    add_divider()

    # ---- TEAM ----
    add_section_label("ĐỘI NGŨ THỰC HIỆN")
    add_line(
        "Về đội ngũ thực hiện, dưới sự hướng dẫn tận tình của Cô Đoàn Thụy Kim Phượng và Thầy Trần Minh Phụng, "
        "toàn bộ mã nguồn Next.js, mô hình 3D WebGL và thuật toán AI đều do 4 học sinh nhóm em tự tay lập trình "
        "100% mà không thuê ngoài một dòng code nào. Nhóm em làm việc theo mô hình Agile, quản lý code qua Git "
        "branching chuyên nghiệp và đã hoàn thành hơn 42 task công việc qua 3 Sprint phát triển. Chúng em cũng đã "
        "hoàn tất hồ sơ đăng ký Bản quyền Tác giả cho phần mềm để bảo hộ sản phẩm Make-in-HUIT này."
    )

    add_divider()

    # ---- CLOSING ----
    add_section_label("LỜI KẾT")
    add_line(
        "Thưa Ban Giám khảo và quý thầy cô, PC Master Builder của chúng em không chỉ là một sản phẩm công nghệ "
        "giáo dục — đây là câu trả lời của chính những học sinh đang chịu thiệt thòi cho bài toán mà hệ thống "
        "giáo dục chưa giải được."
    )
    doc.add_paragraph()
    add_line(
        "Chúng em tin rằng mỗi học sinh Việt Nam, dù ở thành phố hay vùng xa, đều xứng đáng được học phần cứng "
        "máy tính bằng công nghệ 3D hiện đại nhất — và PC Master Builder chính là cách chúng em biến điều đó "
        "thành sự thật."
    )
    doc.add_paragraph()
    add_line(
        "Em xin chân thành cảm ơn Ban Giám khảo và quý thầy cô đã lắng nghe. "
        "Nhóm em sẵn sàng trả lời mọi câu hỏi phản biện!",
        bold=True, size=12, color=COLOR_GREEN
    )

    doc.add_paragraph()
    p_link = doc.add_paragraph()
    p_link.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_link = p_link.add_run("🌐 Trải nghiệm ngay: pc-master-lms.vercel.app")
    r_link.font.name = "Arial"
    r_link.font.size = Pt(11)
    r_link.font.bold = True
    r_link.font.color.rgb = COLOR_BLUE

    out = "Loi_Thuyet_Trinh_Chinh_Thuc_PC_Master_HUIT_2026.docx"
    doc.save(out)
    print(f"OK - File thuyet trinh chinh thuc da tao: {out}")

build_speech_word()
