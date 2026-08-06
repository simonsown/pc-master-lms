import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def create_student_speech_docx():
    doc = docx.Document()

    # Page Margins
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    COLOR_NAVY = RGBColor(3, 31, 59)       # #031F3B
    COLOR_PRIMARY = RGBColor(8, 158, 96)   # #089E60
    COLOR_DARK = RGBColor(30, 41, 59)      # #1E293B
    COLOR_BLUE = RGBColor(14, 116, 144)    # #0E7490

    # Header
    p_t = doc.add_paragraph()
    p_t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_t = p_t.add_run("BÀI NÓI THUYẾT TRÌNH 5 PHÚT DÀNH CHO HỌC SINH (CHUẨN 100 ĐIỂM HUIT 2026)")
    r_t.font.name = 'Arial'
    r_t.font.size = Pt(16)
    r_t.font.bold = True
    r_t.font.color.rgb = COLOR_NAVY

    p_s = doc.add_paragraph()
    p_s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_s = p_s.add_run("Dự án: PC Master Builder (PC Master LMS) | Giọng nói học sinh chân thực - Mạch nối mượt mà - Đủ ý chắc chắn 100 điểm")
    r_s.font.name = 'Arial'
    r_s.font.size = Pt(10.5)
    r_s.font.italic = True
    r_s.font.color.rgb = COLOR_PRIMARY

    doc.add_paragraph()

    speech_sections = [
        ("[0:00 - 0:40] PHẦN 1: MỞ ĐẦU & NÊU NỖI ĐẦU THỰC TẾ CỦA HỌC SINH",
         "Kính thưa Ban Giám khảo và quý thầy cô, em tên là Khánh Sơn, cùng bạn Quốc An đại diện nhóm học sinh trường THPT Nguyễn Công Trứ đến với cuộc thi Khởi nghiệp HUIT 2026.\n\n"
         "Thưa quý vị, là học sinh Lớp 10 theo chương trình GDPT 2018 mới, chúng em bắt buộc phải học phần thực hành lắp ráp máy tính. Tuy nhiên, thực tế tại trường em cũng như hơn 85% các trường THPT hiện nay đều không có phòng lab phần cứng do chi phí trang bị 40 máy thật lên tới hơn 800 triệu đồng. Mỗi lần tháo lắp máy thật, tụi em rất sợ làm gãy chân socket CPU hay làm cháy linh kiện do tĩnh điện, nên cuối cùng hầu như học sinh chúng em chỉ được 'học chay' qua hình vẽ 2D khô khan trong sách.\n\n"
         "Chính vì thế, khi ra ngoài tự đi mua máy tính học tập, học sinh và phụ huynh tụi em hoàn toàn bị bất đồng thông tin và rất dễ bị các cửa hàng đôn giá linh kiện tồn kho."),

        ("[0:40 - 1:30] PHẦN 2: GIẢI PHÁP 3D & ĐỘT PHÁ CÔNG NGHỆ DO HỌC SINH TỰ LÀM",
         "Vì thế, nhóm chúng em đã quyết định tự tay xây dựng PC Master Builder - một phòng lab thực hành ảo 3D ngay trên Web, giúp các bạn học sinh thực hành hoàn toàn miễn phí mà không tốn 1 đồng thiết bị cứng.\n\n"
         "Điểm đặc biệt nhất mà nhóm em tự làm được là Công nghệ cử chỉ tay Hand-Tracking qua Webcam: Bạn học sinh chỉ cần giơ bàn tay trước camera máy tính là có thể co ngón tay để bốc, xoay 360 độ và gắn từng cây RAM, CPU, VGA vào bo mạch chủ như thực tế mà KHÔNG CẦN MUA KÍNH VR ĐẮT TIỀN.\n\n"
         "Không chỉ dừng lại ở đó, nhóm em còn tích hợp Trợ lý AI Chống Chặt Chém tự động tra cứu giá niêm yết từ các đại lý lớn và cảnh báo khi linh kiện bị đẩy giá hoặc không tương thích socket."),

        ("[1:30 - 2:30] PHẦN 3: SỐ LIỆU THỰC NGHỆM THẬT & PHỤNG SỰ 10 MỤC TIÊU SDGS",
         "Sản phẩm của chúng em không phải là mô hình lý thuyết suông. Tụi em đã đưa ứng dụng chạy thực tế tại pc-master-lms.vercel.app và cho 76 bạn học sinh lớp 10A1 và 10C9 trường em trải nghiệm trực tiếp trong tiết học.\n\n"
         "Kết quả thực tế cho thấy: Điểm thi thực hành của các bạn tăng trung bình 28%, tỷ lệ hoàn thành bài học đạt 94% và 92% các bạn phản hồi học dễ hiểu hơn hẳn phương pháp cũ.\n\n"
         "Đặc biệt hơn, dự án của nhóm học sinh chúng em đóng góp thiết thực vào 10 Mục tiêu Phát triển Bền vững (UN SDGs) của Liên Hợp Quốc: Nổi bật là SDG 4 đưa giáo dục 3D chất lượng đến mọi bạn học sinh; SDG 9 đổi mới công nghệ AI; SDG 12 triệt tiêu 100% rác thải điện tử E-waste do thực hành lỗi; SDG 8 định hướng nghề nghiệp Vi mạch - Bán dẫn; SDG 10 thu hẹp khoảng cách số cho các bạn vùng xa; cùng các mục tiêu SDG 13 bảo vệ khí hậu, SDG 5 bình đẳng giới STEM cho nữ sinh, SDG 3 an toàn điện, SDG 11 trường học thông minh và SDG 17 hợp tác nhà trường - đại lý."),

        ("[2:30 - 3:30] PHẦN 4: MÔ HÌNH THỰC TẾ & MỨC GIÁ 20K VỪA TÚI TIỀN HỌC SINH",
         "Về mô hình phát triển, là học sinh nên tụi em hiểu rất rõ tâm lý các bạn: Tụi em chọn mức giá gói Pro chỉ 20.000 VNĐ/tháng - đúng bằng giá 1 ly trà sữa bình dân để bạn nào cũng có thể tự đăng ký học. Ngoài ra, nhóm mở rộng gói B2B LMS cho nhà trường với giá 15-25 triệu/năm, giúp nhà trường tiết kiệm hàng trăm triệu tiền mua máy thật.\n\n"
         "Nhờ sử dụng hạ tầng Cloud Serverless tinh gọn, chi phí duy trì hệ thống chỉ tốn khoảng 2 triệu/tháng. Dự án đạt điểm huề vốn rất nhanh chỉ với 250 bạn học sinh trả phí và đạt biên lợi nhuận ròng trên 75%.\n\n"
         "Nếu may mắn giành 50 triệu tiền thưởng cuộc thi HUIT, tụi em sẽ dành 40% nâng cấp server AI, 40% đến giới thiệu sản phẩm cho 10 trường THPT bạn và 20% đăng ký bản quyền tác giả."),

        ("[3:30 - 4:20] PHẦN 5: NĂNG LỰC NHÓM HỌC SINH & ĐĂNG KÝ BẢN QUYỀN IP",
         "Nhờ sự hướng dẫn tận tình của Cô Kim Phượng và Thầy Minh Phụng, toàn bộ mã nguồn Web Next.js, mô hình 3D và thuật toán AI đều do 4 học sinh nhóm em tự tay lập trình 100% mà không tốn tiền thuê ngoài. Nhóm em cũng đã hoàn tất đăng ký Bản quyền Tác giả cho phần mềm để bảo hộ sản phẩm."),

        ("[4:20 - 5:00] PHẦN 6: LỜI KẾT CHẮC CHẮN ĐẠT ĐIỂM TỐI ĐA",
         "Chính vì tất cả những lý do trên, dự án PC Master LMS của nhóm học sinh chúng em tự tin xứng đáng giành vị trí Quán quân HUIT 2026: Bởi đây là một sản phẩm Make-in-HUIT tự chủ công nghệ 100%, giải đúng nỗi đau thực tế của học sinh chúng em bằng công nghệ 3D/AI hiện đại, phụng sự 10 mục tiêu phát triển bền vững và có tính khả thi nhân rộng rất cao!\n\n"
         "Em xin chân thành cảm ơn Ban Giám khảo và quý thầy cô đã lắng nghe!")
    ]

    for title, text in speech_sections:
        p_h = doc.add_paragraph()
        p_h.paragraph_format.space_before = Pt(14)
        p_h.paragraph_format.space_after = Pt(4)
        r_h = p_h.add_run(title)
        r_h.font.name = 'Arial'
        r_h.font.size = Pt(12)
        r_h.font.bold = True
        r_h.font.color.rgb = COLOR_NAVY

        # Speech Box
        tbl = doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = tbl.cell(0, 0)
        cell.width = Inches(6.8)
        shd = parse_xml(r'<w:shd {} w:fill="F8FAFC"/>'.format(nsdecls('w')))
        cell._tc.get_or_add_tcPr().append(shd)

        p_c = cell.paragraphs[0]
        p_c.paragraph_format.space_before = Pt(6)
        p_c.paragraph_format.space_after = Pt(6)

        r_c = p_c.add_run(text)
        r_c.font.name = 'Arial'
        r_c.font.size = Pt(10)
        r_c.font.color.rgb = COLOR_DARK

        doc.add_paragraph()

    filename = "Loi_Noi_Thuyet_Trinh_5_Phut_Hoc_Sinh_HUIT_2026.docx"
    doc.save(filename)
    print(f"Student speech created successfully: {filename}")

if __name__ == '__main__':
    create_student_speech_docx()
